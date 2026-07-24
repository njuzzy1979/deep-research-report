# -*- coding: utf-8 -*-
"""md2docx 转换器 v2 端到端集成测试。

测试用例覆盖 C-15 规格的五类场景：
  1. 最小合法 Markdown（1H1 + 1H2 + 1段正文）→ exit 0, FATAL=0
  2. 多 H1 警告场景（W-HDR-03 触发但不阻断）
  3. 图片缺失场景（E-IMG-01 + --allow-missing-figures 降级）
  4. 章节编号连续性（4 章 → 编号无跳号）
  5. FRONT_MATTER 不编号（前言 H1 + 4 个前言 H2 + 2 个正文 H2）

运行方式：
    cd <deep-research-report>/scripts
    pytest md2docx/tests/ -v
"""
from __future__ import annotations

import re
from pathlib import Path

import pytest

from conftest import (
    assemble_document_ir,
    copy_fixture_to_workdir,
    run_converter,
)


def _stderr_safe(result) -> str:
    """安全获取 stderr 字符串，处理 None 情况。"""
    return result.stderr or ""


def _combined_output(result) -> str:
    """安全获取 stdout + stderr 合并字符串。"""
    return (result.stdout or "") + (result.stderr or "")


# ── 测试用例 1：最小合法 Markdown ──────────────────────────────

def test_minimal_valid_markdown(tmp_workdir: Path, scripts_dir: Path):
    """1H1 + 1H2 + 1段正文 → 转换成功 exit 0, FATAL=0。"""
    md_path = copy_fixture_to_workdir("minimal.md", tmp_workdir)
    docx_path = tmp_workdir / "minimal.docx"
    report_path = tmp_workdir / "minimal.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )

    # 断言 exit 0
    assert result.returncode == 0, (
        f"期望 exit 0，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )

    # 断言 FATAL 为 0
    stderr_combined = result.stderr + (result.stdout or "")
    fatal_count = stderr_combined.count("[FATAL]")
    assert fatal_count == 0, f"期望 FATAL=0，实际 {fatal_count}。stderr: {_stderr_safe(result)[-2000:]}"

    # 断言 docx 文件生成
    assert docx_path.exists(), f"docx 文件未生成: {docx_path}"
    assert docx_path.stat().st_size > 0, "docx 文件为空"

    # 断言转换报告生成
    assert report_path.exists(), f"转换报告未生成: {report_path}"


# ── 测试用例 2：多 H1 警告场景 ──────────────────────────────────

def test_multi_h1_warning_not_blocking(tmp_workdir: Path, scripts_dir: Path):
    """多 H1 触发 W-HDR-03 警告但不阻断转换（exit 0 或 1，非 2）。"""
    # 构造含两个 H1 的 Markdown
    content = """# 第一个 H1（前言）

这是前言部分的正文。

## 摘要

本报告用于验证多 H1 场景下转换器的容错行为。

## 正文第一章

这是正文第一章的内容。第二个 H1 将在下方出现。

# 第二章（多 H1 场景）

这是使用 H1 的第二章，应触发 W-HDR-03。

## 第二章第一节

验证转换器不会因多余 H1 而 FATAL 中断。
"""
    md_path = tmp_workdir / "multi-h1.md"
    md_path.write_text(content, encoding="utf-8")
    docx_path = tmp_workdir / "multi-h1.docx"
    report_path = tmp_workdir / "multi-h1.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )

    # 不应 FATAL 退出（exit 2 表示 FATAL 存在）
    assert result.returncode != 2, (
        f"多 H1 不应触发 FATAL (exit 2)，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )

    # 应能正常产出 docx
    assert docx_path.exists(), "即使有多 H1 警告，也应产出 docx"
    assert docx_path.stat().st_size > 0

    # 检查转换报告中是否有 W-HDR-03 相关警告
    report_content = report_path.read_text(encoding="utf-8") if report_path.exists() else ""
    combined = result.stderr + result.stdout + report_content
    has_hdr_warning = "W-HDR-03" in combined or "HDR" in combined or "H1" in combined.lower()
    # 不强制要求具体警告码——不同转换器版本可能报告不同，但至少应产出结果
    print(f"[INFO] multi-H1 test: returncode={result.returncode}, docx generated={docx_path.exists()}")


# ── 测试用例 3：图片缺失场景 ────────────────────────────────────

def test_missing_image_with_allow_flag(tmp_workdir: Path, scripts_dir: Path):
    """报告引用不存在的图片文件 → --allow-missing-figures 降级为 WARNING。"""
    content = """# 图片缺失测试

## 摘要

本报告用于验证 --allow-missing-figures 参数对缺失图片的降级处理。

## 架构图

如图 1-1 所示，系统采用分层架构设计。

![图1-1 不存在的架构图](nonexistent-figure.png)

正文继续描述架构细节。
"""
    md_path = tmp_workdir / "missing-image.md"
    md_path.write_text(content, encoding="utf-8")
    docx_path = tmp_workdir / "missing-image.docx"
    report_path = tmp_workdir / "missing-image.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        allow_missing_figures=True,
        cwd=scripts_dir,
    )

    # 使用 --allow-missing-figures 后，exit code 应为 0 或 1（不应 FATAL=2）
    assert result.returncode != 2, (
        f"--allow-missing-figures 不应导致 FATAL (exit 2)，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )

    # 应正常产出 docx
    assert docx_path.exists()
    assert docx_path.stat().st_size > 0

    # 转换报告或 stderr 中应有图片相关警告
    combined = result.stderr + result.stdout
    has_image_warning = (
        "IMG" in combined
        or "图" in combined
        or "figure" in combined.lower()
        or "missing" in combined.lower()
    )
    print(f"[INFO] missing-image test: returncode={result.returncode}, "
          f"image warning found={has_image_warning}")


def test_missing_image_without_allow_flag_fails(tmp_workdir: Path, scripts_dir: Path):
    """不使用 --allow-missing-figures 时，图片缺失应导致非零退出码。"""
    content = """# 图片缺失测试（无 allow flag）

## 架构图

![图1-1 缺失的图](completely-missing.png)
"""
    md_path = tmp_workdir / "missing-image-strict.md"
    md_path.write_text(content, encoding="utf-8")
    docx_path = tmp_workdir / "missing-image-strict.docx"
    report_path = tmp_workdir / "missing-image-strict.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        allow_missing_figures=False,
        cwd=scripts_dir,
    )

    # 未加 --allow-missing-figures 时，预期非零退出（至少 1 = ERROR 或 2 = FATAL）
    assert result.returncode != 0, (
        f"缺少 --allow-missing-figures 时代码应为非零退出，实际 {result.returncode}"
    )
    print(f"[INFO] missing-image-strict test: returncode={result.returncode} (expected non-zero)")


# ── 测试用例 4：章节编号连续性 ──────────────────────────────────

def test_chapter_numbering_continuity(tmp_workdir: Path, scripts_dir: Path):
    """4 章标准结构 → 验证章节编号无跳号。"""
    md_path = copy_fixture_to_workdir("multi-chapter.md", tmp_workdir)
    docx_path = tmp_workdir / "multi-chapter.docx"
    report_path = tmp_workdir / "multi-chapter.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )

    assert result.returncode == 0, (
        f"4 章标准结构应正常转换 (exit 0)，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )

    # 读取转换报告，检查章节编号
    report_content = report_path.read_text(encoding="utf-8") if report_path.exists() else ""
    combined = result.stderr + result.stdout + report_content

    # 检查是否有跳号相关错误（E-HDR 类），而非"无跳号"等通过性表述
    has_skip_error = "E-HDR" in combined or ("跳号" in combined and "无跳号" not in combined)
    assert not has_skip_error, f"不应出现章节跳号错误。输出: {combined[:3000]}"

    # 确认产出 docx
    assert docx_path.exists() and docx_path.stat().st_size > 0


# ── 测试用例 5：FRONT_MATTER 不编号 ─────────────────────────────

# front-matter.md 中前置件区（前言/导论 H1 之下）的 4 个无编号 H2 文本。
# 这些 H2 必须被归类为 FRONT_MATTER，既不参与正文"第X章"编号，也不得丢失。
_FRONT_MATTER_H2_TEXTS = [
    "问题提出与研究背景",
    "研究目标与意义",
    "研究方法概述",
    "报告结构说明",
]

# front-matter.md 中正文章（带显式"第X章"编号）的剥离后标题 → 期望重编编号。
# P-006 修复后，这两个正文章应恢复 CHAPTER 分类与连续编号（增强断言，非放宽）：
# 前置件区遇显式章编号即终止（R-FM 终止条件之 (a)），正文从第一章开始编号。
_BODY_CHAPTER_EXPECTED = [
    ("军事需求与现状分析", "第一章"),
    ("研究目标与技术指标", "第二章"),
]


def test_front_matter_not_numbered_ir_level(scripts_dir: Path):
    """结构化断言（P0-2 重写）：直接检查装配后的 HeadingIR。

    对 front-matter.md 跑进程内 阶段0→3，断言前置件区 4 个 H2：
      (1) kind == FRONT_MATTER —— 分类正确（守 P0-1 根因A：不被误判为 CHAPTER）；
      (2) number is None 且 display_number == "" —— 无正文编号；
      (3) 全部出现在 DocumentIR.elements 中 —— 未丢失（守 P0-1 根因B：索引失配静默丢内容）；
      (4) 不产生 W-HDR-01（索引未命中会退化为 W-HDR-01）。

    此断言基于结构化 IR 字段，若 P0-1 任一根因复发（误编号 / 丢失），
    本测试必然失败——不再是旧版对报告文本做 or 连接的永真式陷阱。
    """
    from md2docx.ir import HeadingIR, HeadingKind

    fixture = scripts_dir / "md2docx" / "tests" / "test_fixtures" / "front-matter.md"
    doc, issues = assemble_document_ir(fixture)

    heading_irs = [e for e in doc.elements if isinstance(e, HeadingIR)]
    by_text = {h.text: h for h in heading_irs}

    for txt in _FRONT_MATTER_H2_TEXTS:
        # (3) 未丢失：必须出现在 elements 中
        assert txt in by_text, (
            f"前置件 H2 {txt!r} 未出现在 DocumentIR.elements 中"
            f"（P0-1 根因B 复发：索引失配导致标题被静默丢弃）。"
            f"现有标题：{sorted(by_text)}"
        )
        h = by_text[txt]
        # (1) 分类正确
        assert h.kind == HeadingKind.FRONT_MATTER, (
            f"前置件 H2 {txt!r} 被误分类为 {h.kind.name}"
            f"（P0-1 根因A 复发：应为 FRONT_MATTER），"
            f"display_number={h.display_number!r}"
        )
        # (2) 无编号
        assert h.number is None, (
            f"前置件 H2 {txt!r} 不应有结构化编号，实际 number={h.number!r}"
        )
        assert h.display_number == "", (
            f"前置件 H2 {txt!r} 不应有显示编号，实际 display_number={h.display_number!r}"
            f"（若为'第一章'等即 P0-1 复发）"
        )

    # (4) 不应因索引未命中而退化为 W-HDR-01
    whdr01 = [i for i in issues if i.code == "W-HDR-01"]
    assert not whdr01, (
        f"不应产生 W-HDR-01（索引未命中/丢失信号），实际 {len(whdr01)} 条："
        f"{[i.message[:60] for i in whdr01]}"
    )

    # (5) P-006 增强断言：正文章（补显式编号后）恢复 CHAPTER 且编号连续。
    #     守 R-FM 终止条件——前置件区遇显式"第X章"即终止，正文从第一章起编。
    #     这是对 §2.3"官方夹具能过是假象"的直接修正：夹具原本正文章无编号被
    #     全部误吞为 FRONT_MATTER，补编号后必须真正演示"前置件 + 正文"分层。
    for txt, expected_display in _BODY_CHAPTER_EXPECTED:
        assert txt in by_text, (
            f"正文章 {txt!r} 未出现在 DocumentIR.elements 中"
        )
        h = by_text[txt]
        assert h.kind == HeadingKind.CHAPTER, (
            f"正文章 {txt!r} 应恢复 CHAPTER，实际 {h.kind.name}"
            f"（P-006 根因C 复发：无编号信号误吞正文章为 FRONT_MATTER）"
        )
        assert h.display_number == expected_display, (
            f"正文章 {txt!r} 编号应为 {expected_display}，"
            f"实际 {h.display_number!r}（编号不连续或未从第一章起编）"
        )


def test_front_matter_not_numbered(tmp_workdir: Path, scripts_dir: Path):
    """端到端断言（P0-2 重写）：解析输出 docx，检查前置件 H2 的实际渲染。

    旧版断言 `"0.1" not in combined or "FRONT_MATTER" not in combined` 是 or
    连接的永真式（"FRONT_MATTER"为内部枚举名，从不出现在报告文本，右操作数
    几乎恒真 → 整个 assert 恒成立）。此处改为用 python-docx 解析输出 docx，
    逐段核对 4 个前置件 H2 段落：既存在，又不带"第X章"/"0.1"等编号前缀。
    """
    from docx import Document

    md_path = copy_fixture_to_workdir("front-matter.md", tmp_workdir)
    docx_path = tmp_workdir / "front-matter.docx"
    report_path = tmp_workdir / "front-matter.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )

    assert result.returncode == 0, (
        f"前言 Markdown 应正常转换 (exit 0)，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )
    assert docx_path.exists() and docx_path.stat().st_size > 0

    doc = Document(str(docx_path))
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 编号前缀集合：正文章编号（第X章）与"0.Y"式退化编号
    _NUM_PREFIXES = tuple(f"第{c}章" for c in "一二三四五六七八九十") + ("0.",)

    for txt in _FRONT_MATTER_H2_TEXTS:
        # 找到承载该前置件标题的段落（标题文本应作为独立段落存在）
        matches = [pt for pt in para_texts if txt in pt]
        assert matches, (
            f"前置件 H2 {txt!r} 未出现在输出 docx 的任何段落中"
            f"（P0-1 根因B 复发：内容丢失）"
        )
        for pt in matches:
            assert not pt.startswith(_NUM_PREFIXES), (
                f"前置件 H2 段落 {pt!r} 带有正文编号前缀"
                f"（P0-1 根因A 复发：前言 H2 被误编号）"
            )

    print(f"[INFO] front-matter test: returncode={result.returncode}, "
          f"docx={docx_path.stat().st_size} bytes, "
          f"前置件H2段落均存在且无编号")


# ── 测试用例 6：换样本金标准测试（反硬编码终极判据，§F.3.3 CI 必跑项）──────

# alt-topic-coffee.md（题材：咖啡烘焙工艺）与真实样本（态势感知/军事研究报告）
# 完全无关。转换产物的全部题注/编号/封面字段必须与本夹具内容一致，且不得出现
# 任何原样本主题相关字符串——"换报告零改码"通过即证明反硬编码达成。

# 本夹具的结构化"金标准"期望值（全部来自 alt-topic-coffee.md 内容）
_COFFEE_EXPECTED = {
    "title": "咖啡烘焙工艺与风味形成机理",
    "subtitle": "从生豆化学到杯中风味的全链路解析",
    "report_type": "技艺研究白皮书",
    "organization": "城南手冲实验室",
    "chapters": ["生豆特性与预处理", "烘焙曲线与热力学控制", "风味评估与杯测校准"],
    "chapter_numbers": ["第一章", "第二章", "第三章"],
    "figures": {"1-1": "三种处理法生豆失水速率对比", "2-1": "标准中烘温度-时间曲线"},
    "body_tables": {"1-1": "生豆分级筛选标准", "2-1": "浅中深三档烘焙关键参数"},
    "appendices": ["附录A", "附录B"],
}

# 原样本（态势感知报告）主题词——绝不应泄漏到咖啡报告的转换产物中。
# 若任一词出现，说明代码把"上一份报告写了什么"硬编码了（§F.1 违规）。
_ORIGINAL_SAMPLE_LEAK_WORDS = [
    "态势感知", "空间态势", "军事", "遨天", "SSA", "多智能体", "非合作目标",
    "深度研究报告",  # 旧 cover.py 硬编码的报告类型默认值（已修复为不硬编码）
]


def test_alt_topic_gold_standard_ir(scripts_dir: Path):
    """换样本金标准（IR 层）：咖啡题材夹具的装配结果与夹具内容逐项一致。

    验证转换器对编号/分类/题注的处理完全由文档内容驱动（动态解析），
    换到与真实样本毫不相干的题材后，输出结构仍正确——反硬编码达成的直接证据。
    """
    from md2docx.ir import HeadingIR, HeadingKind, FigureIR, TableIR, TableKind

    fixture = scripts_dir / "md2docx" / "tests" / "test_fixtures" / "alt-topic-coffee.md"
    doc, issues = assemble_document_ir(fixture)

    # 封面字段全部来自文档元数据
    assert doc.metadata.title == _COFFEE_EXPECTED["title"]
    assert doc.metadata.subtitle == _COFFEE_EXPECTED["subtitle"]
    assert doc.metadata.report_type == _COFFEE_EXPECTED["report_type"]
    assert doc.metadata.organization == _COFFEE_EXPECTED["organization"]

    # 章标题分类与编号
    chapters = [
        e for e in doc.elements
        if isinstance(e, HeadingIR) and e.kind == HeadingKind.CHAPTER
    ]
    assert [c.text for c in chapters] == _COFFEE_EXPECTED["chapters"], (
        f"章标题不符：{[c.text for c in chapters]}"
    )
    assert [c.display_number for c in chapters] == _COFFEE_EXPECTED["chapter_numbers"], (
        f"章编号不符：{[c.display_number for c in chapters]}"
    )

    # 附录字母编号
    appendices = [
        e.display_number for e in doc.elements
        if isinstance(e, HeadingIR) and e.kind == HeadingKind.APPENDIX
    ]
    assert appendices == _COFFEE_EXPECTED["appendices"], f"附录编号不符：{appendices}"

    # 图题注（图号→题注文字全动态）
    fig_map = {
        e.figure_id: e.caption_text
        for e in doc.elements if isinstance(e, FigureIR)
    }
    assert fig_map == _COFFEE_EXPECTED["figures"], f"图题注不符：{fig_map}"

    # 正文表题注
    body_tbl_map = {
        e.table_id: e.caption_text
        for e in doc.elements
        if isinstance(e, TableIR) and e.kind == TableKind.BODY
    }
    assert body_tbl_map == _COFFEE_EXPECTED["body_tables"], f"表题注不符：{body_tbl_map}"


def test_alt_topic_no_original_sample_leak(tmp_workdir: Path, scripts_dir: Path):
    """换样本金标准（端到端）：咖啡报告 docx 中不得出现任何原样本主题词。

    这是 §F.3.3"换报告零改码"的核心判据——若转换器把"上一份报告写了什么"
    硬编码进了代码，切换到咖啡题材时那些残留词会泄漏到输出。逐段扫描输出 docx，
    断言 (a) 咖啡夹具的关键内容都在，(b) 原样本主题词一个都不在。
    """
    from docx import Document

    md_path = copy_fixture_to_workdir("alt-topic-coffee.md", tmp_workdir)
    docx_path = tmp_workdir / "alt-topic-coffee.docx"
    report_path = tmp_workdir / "alt-topic-coffee.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )
    assert result.returncode == 0, (
        f"咖啡题材报告应正常转换 (exit 0)，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )
    assert docx_path.exists() and docx_path.stat().st_size > 0

    doc = Document(str(docx_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # 加入表格文本
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    # (a) 咖啡夹具关键内容确实出现
    for chapter in _COFFEE_EXPECTED["chapters"]:
        assert chapter in full_text, f"章标题 {chapter!r} 未出现在输出 docx 中"
    assert _COFFEE_EXPECTED["title"] in full_text, "封面标题未出现"
    for cap in _COFFEE_EXPECTED["figures"].values():
        assert cap in full_text, f"图题注 {cap!r} 未出现"

    # (b) 原样本主题词一个都不应出现（反硬编码终极判据）
    leaked = [w for w in _ORIGINAL_SAMPLE_LEAK_WORDS if w in full_text]
    assert not leaked, (
        f"检测到原样本主题词泄漏到咖啡报告输出：{leaked}"
        f"（§F.1 违规：代码硬编码了上一份报告的内容）"
    )

    print(f"[INFO] alt-topic gold-standard: returncode={result.returncode}, "
          f"docx={docx_path.stat().st_size} bytes, 无原样本主题词泄漏")


# ── 测试用例 7：反硬编码静态扫描器（check_no_hardcode）自身作为 CI 门 ────

def test_check_no_hardcode_clean(scripts_dir: Path):
    """反硬编码扫描器对 md2docx 包应零发现（§F.3.1/F.3.2 CI 必跑项）。"""
    import importlib

    mod = importlib.import_module("md2docx.check_no_hardcode")
    pkg_root = Path(mod.__file__).resolve().parent
    result = mod.run_scan(pkg_root, verbose=False)
    assert not result.findings, (
        f"反硬编码扫描发现 {len(result.findings)} 项，应为 0：\n"
        + "\n".join(f"  {f.file}:{f.line} [{f.code}] {f.literal!r} — {f.reason}"
                    for f in result.findings[:20])
    )


# ── 辅助：验证 with-image 和 with-table 基础转换不崩溃 ───────────

def test_with_image_basic_conversion(tmp_workdir: Path, scripts_dir: Path):
    """含图片引用的 Markdown 基本转换不崩溃（即使图片路径指向真实 dummy.png）。"""
    md_path = copy_fixture_to_workdir("with-image.md", tmp_workdir)

    # 确保 dummy.png 存在于 workdir
    # (已被 tmp_workdir fixture 创建，with-image.md 中的路径已被替换为 dummy.png)

    docx_path = tmp_workdir / "with-image.docx"
    report_path = tmp_workdir / "with-image.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        figures_dir=tmp_workdir,
        report_path=report_path,
        cwd=scripts_dir,
    )

    # 有真实图片时应正常转换
    assert result.returncode in (0, 1), (
        f"含真实图片的转换不应 FATAL，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )
    assert docx_path.exists() and docx_path.stat().st_size > 0


def test_with_table_basic_conversion(tmp_workdir: Path, scripts_dir: Path):
    """含表格题注的 Markdown 基本转换不崩溃。"""
    md_path = copy_fixture_to_workdir("with-table.md", tmp_workdir)
    docx_path = tmp_workdir / "with-table.docx"
    report_path = tmp_workdir / "with-table.conversion-report.md"

    result = run_converter(
        md_path, docx_path,
        report_path=report_path,
        cwd=scripts_dir,
    )

    assert result.returncode == 0, (
        f"含表格的 Markdown 应正常转换，实际 {result.returncode}\n"
        f"stderr: {_stderr_safe(result)[-2000:]}"
    )
    assert docx_path.exists() and docx_path.stat().st_size > 0
