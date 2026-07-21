"""表格渲染模块（C-11）：将 TableIR 渲染为符合 V3.0 §5.1 规范的 Word 表格。

全框线、交替底纹、表头加粗重复、宽度 90% 居中。表题注在表上方，
来源行在表下方。正文表（BODY）有书签，附录表（APPENDIX）无书签。

本模块由 render/document.py（C-07c，未来任务）调用——document 遍历
DocumentIR.elements，对遇到的每个 TableIR 元素委托给本模块渲染。
"""
from __future__ import annotations

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..config import (
    COLOR_TABLE_ALT_SHADE,
    TABLE_BORDERS,
    BehaviorFlags,
)
from ..ir import InlineRun, TableIR, TableKind
from ..issues import IssueCollector
from .oxml_helpers import (
    make_bookmark_end,
    make_bookmark_start,
    make_shd,
    make_tblBorders,
    make_tcBorders_bottom,
    set_eastAsia_font,
)

# 表头与表体的硬编码字体名（不经过 config.StyleSpec 通路，因为单元格内的 run
# 级字体设置是段落级样式无法完全覆盖的——python-docx 对 eastAsia 字体不会
# 从段落样式自动继承到 run，需在 run 级显式设置）。
_HEADER_CJK_FONT = "微软雅黑"  # V3.0 §5.1
_HEADER_LATIN_FONT = "Times New Roman"
_BODY_CJK_FONT = "宋体"  # V3.0 §5.1（历史问题13：表体字号 10.5pt，非 12pt）
_BODY_LATIN_FONT = "Times New Roman"

# 来源行前缀（M4）
_SOURCE_PREFIX = "数据来源："

# 表格宽度比率：pct50 单位下 100% = 5000，90% = 4500
_TABLE_WIDTH_PCT50 = 4500

# 书签名前缀（正文表 "TabX_Y"，04-interface-spec.md §6）
_BOOKMARK_PREFIX = "Tab"


# ===========================================================================
# 内部辅助函数
# ===========================================================================


def _render_cell_runs(
    cell,
    runs: list[InlineRun],
    cjk_font: str,
    latin_font: str,
    size_pt: float,
    default_bold: bool,
) -> None:
    """将 InlineRun 列表渲染到表格单元格的默认段落中。

    先清空默认段落（python-docx 创建的空白段落会残留一个空 run），
    然后为每个 InlineRun 创建一个 w:r 元素，应用格式属性。

    Args:
        cell: python-docx Cell 对象
        runs: 该单元格的 InlineRun 列表
        cjk_font: 中文字体名
        latin_font: 西文字体名
        size_pt: 字号（pt）
        default_bold: 默认是否加粗（表头 True，表体 False）
    """
    p = cell.paragraphs[0]
    p.clear()

    for run_data in runs:
        r = p.add_run(run_data.text)
        # 西文字体
        r.font.name = latin_font
        # 中文字体（必须 run 级显式设置，python-docx 不会从段落样式继承 eastAsia）
        rPr = r._r.get_or_add_rPr()
        set_eastAsia_font(rPr, cjk_font)
        # 字号
        r.font.size = Pt(size_pt)
        # 字重：InlineRun.bold 为 True 则加粗，否则用默认
        r.font.bold = run_data.bold or default_bold
        # 斜体
        if run_data.italic:
            r.font.italic = True
        # 行内代码：西文字体切换为 Consolas（V3.0 §10.3；eastAsia 仍用原字体）
        if run_data.code:
            r.font.name = "Consolas"
        # 上标
        if run_data.superscript:
            r.font.superscript = True
        # 超链接：通过 OXML 手写（python-docx 的段落级 add_hyperlink 无法
        # 在已有段落的 run 序列中按位置插入）
        if run_data.link_url:
            _replace_run_with_hyperlink(p, r, run_data)


def _replace_run_with_hyperlink(
    paragraph,
    run,
    run_data: InlineRun,
) -> None:
    """把段落中最后一个 run 替换为超链接 w:hyperlink 元素。

    python-docx 不支持在 run 序列中间插入超链接，因此采用"先添加普通 run
    再将其替换为 hyperlink"的策略：移除最后一个 w:r，在同一位置插入
    w:hyperlink，其中包含一个等效的 w:r（含格式属性 + w:t 文本）。

    Args:
        paragraph: python-docx Paragraph 对象
        run: 刚添加的普通 run（将被替换）
        run_data: 原始 InlineRun（含 link_url）
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    p_elem = paragraph._p
    run_elem = run._r

    # 为超链接目标建立关系
    r_id = paragraph.part.relate_to(
        run_data.link_url, RT.HYPERLINK, is_external=True
    )

    # 构建 w:hyperlink 元素
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # 在 hyperlink 内部重建一个等效的 w:r（复制原 run 的属性）
    new_r = OxmlElement("w:r")
    # 复制 rPr（如果原 run 有的话）
    orig_rPr = run_elem.find(qn("w:rPr"))
    if orig_rPr is not None:
        # 深拷贝一份 rPr，避免 lxml 引用问题
        from copy import deepcopy

        new_rPr = deepcopy(orig_rPr)
        new_r.append(new_rPr)

    # 添加 w:t 文本
    t = OxmlElement("w:t")
    t.text = run_data.text
    t.set(qn("xml:space"), "preserve")
    new_r.append(t)

    hyperlink.append(new_r)

    # 替换：用 hyperlink 替换原 run_elem
    p_elem.replace(run_elem, hyperlink)


def _set_table_width_pct50(tblPr) -> None:
    """设置表格宽度为页面 90%（pct50 单位）。

    在 tblPr 子元素最前面插入 w:tblW，避免影响已有属性。
    pct50 单位下：5000 = 100%，4500 = 90%。
    """
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(_TABLE_WIDTH_PCT50))
    tblW.set(qn("w:type"), "pct50")
    tblPr.insert(0, tblW)


def _set_tblHeader_repeat(row) -> None:
    """设置表头行为 tblHeader，使长表跨页时表头自动重复。

    OOXML 中 <w:tblHeader/> 为布尔属性元素，存在即表示 true。
    同时设置 w:val="true" 作为显式兼容写法（某些旧版 Word 仅识别
    带 val 属性的形式——这是防御性双写，非规范严格要求）。
    """
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _apply_cell_shading(cell, fill_hex_no_hash: str) -> None:
    """对单元格应用底纹（交替行灰底）。

    Args:
        cell: python-docx Cell 对象
        fill_hex_no_hash: 填充色（无 # 号），如 "F2F2F2"
    """
    tcPr = cell._tc.get_or_add_tcPr()
    make_shd(tcPr, fill_hex_no_hash)


def _apply_header_bottom_border(cell, sz: int, color_hex_no_hash: str) -> None:
    """对表头单元格设置底部 1pt 分隔线，覆盖表级 insideH 的 0.5pt。

    Args:
        cell: python-docx Cell 对象
        sz: 线宽（八分之一磅单位），8 = 1pt
        color_hex_no_hash: 颜色（无 # 号），如 "000000"
    """
    tcPr = cell._tc.get_or_add_tcPr()
    make_tcBorders_bottom(tcPr, sz=sz, color=color_hex_no_hash)


def _hex_strip(hex_with_hash: str) -> str:
    """去掉颜色十六进制字符串的 # 前缀（OXML 的 w:color/w:fill 不接受 # 号）。"""
    return hex_with_hash.lstrip("#")


# ===========================================================================
# 公开 API
# ===========================================================================


def render_table(
    doc: Document,
    table_ir: TableIR,
    styles: dict,
    issues: IssueCollector,
    flags: BehaviorFlags,
    bookmark_id_start: int = 0,
) -> int:
    """渲染表格，返回消耗的 bookmark 数量（供调用方全局 ID 递增）。

    渲染流程（V3.0 §5.1 / 04-interface-spec.md §2.6）：
        1. 表题注（表上方，Caption Table 样式）——仅正文表
        2. 表格主体（居中 90% 宽、全框线、tblHeader）
        3. 表头行（Table Header 样式、加粗、底部 1pt 分隔线）
        4. 数据行（Table Body 样式、交替行底纹 #F2F2F2）
        5. 来源行（表下方，Table Source 样式）——可选

    Args:
        doc: python-docx Document 对象
        table_ir: 表格中间表示（来自 assemble/* 产出）
        styles: 样式名→样式对象字典（render/styles.py build_styles() 产出）
        issues: IssueCollector（本函数不产生 Issue，保留参数以统一签名）
        flags: BehaviorFlags（影响 table_first_col_left_align 等开关）
        bookmark_id_start: 起始书签 ID；正文表若含书签则使用此值，
            附录表不消费

    Returns:
        int: 本表格消耗的书签 ID 数量（0 或 1）；调用方应将返回值加到
            自身的 bookmark_id 计数器上
    """
    is_body = table_ir.kind == TableKind.BODY
    bookmark_consumed = 0
    border_color = _hex_strip(TABLE_BORDERS.color_hex)
    shade_color = _hex_strip(COLOR_TABLE_ALT_SHADE)

    # ------------------------------------------------------------------
    # 1. 表题注（表上方）
    # ------------------------------------------------------------------
    if table_ir.caption_text is not None:
        caption_para = doc.add_paragraph()
        caption_style = styles.get("Caption Table")
        if caption_style is not None:
            caption_para.style = caption_style

        # 正文表：发射书签（bookmarkStart/End 对，锚定在题注段落上）
        if is_body and table_ir.bookmark_name is not None:
            bm_id = bookmark_id_start
            make_bookmark_start(caption_para, bm_id, table_ir.bookmark_name)
            make_bookmark_end(caption_para, bm_id, table_ir.bookmark_name)
            bookmark_consumed = 1

        # 题注文本格式："表X-Y 题注内容"
        full_caption = f"表{table_ir.table_id} {table_ir.caption_text}"
        caption_para.add_run(full_caption)

    # ------------------------------------------------------------------
    # 2. 表格主体
    # ------------------------------------------------------------------
    n_rows = 1 + len(table_ir.body_rows)
    n_cols = table_ir.n_cols
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # 表格居中（M4）
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 表格宽度 = 页面 90%
    tblPr = table._tbl.tblPr
    _set_table_width_pct50(tblPr)

    # 全框线（顶/底 1.5pt、竖线 0.5pt、内部横线 0.5pt）
    make_tblBorders(tblPr, TABLE_BORDERS)

    # ------------------------------------------------------------------
    # 3. 表头行
    # ------------------------------------------------------------------
    header_style = styles.get("Table Header")
    for col_idx, header_runs in enumerate(table_ir.header_cells):
        cell = table.cell(0, col_idx)

        # 单元格垂直居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 应用段落样式
        if header_style is not None:
            cell.paragraphs[0].style = header_style

        # 渲染表头单元格文本（含加粗等格式属性）
        _render_cell_runs(
            cell,
            header_runs,
            cjk_font=_HEADER_CJK_FONT,
            latin_font=_HEADER_LATIN_FONT,
            size_pt=10.0,
            default_bold=True,
        )

        # 表头下 1pt 分隔线（单元格级 tcBorders 覆盖表级 insideH 0.5pt，M5）
        _apply_header_bottom_border(
            cell,
            sz=TABLE_BORDERS.header_bottom_sz,
            color_hex_no_hash=border_color,
        )

    # 设置表头行为 tblHeader（长表跨页自动重复表头）
    _set_tblHeader_repeat(table.rows[0])

    # ------------------------------------------------------------------
    # 4. 数据行
    # ------------------------------------------------------------------
    body_style = styles.get("Table Body")
    first_col_left = flags.table_first_col_left_align  # M4 开关（默认 False）

    for row_idx, row_data in enumerate(table_ir.body_rows):
        for col_idx, cell_runs in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)

            # 单元格垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # 应用段落样式
            if body_style is not None:
                cell.paragraphs[0].style = body_style

            # V3.0 字面全居中（含首列，R10），但 BehaviorFlags 开关允许
            # 首列左对齐（table_first_col_left_align=True 时）。
            if first_col_left and col_idx == 0:
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 渲染数据单元格文本
            _render_cell_runs(
                cell,
                cell_runs,
                cjk_font=_BODY_CJK_FONT,
                latin_font=_BODY_LATIN_FONT,
                size_pt=10.5,
                default_bold=False,
            )

            # 交替行底纹（M4：偶数数据行（row_idx 为奇数）加灰底 #F2F2F2）
            if row_idx % 2 == 1:
                _apply_cell_shading(cell, shade_color)

    # ------------------------------------------------------------------
    # 5. 来源行（表下方）
    # ------------------------------------------------------------------
    if table_ir.source_note is not None and len(table_ir.source_note) > 0:
        source_para = doc.add_paragraph()
        source_style = styles.get("Table Source")
        if source_style is not None:
            source_para.style = source_style

        # 前缀 "数据来源："（M4：Table Source 样式已定义 9pt italic #555555）
        source_para.add_run(_SOURCE_PREFIX)

        # 逐条渲染 source_note 中的 InlineRun（可能含超链接等格式）
        for run_data in table_ir.source_note:
            r = source_para.add_run(run_data.text)
            if run_data.link_url:
                _replace_run_with_hyperlink(source_para, r, run_data)

    return bookmark_consumed
