"""封面渲染模块（C-07c）：生成符合 V3.1 §7 规范的研究报告封面。

不含密级字段（🚫严禁标密），只含五要素+分隔线。
布局（自上而下）：顶部留白 6cm → 标题(28pt Bold) → 副标题(14pt,可选)
→ 报告类型(16pt) → 分隔线(1pt 5cm宽居中) → 机构名(14pt Bold) → 版本+日期(11pt)

封面各段落以 Normal 样式为基础，通过 run 级别属性覆盖字号/加粗/颜色等。
分隔线使用 oxml_helpers.make_pBdr_bottom() 在空段落底部生成边框线。
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..config import (
    COLOR_BLACK,
    COVER_ELEMENTS,
    COVER_SEPARATOR,
    COVER_SEPARATOR_SPACING_PT,
    COVER_TOP_SPACING_CM,
    MARGIN_LEFT_CM,
    MARGIN_RIGHT_CM,
    PAGE_WIDTH_CM,
)
from ..ir import MetadataIR
from .oxml_helpers import make_pBdr_bottom


def _add_cover_paragraph(
    doc: Document,
    text: str,
    size_pt: float,
    bold: bool,
    color_hex: str,
    space_before_pt: float,
    space_after_pt: float,
    line_spacing: float | None = None,
) -> None:
    """添加封面段落：居中、指定字号/字重/颜色，段前段后间距。

    封面元素使用微软雅黑作为中文字体（与 Heading 系列一致），西文使用
    python-docx 默认字体（Calibri），由 python-docx 自动处理字体回退。

    Args:
        doc: python-docx Document 对象
        text: 段落文本（单行）
        size_pt: 字号（磅）
        bold: 是否加粗
        color_hex: 颜色十六进制（含 # 号，如 "#000000"）
        space_before_pt: 段前间距（磅）
        space_after_pt: 段后间距（磅）
        line_spacing: 行距倍数（None 表示不设置，使用默认单倍行距）
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing

    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex.lstrip("#"))

    # 中文字体（封面元素统一使用微软雅黑）
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.insert(0, rFonts)


def _add_separator(doc: Document) -> None:
    """添加封面分隔线：1pt 黑色横线，约 5cm 宽，居中。

    通过空段落 + pBdr bottom 边框实现横线效果。
    段落宽度通过左右缩进控制，使其在页面上居中且约 5cm 宽。
    分隔线与上下元素各间距 12pt。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(COVER_SEPARATOR_SPACING_PT)
    pf.space_after = Pt(COVER_SEPARATOR_SPACING_PT)

    # 计算左右缩进：使段落宽度（即边框线宽度）≈ 5cm 且居中
    usable_width = PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM
    indent = (usable_width - COVER_SEPARATOR.width_cm) / 2
    pf.left_indent = Cm(indent)
    pf.right_indent = Cm(indent)

    # 空段落底部置 1pt 黑色边框（段落无文字，只呈现横线效果）
    pPr = p._p.get_or_add_pPr()
    make_pBdr_bottom(pPr, sz=COVER_SEPARATOR.sz, color=COVER_SEPARATOR.color_hex, space=1)


def render_cover(doc: Document, metadata: MetadataIR, styles: dict,
                 cover_metadata: dict | None = None) -> None:
    """在 Document 的当前节中渲染封面。

    封面布局（自上而下）：
        1. 顶部留白 6cm（空段落 space_before=Cm(6.0)）
        2. 报告标题（28pt Bold #000000 居中，行距 1.5）
        3. 英文副标题（14pt #000000 居中，metadata.subtitle 为 None 时跳过）
        4. 报告类型（16pt #000000 居中，默认"深度研究报告"）
        5. 分隔线（1pt 黑色，约 5cm 宽，居中）
        6. 机构名（14pt Bold #000000 居中，默认"遨天科技"）
        7. 版本 + 日期（11pt #000000 居中，格式 V{version} | {date}）

    不含密级字段（🚫严禁标密）。封面段落以 Normal 样式为基础
    （doc.add_paragraph() 默认已使用 Normal），run 级别覆盖格式属性。

    封面字段优先级（改进 14）：cover_metadata > MetadataIR > 默认值。
    cover_metadata 来自 --cover 参数指定的封面 MD 文件的 YAML frontmatter。

    Args:
        doc: python-docx Document 对象
        metadata: 文档头元数据（ir.MetadataIR）
        styles: 样式名 → 样式对象的映射（封面段落以 Normal 为基础，
            通过 run 级别属性覆盖；styles 参数保留供未来扩展使用）
        cover_metadata: 封面独立元数据 dict（可选）；键名与 cover.md YAML 一致：
            title / title_en / report_type / org / date / version / header_short
    """
    # 封面字段解析辅助：cover_metadata > MetadataIR > 默认值
    cm = cover_metadata or {}

    def _cv(cover_key: str, metadata_val: str | None, default: str = "") -> str:
        """取封面字段：cover_metadata[cover_key] > metadata_val > default。"""
        if cm.get(cover_key):
            return str(cm[cover_key])
        return metadata_val or default

    # 1. 顶部留白 6cm（I11 裁决：单个空段落 space_before=Cm(6.0)）
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Cm(COVER_TOP_SPACING_CM)

    # 2. 报告标题（28pt Bold，行距 1.5）
    title_spec = COVER_ELEMENTS["title"]
    cover_title = _cv("title", metadata.title)
    _add_cover_paragraph(
        doc,
        cover_title,
        title_spec.size_pt,
        title_spec.bold,
        title_spec.color_hex,
        title_spec.space_before_pt,
        title_spec.space_after_pt,
        line_spacing=1.5,
    )

    # 3. 英文副标题（可选：cover.title_en 或 metadata.subtitle 为 None 时跳过整段）
    cover_subtitle = _cv("title_en", metadata.subtitle)
    if cover_subtitle:
        sub_spec = COVER_ELEMENTS["subtitle"]
        _add_cover_paragraph(
            doc,
            cover_subtitle,
            sub_spec.size_pt,
            sub_spec.bold,
            sub_spec.color_hex,
            sub_spec.space_before_pt,
            sub_spec.space_after_pt,
        )

    # 4. 报告类型（16pt，默认文案"深度研究报告"）
    type_spec = COVER_ELEMENTS["report_type"]
    report_type_text = _cv("report_type", metadata.report_type, "深度研究报告")
    _add_cover_paragraph(
        doc,
        report_type_text,
        type_spec.size_pt,
        type_spec.bold,
        type_spec.color_hex,
        type_spec.space_before_pt,
        type_spec.space_after_pt,
    )

    # 5. 分隔线（1pt 黑色，约 5cm 宽，居中）
    _add_separator(doc)

    # 6. 机构名（14pt Bold，默认"遨天科技"）
    org_spec = COVER_ELEMENTS["organization"]
    org_text = _cv("org", metadata.organization, "遨天科技")
    _add_cover_paragraph(
        doc,
        org_text,
        org_spec.size_pt,
        org_spec.bold,
        org_spec.color_hex,
        org_spec.space_before_pt,
        org_spec.space_after_pt,
    )

    # 7. 版本 + 日期（11pt，格式 V{version} | {date}）
    vd_spec = COVER_ELEMENTS["version_date"]
    cover_version = _cv("version", metadata.version, "V1.0")
    cover_date = _cv("date", metadata.date)
    if cover_date:
        vd_line = f"{cover_version} | {cover_date}"
    else:
        vd_line = cover_version
    _add_cover_paragraph(
        doc,
        vd_line,
        vd_spec.size_pt,
        vd_spec.bold,
        vd_spec.color_hex,
        vd_spec.space_before_pt,
        vd_spec.space_after_pt,
    )
