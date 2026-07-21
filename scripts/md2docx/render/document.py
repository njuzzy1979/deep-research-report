"""渲染主调度器与 Document 组装（C-07c）。

将 DocumentIR 转换为完整的 .docx 文件。本模块是渲染管道的唯一入口，
负责：Document 创建、节（Section）管理、元素分派。

G-02 硬约束：add_page_break 调用**仅此文件**存在（PageBreakIR 唯一消费点）。
V-04 约束：分页符只在此处消费，其他渲染模块无分页能力。
"""
from __future__ import annotations

from docx import Document

from ..config import (
    FOOTER_DISTANCE_CM,
    HEADER_DISTANCE_CM,
    MARGIN_BOTTOM_CM,
    MARGIN_LEFT_CM,
    MARGIN_RIGHT_CM,
    MARGIN_TOP_CM,
    PAGE_HEIGHT_CM,
    PAGE_WIDTH_CM,
    BehaviorFlags,
)
from ..ir import (
    DocumentIR,
    FigureIR,
    HeadingIR,
    ListBlockIR,
    PageBreakIR,
    ParagraphIR,
    QuoteIR,
    SectionKind,
    SectionPlan,
    TableIR,
    TableKind,
)
from ..issues import Issue, IssueCollector, Level
from . import oxml_helpers as _oxml
from .cover import render_cover
from .figures import render_figure
from .headerfooter import setup_section_headers_footers
from .headings import render_heading
from .lists import render_bullet_list, render_numbered_list
from .paragraphs import render_paragraph, render_quote
from .styles import register_styles
from .tables import render_table
from .toc import render_chart_directory, render_toc, should_render_chart_directory

# ---------------------------------------------------------------------------
# 内部辅助：节设置
# ---------------------------------------------------------------------------


def _setup_page_dimensions(section) -> None:
    """设置节级页面尺寸：A4 + 标准边距 + 页眉/页脚距边。"""
    from docx.shared import Cm

    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)
    section.header_distance = Cm(HEADER_DISTANCE_CM)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)


# ---- 节页码格式定义 ----
# (section_kind → format_str, restart)
_SECTION_PG_CONFIG = {
    SectionKind.COVER: (None, False),
    SectionKind.ABSTRACT: ("lowerRoman", True),
    SectionKind.TOC: ("lowerRoman", False),
    SectionKind.BODY: ("decimal", True),
}


def _apply_section_pg(section, section_spec) -> None:
    """在当前节上设置页码格式（仅 document.py 设置，避免 python-docx sectPr 索引错位）。

    headerfooter.py 负责页眉/页脚内容，不重复调用 set_pgNumType。
    """
    fmt_str, restart = _SECTION_PG_CONFIG.get(
        section_spec.kind, (None, False)
    )
    if fmt_str is None:
        return
    start = 1 if restart else None
    _oxml.set_pgNumType(section._sectPr, fmt_str, start)


def _compute_section_ranges(
    section_plan: SectionPlan, total_elements: int
) -> dict[int, tuple[int, int]]:
    """计算每个节在 ir.elements 中消费的元素范围。

    COVER 节和 TOC 节的内容由渲染器自动生成，不从 ir.elements 消费。
    仅返回 ABSTRACT 和 BODY 节的元素范围。

    返回：{plan_index: (start_inclusive, end_exclusive)}
    """
    sections = section_plan.sections
    consuming: list[tuple[int, int]] = []
    for i, sec in enumerate(sections):
        if sec.kind in (SectionKind.ABSTRACT, SectionKind.BODY):
            consuming.append((i, sec.start_element_index))

    ranges: dict[int, tuple[int, int]] = {}
    for idx, (plan_idx, start) in enumerate(consuming):
        end = consuming[idx + 1][1] if idx + 1 < len(consuming) else total_elements
        ranges[plan_idx] = (start, end)

    return ranges


# ---------------------------------------------------------------------------
# 内部辅助：元素分派
# ---------------------------------------------------------------------------


def _dispatch_body_element(
    doc: Document,
    element,
    styles_map: dict,
    flags: BehaviorFlags,
    issues: IssueCollector,
    bookmark_id_counter: list[int],  # mutable list as counter
) -> None:
    """按元素类型分派到对应渲染模块。

    G-02：PageBreakIR 的唯一消费点——doc.add_page_break() 只在此函数中调用。
    """
    if isinstance(element, HeadingIR):
        render_heading(doc, element, styles_map)

    elif isinstance(element, ParagraphIR):
        render_paragraph(doc, element, styles_map)

    elif isinstance(element, FigureIR):
        render_figure(doc, element, styles_map, _oxml, issues, flags)

    elif isinstance(element, TableIR):
        used = render_table(doc, element, styles_map, issues, flags,
                            bookmark_id_start=bookmark_id_counter[0])
        bookmark_id_counter[0] += used

    elif isinstance(element, ListBlockIR):
        if element.ordered:
            render_numbered_list(doc, element, styles_map)
        else:
            render_bullet_list(doc, element, styles_map)

    elif isinstance(element, QuoteIR):
        render_quote(doc, element, styles_map)

    elif isinstance(element, PageBreakIR):
        # G-02/V-04：全项目唯一 add_page_break 调用点
        doc.add_page_break()

    else:
        # G-07 零静默：未识别元素类型降级为 INFO Issue
        issues.append(
            Issue(
                level=Level.INFO,
                code="I-DOC-01",
                stage="render",
                message=(
                    f"未识别的元素类型：{type(element).__name__}，"
                    f"已跳过渲染"
                ),
                source_line=getattr(element, "source_line", None),
                suggestion=(
                    "若该元素类型确需支持，请在 ir.py 的 BlockIR 联合类型中注册，"
                    "并在 _dispatch_body_element 中添加对应渲染分支"
                ),
            )
        )


# ===========================================================================
# 公开 API
# ===========================================================================


def render_document(
    ir: DocumentIR,
    output_path: str,
    flags: BehaviorFlags,
    styles_map: dict | None = None,
    issues: IssueCollector | None = None,
) -> str:
    """渲染管道的核心入口：将 DocumentIR 转换为 .docx 文件。

    按四节方案（或三节降级方案）依次渲染：
    Sec0 COVER → Sec1 ABSTRACT → Sec2 TOC → Sec3 BODY

    Args:
        ir: 文档中间表示
        output_path: 输出 .docx 文件路径
        flags: 行为开关
        styles_map: 预构建的样式字典；None 时自动调用 register_styles()
        issues: IssueCollector 实例；None 时静默略过

    Returns:
        输出文件路径（即 output_path）
    """
    if issues is None:
        issues = IssueCollector()

    doc = Document()

    # ---- 阶段 5.1：注册样式 ----
    if styles_map is None:
        styles_map = register_styles(doc)

    # ---- 阶段 5.2：设置页面尺寸（作用于默认首节） ----
    _setup_page_dimensions(doc.sections[0])

    # ---- 阶段 5.3：预计算节信息 ----
    section_plan = ir.section_plan
    if not section_plan.sections:
        doc.save(output_path)
        return output_path

    element_ranges = _compute_section_ranges(section_plan, len(ir.elements))

    # ---- 阶段 5.5-5.8：逐节渲染 ----
    bookmark_id_counter = [0]

    for sec_idx, section_spec in enumerate(section_plan.sections):
        if sec_idx > 0:
            new_sec = doc.add_section()
            _setup_page_dimensions(new_sec)

        if section_spec.kind == SectionKind.COVER:
            # 阶段 5.5：封面渲染（委托 cover.py）
            render_cover(doc, ir.metadata, styles_map)

        elif section_spec.kind == SectionKind.ABSTRACT:
            # 阶段 5.6：摘要渲染（委托 headings.py + paragraphs.py）
            elem_range = element_ranges.get(sec_idx, (0, 0))
            for element in ir.elements[elem_range[0]:elem_range[1]]:
                _dispatch_body_element(
                    doc, element, styles_map, flags, issues, bookmark_id_counter
                )

        elif section_spec.kind == SectionKind.TOC:
            # 阶段 5.7：目录渲染（委托 toc.py）
            render_toc(doc, ir, styles_map, _oxml)
            # 图表目录（G-02：add_page_break 唯一消费点）
            if should_render_chart_directory(ir):
                doc.add_page_break()
                figures = [e for e in ir.elements if isinstance(e, FigureIR)]
                body_tables = [
                    e for e in ir.elements
                    if isinstance(e, TableIR) and e.kind == TableKind.BODY and e.table_id is not None
                ]
                toc_heading_style = styles_map.get("TOC Heading")
                render_chart_directory(doc, figures, body_tables, toc_heading_style, _oxml)

        elif section_spec.kind == SectionKind.BODY:
            # 阶段 5.8：正文 + 附录渲染
            elem_range = element_ranges.get(sec_idx, (0, len(ir.elements)))
            for element in ir.elements[elem_range[0]:elem_range[1]]:
                _dispatch_body_element(
                    doc, element, styles_map, flags, issues, bookmark_id_counter
                )

    # ---- 延后设置页码格式：在所有节创建完毕后，doc.sections[i]._sectPr 才能正确解析 ----
    for sec_idx, section_spec in enumerate(section_plan.sections):
        _apply_section_pg(doc.sections[sec_idx], section_spec)

    # ---- 配置页眉页脚（委托 headerfooter.py） ----
    report_title = ir.metadata.title_short or ir.metadata.title or ""
    # 将 doc.sections 包装为 list 传入（python-docx Sections 对象已支持索引）
    setup_section_headers_footers(list(doc.sections), report_title)

    # ---- 阶段 5.9：保存 ----
    doc.save(output_path)
    return output_path


def render(
    document_ir: DocumentIR,
    options,  # RunOptions
    flags: BehaviorFlags,
    issues: IssueCollector,
) -> str:
    """pipeline.py 兼容入口（C-04 约定签名）。

    本函数是 render_document() 的薄包装，从 RunOptions 中提取 output_path。
    """
    return render_document(
        ir=document_ir,
        output_path=options.output_path,
        flags=flags,
        issues=issues,
    )
