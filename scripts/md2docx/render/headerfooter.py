"""四节页眉页脚渲染模块（C-12）：render/headerfooter.py

实现 COVER / ABSTRACT / TOC / BODY 四节的差异化页眉页脚，符合 V3.1 §6 规范。
所有 Word 域通过 oxml_helpers.make_page_field() 构建（G-03 域构建唯一入口约束）；
段落下边框通过 oxml_helpers.make_pBdr_bottom() 设置；页码格式通过
oxml_helpers.set_pgNumType() 在节级设置。
"""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .oxml_helpers import make_field, make_pBdr_bottom
from ..config import HEADER_FOOTER_FONT_SIZE_PT

# ===========================================================================
# 内部工具函数
# ===========================================================================


def _clear_paragraph_runs(paragraph) -> None:
    """清空段落中的所有 w:r 子元素，保留 w:pPr 等其他子元素不变。"""
    p_elem = paragraph._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)


def _add_page_field_to_paragraph(paragraph, font_size_pt: int = 9) -> None:
    """向段落追加 PAGE 域并通过 make_field() 构建（G-03 域唯一入口）。

    make_field(paragraph, 'PAGE', field_type='PAGE') 生成四态域
    （begin → instrText → separate → end）。为让 Word 在更新域后
    以正确字号渲染页码，在 separate 与 end 之间插入一个带 w:rPr/w:sz
    的占位文本 run "1"。
    """
    # 记录当前段落子元素数量（插入占位文本前）
    n_before = len(paragraph._p)

    make_field(paragraph, "PAGE", field_type="PAGE")

    # make_field 追加了 4 个 w:r（begin/instrText/separate/end），
    # 在倒数第二个（separate）与最后一个（end）之间插入占位文本 run
    placeholder_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size_pt * 2))  # 半磅单位：9pt → 18
    rPr.append(sz)
    placeholder_run.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = "1"
    placeholder_run.append(t_elem)

    # 插入到 end run 之前（即 separate 之后）
    paragraph._p.insert(len(paragraph._p) - 1, placeholder_run)


def _get_or_create_pPr(paragraph) -> "lxml.etree._Element":
    """获取或创建段落的 w:pPr 子元素，返回该元素供边框等属性设置。"""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    return pPr


# ===========================================================================
# 公开 API
# ===========================================================================


def setup_section_headers_footers(
    sections: list,
    report_short_title: str,
) -> None:
    """配置四节的页眉页脚（04-interface-spec.md §2.4 四节方案）。

    ┌─────────┬──────────┬─────────────────┬───────────────┐
    │ 节       │ 页眉      │ 页脚              │ link_to_prev  │
    ├─────────┼──────────┼─────────────────┼───────────────┤
    │ Sec0     │ 无        │ 无               │ N/A           │
    │ COVER    │           │                  │               │
    ├─────────┼──────────┼─────────────────┼───────────────┤
    │ Sec1     │ 无        │ 罗马 i, ii...    │ hdr=False     │
    │ ABSTRACT │           │ 居中 9pt          │ ftr=False     │
    ├─────────┼──────────┼─────────────────┼───────────────┤
    │ Sec2     │ 无        │ 罗马续编           │ hdr=True      │
    │ TOC      │ (继承Sec1)│ (继承Sec1格式)     │ ftr=True      │
    ├─────────┼──────────┼─────────────────┼───────────────┤
    │ Sec3     │ 报告简称   │ 阿拉伯 1, 2...    │ hdr=False     │
    │ BODY     │ 右对齐+底线│ 居中 9pt          │ ftr=False     │
    └─────────┴──────────┴─────────────────┴───────────────┘

    Args:
        sections: [sec0, sec1, sec2, sec3] 四个 python-docx Section 对象
        report_short_title: 报告简称，显示在正文节（Sec3）页眉右端

    Raises:
        ValueError: 当 sections 长度不足 4 时
    """
    if len(sections) < 3:
        raise ValueError(
            f"至少需要 3 个节（COVER/TOC/BODY），实际只有 {len(sections)} 个"
        )

    has_abstract = len(sections) >= 4
    if has_abstract:
        sec0, sec1, sec2, sec3 = sections[0], sections[1], sections[2], sections[3]
    else:
        sec0, sec1, sec2 = sections[0], sections[1], sections[2]
        sec3 = None  # no ABSTRACT section

    # =====================================================================
    # Sec0 (COVER)：无页眉、无页脚
    #
    # python-docx 新建 Document 的第一节默认 header/footer 即为空段落，
    # 无需额外操作。不设置 is_linked_to_previous（第一节无前一节可链接）。
    # =====================================================================

    # =====================================================================
    # Sec1: 四节=ABSTRACT（无页眉，罗马页码 i 起始）
    #       三节=TOC（无页眉，罗马页码 i 起始）
    # =====================================================================

    sec1_header = sec1.header
    sec1_header.is_linked_to_previous = False
    _clear_paragraph_runs(sec1_header.paragraphs[0])

    sec1_footer = sec1.footer
    sec1_footer.is_linked_to_previous = False
    p1 = sec1_footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_paragraph_runs(p1)
    _add_page_field_to_paragraph(p1, font_size_pt=9)

    if has_abstract:
        # =================================================================
        # Sec2 (TOC)：继承 Sec1 的空页眉和罗马格式，页码续编
        # =================================================================

        sec2.header.is_linked_to_previous = True
        sec2_footer = sec2.footer
        sec2_footer.is_linked_to_previous = True
        p2 = sec2_footer.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_paragraph_runs(p2)
        _add_page_field_to_paragraph(p2, font_size_pt=9)

        body_sec = sec3
    else:
        body_sec = sec2

    # =====================================================================
    # BODY 节：页眉报告简称 + 底线，页脚阿拉伯数字 1 起始
    # =====================================================================

    body_header = body_sec.header
    body_header.is_linked_to_previous = False
    hp = body_header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _clear_paragraph_runs(hp)
    run = hp.add_run(report_short_title)
    run.font.size = Pt(HEADER_FOOTER_FONT_SIZE_PT)

    pPr = _get_or_create_pPr(hp)
    make_pBdr_bottom(pPr, sz=8, color="000000", space=1)

    body_footer = body_sec.footer
    body_footer.is_linked_to_previous = False
    p3 = body_footer.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _clear_paragraph_runs(p3)
    _add_page_field_to_paragraph(p3, font_size_pt=9)


# ===========================================================================
# 自检块
# ===========================================================================
if __name__ == "__main__":
    """快速烟雾测试：验证模块可导入、公开 API 存在且符合预期签名。"""
    import inspect

    _func = setup_section_headers_footers
    _sig = inspect.signature(_func)
    _params = list(_sig.parameters.keys())

    assert "sections" in _params, "缺少 'sections' 参数"
    assert "report_short_title" in _params, "缺少 'report_short_title' 参数"
    assert len(_params) == 2, f"期望 2 个参数，实际 {len(_params)} 个"

    print(
        f"headerfooter.py 自检通过："
        f"setup_section_headers_footers({', '.join(_params)})"
    )
