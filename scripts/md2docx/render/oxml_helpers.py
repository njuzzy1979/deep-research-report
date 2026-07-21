"""OXML 底层积木函数模块 —— render 层的 OXML 原子操作集。

G-03 硬约束：``w:fldChar`` 字符串字面量**仅此文件**存在；全项目所有域代码生成
必须通过本文件的 :func:`make_field` 函数。

本模块是 md→docx 转换器渲染层的基础模块。python-docx 1.2.0 对许多 Word 高级
功能（域代码、书签、页码类型、段落边框、表格边框、底纹、超链接、首行缩进字符数）
只提供有限 API 或不提供 API，必须通过原始 OXML 操作实现。

所有 OXML 元素创建使用 ``OxmlElement(tag)`` + ``qn()`` 处理命名空间。
命名空间前缀 ``w:`` 对应
``http://schemas.openxmlformats.org/wordprocessingml/2006/main``。

设计依据：
- 04-interface-spec.md §2.4（分节页码）/ §2.5（TOC 域四态）/ §2.6（表格全框线）
  / §2.10（超链接）
- 00-master-design.md M4（封面分隔线）/ M5（表头下分隔线优先级）
- G-04：Pt() 和 Cm() 的 import 仅允许存在于 config.py、render/styles.py、本文件
"""
from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from ..config import (
    BODY_FIRST_LINE_CHARS,
    COLOR_BLACK,
    COLOR_HYPERLINK,
    COVER_SEPARATOR,
    TABLE_BORDERS,
)


# =============================================================================
# 内部辅助
# =============================================================================

def _strip_color_prefix(color_str: str) -> str:
    """去掉颜色字符串的 # 前缀（OOXML 颜色属性不接受 # 前缀）。"""
    return color_str.lstrip('#') if color_str else color_str


def _make_border_element(side_name: str, attrs: dict):
    """创建单个边框子元素（w:left / w:right / w:top / w:bottom）。

    Args:
        side_name: 边框方向名（'left' / 'right' / 'top' / 'bottom'）。
        attrs: 边框属性字典，键为 OOXML 属性名（不带 w: 前缀），
            值为字符串。颜色值应已去掉 # 前缀。

    Returns:
        OxmlElement（如 ``<w:left w:val="single" w:sz="4" .../>``）。
    """
    el = OxmlElement(f'w:{side_name}')
    for attr_name, attr_value in attrs.items():
        el.set(qn(f'w:{attr_name}'), str(attr_value))
    return el


# =============================================================================
# 1. 域代码生成（G-03：fldChar 字面量仅此文件出现）
# =============================================================================

def make_field(paragraph, instr_text: str, field_type: str = 'PAGE'):
    """生成 Word 域代码的四态结构，追加到指定段落。

    Word 域代码的四态结构（begin → instrText → separate → end）是 Word 识别
    域代码的必要条件。缺少 separate 会导致 Word 不识别为域（历史缺陷2的教训）。

    四态对应的 XML 结构::

        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> INSTRUCTION </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>

    Args:
        paragraph: 要追加域的段落对象（python-docx Paragraph）。
        instr_text: 域指令文本。常用值：

            * TOC 目录: ``'TOC \\o "1-3" \\h \\z \\u'``
            * PAGE 页码: ``'PAGE'``
            * PAGEREF 书签页码: ``'PAGEREF BookmarkName \\h'``

        field_type: 域类型标识（``'PAGE'`` / ``'TOC'`` / ``'PAGEREF'``），
            当前仅用于文档说明与调用方可读性，不影响生成的 XML。

    Returns:
        begin run 的 fldChar 元素（OxmlElement），用于后续在 begin 和
        separate 之间插入额外内容（如 TOC 占位文本）。
    """
    # ① begin
    begin_run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    begin_run._r.append(fldChar_begin)

    # ② instrText（xml:space="preserve" 必须——否则含空格的域指令被 Word 截断）
    instr_run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' ' + instr_text + ' '
    instr_run._r.append(instrText)

    # ③ separate（必须——缺少则 Word 不识别为域，历史缺陷2）
    sep_run = paragraph.add_run()
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    sep_run._r.append(fldChar_sep)

    # ④ end
    end_run = paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(fldChar_end)

    return fldChar_begin


# =============================================================================
# 2. 书签
# =============================================================================

def make_bookmark_start(paragraph, bookmark_id: int, bookmark_name: str):
    """在段落中插入书签起始标记。

    bookmarkStart 和 bookmarkEnd 的 id 必须配对（同一个整数），由调用方保证。
    本函数不做全局注册或唯一性校验——id 管理职责在调用方。

    XML 示例::

        <w:bookmarkStart w:id="0" w:name="Fig3_2"/>

    Args:
        paragraph: 目标段落对象。
        bookmark_id: 书签整数 ID（需调用方自行管理唯一性，与 bookmarkEnd 配对）。
        bookmark_name: 书签名称字符串（如 ``'Fig3_2'``）。

    Returns:
        创建的 bookmarkStart OxmlElement。
    """
    bookmark = OxmlElement('w:bookmarkStart')
    bookmark.set(qn('w:id'), str(bookmark_id))
    bookmark.set(qn('w:name'), bookmark_name)
    paragraph._p.append(bookmark)
    return bookmark


def make_bookmark_end(paragraph, bookmark_id: int, bookmark_name: str):
    """在段落中插入书签结束标记。

    bookmarkEnd 在 OOXML 中仅含 ``w:id`` 属性（不含 ``w:name``），与
    bookmarkStart 的 id 必须相同。``bookmark_name`` 参数保留用于 API
    签名一致性，不写入 XML。

    XML 示例::

        <w:bookmarkEnd w:id="0"/>

    Args:
        paragraph: 目标段落对象。
        bookmark_id: 书签整数 ID（必须与对应 bookmarkStart 的 id 相同）。
        bookmark_name: 书签名称（API 一致性参数，不写入 bookmarkEnd 元素）。

    Returns:
        创建的 bookmarkEnd OxmlElement。
    """
    bookmark = OxmlElement('w:bookmarkEnd')
    bookmark.set(qn('w:id'), str(bookmark_id))
    paragraph._p.append(bookmark)
    return bookmark


# =============================================================================
# 3. 页码格式
# =============================================================================

def set_pgNumType(sectPr, format_str: str, start_num: int | None = None):
    """设置节的页码类型（罗马数字/阿拉伯数字）和起始值。

    典型用法：

    * 第二节（摘要/目录）：``fmt='lowerRoman', start=1`` 以重新开始罗马计数。
    * 第三节（正文）：``fmt='decimal', start=1`` 以重新开始阿拉伯计数。

    XML 示例::

        <w:pgNumType w:fmt="lowerRoman" w:start="1"/>

    Args:
        sectPr: 节的 sectPr XML 元素（如 ``document.sections[0]._sectPr``）。
        format_str: 页码格式字符串。支持 ``'lowerRoman'`` / ``'upperRoman'`` /
            ``'decimal'`` / ``'chineseCounting'`` 等（直接写入 ``w:fmt`` 属性）。
        start_num: 起始页码（整数），为 None 时不设置 ``w:start`` 属性
            （页码续接上一节）。
    """
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_str)
    if start_num is not None:
        pgNumType.set(qn('w:start'), str(start_num))
    sectPr.append(pgNumType)


# =============================================================================
# 4. 段落边框与页面边框
# =============================================================================

def set_paragraph_borders(paragraph, **borders):
    """设置段落边框（可用于 Quote 样式左侧竖线等）。

    在段落的 ``pPr`` 中添加或替换 ``w:pBdr`` 元素。已存在的 ``pBdr`` 会被
    移除后重建，避免多次调用产生重复边框元素。

    Args:
        paragraph: 目标段落对象。
        **borders: 关键字参数，键为边框方向名（``'left'`` / ``'right'`` /
            ``'top'`` / ``'bottom'``），值为属性字典。颜色值可带或不带 ``#``
            前缀（本函数自动处理）。

    Example:
        设置 Quote 样式左侧灰色竖线::

            set_paragraph_borders(paragraph,
                left={'val': 'single', 'sz': '4', 'space': '4', 'color': '#BFBFBF'})

        设置页眉底部 1pt 黑色底线::

            set_paragraph_borders(paragraph,
                bottom={'val': 'single', 'sz': '8', 'space': '1', 'color': '#000000'})

    XML 示例::

        <w:pBdr>
          <w:left w:val="single" w:sz="4" w:space="4" w:color="BFBFBF"/>
        </w:pBdr>
    """
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:pBdr')):
        pPr.remove(old)

    pBdr = OxmlElement('w:pBdr')
    for side_name, attrs in borders.items():
        cleaned = {
            k: _strip_color_prefix(v) if k == 'color' else str(v)
            for k, v in attrs.items()
        }
        pBdr.append(_make_border_element(side_name, cleaned))
    pPr.append(pBdr)


def set_page_borders(sectPr, **borders):
    """设置页面边框（节级，通过 sectPr 下的 ``w:pgBorders`` 实现）。

    用法同 :func:`set_paragraph_borders`，但作用域为整节而非单段落。

    Args:
        sectPr: 节的 sectPr XML 元素。
        **borders: 同 :func:`set_paragraph_borders` 的边框参数格式。
    """
    for old in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(old)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side_name, attrs in borders.items():
        cleaned = {
            k: _strip_color_prefix(v) if k == 'color' else str(v)
            for k, v in attrs.items()
        }
        pgBorders.append(_make_border_element(side_name, cleaned))
    sectPr.append(pgBorders)


# =============================================================================
# 5. 段落底部边框便捷函数（cover.py 依赖，保持向后兼容）
# =============================================================================

def make_pBdr_bottom(pPr_element, sz: int = 8, color: str = "000000", space: int = 1):
    """在段落属性中添加底部边框（页眉底线 / 封面分隔线用）。

    产出 XML::

        <w:pBdr>
          <w:bottom w:val="single" w:sz="8" w:space="1" w:color="000000"/>
        </w:pBdr>

    用于 headerfooter.py 中页眉段落的下划线（V3.0 §6.1）以及 cover.py 中
    封面分隔线（M4）。

    Args:
        pPr_element: 段落属性元素（w:pPr），即 ``paragraph._p.get_or_add_pPr()``
            的返回值。
        sz: 线宽（八分之一磅单位），默认 8 = 1pt。
        color: 颜色十六进制（可带或不带 # 号），默认 ``"000000"``。
        space: 线与文字间距（磅单位），默认 1。
    """
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), _strip_color_prefix(color))
    pBdr.append(bottom)
    pPr_element.append(pBdr)


# =============================================================================
# 6. 段落底纹通用函数
# =============================================================================

def make_shd(element, fill_hex: str, val: str = "clear"):
    """设置底纹（段落级 pPr 或单元格级 tcPr）。

    产出 XML::

        <w:shd w:val="clear" w:color="auto" w:fill="F2F2F2"/>

    用途：

    - 表格交替行底纹（04-interface-spec.md §2.6，颜色 #F2F2F2）
    - 单元格级或段落级底纹

    Args:
        element: 目标元素（tcPr 或 pPr）。
        fill_hex: 填充色（可带或不带 # 号），如 ``"F2F2F2"`` 或 ``"#F2F2F2"``。
        val: shd 类型，默认 ``"clear"``（清除型填充，仅 fill 生效）。
    """
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), val)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _strip_color_prefix(fill_hex))
    element.append(shd)


# =============================================================================
# 7. 表格边框全框线（04-interface-spec.md §2.6 全框线规格）
# =============================================================================

def make_tblBorders(tblPr_element, border_spec) -> None:
    """设置表格级全框线（tblPr → tblBorders）。

    border_spec 是 config.TableBorderSpec 的实例（dataclass），包含：
    top_sz, bottom_sz, inside_h_sz, inside_v_sz, left_sz, right_sz, color_hex
    （sz 单位均为八分之一磅的整数）。

    产出 XML::

        <w:tblBorders>
          <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
          <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
          <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
          <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tblBorders>

    M5 说明：表头下 1pt 分隔线不在此处设置，而是通过
    :func:`make_tcBorders_bottom` 在单元格级 tcBorders 上覆盖表级 insideH
    （0.5pt），因为 OXML 中单元格级边框优先级高于表格级
    （04-interface-spec.md §2.6 I4 裁决）。

    Args:
        tblPr_element: 表格属性元素（w:tblPr），即 ``table._tbl.tblPr``。
        border_spec: config.TableBorderSpec 实例。
    """
    tblBorders = OxmlElement("w:tblBorders")

    def _add_border(tag: str, sz_val: int, clr: str) -> None:
        elem = OxmlElement(tag)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(sz_val))
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), _strip_color_prefix(clr))
        tblBorders.append(elem)

    _add_border("w:top", border_spec.top_sz, border_spec.color_hex)
    _add_border("w:bottom", border_spec.bottom_sz, border_spec.color_hex)
    _add_border("w:insideH", border_spec.inside_h_sz, border_spec.color_hex)
    _add_border("w:insideV", border_spec.inside_v_sz, border_spec.color_hex)
    _add_border("w:left", border_spec.left_sz, border_spec.color_hex)
    _add_border("w:right", border_spec.right_sz, border_spec.color_hex)

    tblPr_element.append(tblBorders)


def make_tcBorders_bottom(tcPr_element, sz: int = 8, color: str = "000000"):
    """设置单元格级底部边框（表头下分隔线，M4 1pt 差异化）。

    产出 XML::

        <w:tcBorders>
          <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>
        </w:tcBorders>

    作用域仅限当前单元格（tc）；用于表头行的每个单元格，共同构成表头下
    1pt 分隔线效果（覆盖表级 tblBorders.insideH 的 0.5pt）。

    Args:
        tcPr_element: 单元格属性元素（w:tcPr），即 ``cell._tc.tcPr``。
        sz: 线宽（八分之一磅单位），默认 8 = 1pt。
        color: 颜色十六进制（可带或不带 # 号），默认 ``"000000"``。
    """
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), _strip_color_prefix(color))
    tcBorders.append(bottom)
    tcPr_element.append(tcBorders)


def add_table_borders(table):
    """设置表格全框线 + 表头下分隔线（一站式便捷函数）。

    内部调用 :func:`make_tblBorders` 设置表级边框，再遍历表头行每个单元格
    调用 :func:`make_tcBorders_bottom` 设置表头下 1pt 分隔线。

    边框逐项对照（04-interface-spec.md §2.6）：

    ============ ====== ==== ====================================
    边线         粗细   sz   实现层级
    ============ ====== ==== ====================================
    顶线         1.5pt  12   tblPr/w:tblBorders/w:top
    底线         1.5pt  12   tblPr/w:tblBorders/w:bottom
    表头下分隔线 1pt    8    表头行 tcPr/w:tcBorders/w:bottom
    内部横线     0.5pt  4    tblPr/w:tblBorders/w:insideH
    内部竖线     0.5pt  4    tblPr/w:tblBorders/w:insideV
    左/右外框    0.5pt  4    tblPr/w:tblBorders/w:left / w:right
    ============ ====== ==== ====================================

    Args:
        table: python-docx Table 对象。
    """
    # 表级边框
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    make_tblBorders(tblPr, TABLE_BORDERS)

    # 表头下分隔线（单元格级 tcBorders 覆盖表级 insideH）
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(old)
            make_tcBorders_bottom(
                tcPr,
                sz=TABLE_BORDERS.header_bottom_sz,
                color=TABLE_BORDERS.color_hex,
            )


# =============================================================================
# 8. 单元格底纹
# =============================================================================

def set_cell_shading(cell, color: str):
    """设置单元格底纹颜色（用于表格交替行底纹、定义框等）。

    XML 示例::

        <w:shd w:val="clear" w:fill="F2F2F2"/>

    Args:
        cell: python-docx Cell 对象。
        color: 颜色十六进制字符串，如 ``'#F2F2F2'``（带或不带 ``#`` 前缀均可）。
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    make_shd(tcPr, _strip_color_prefix(color))


# =============================================================================
# 9. 超链接
# =============================================================================

def add_hyperlink(paragraph, url: str, text: str):
    """在段落中添加可点击超链接（通过 OPC Hyperlink 关系，非纯文本 URL）。

    超链接样式：蓝色（``#0563C1``，Word 默认超链接色）+ 下划线，
    同时引用 Word 内置 ``Hyperlink`` 字符样式以继承完整超链接行为。

    实现路径（tech#10）：

    1. 通过 ``paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)``
       获取关系 ``r:id``。
    2. 构建 ``<w:hyperlink r:id="...">`` 包裹内部 run。

    Args:
        paragraph: 目标段落对象。
        url: 超链接 URL。
        text: 显示文本。
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # 内部 run：引用 Hyperlink 字符样式 + 显式颜色/下划线兜底
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), _strip_color_prefix(COLOR_HYPERLINK))
    rPr.append(color_el)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# =============================================================================
# 10. 中文字体设置
# =============================================================================

def set_eastAsia_font(rPr_element, font_name: str):
    """在 run 属性或样式 rPr 上设置 w:eastAsia 字体。

    产出效果：``<w:rFonts w:eastAsia="宋体" .../>``

    若 rPr_element 中尚不存在 w:rFonts 子元素则自动创建并插入到首位
    （OXML 中 rFonts 通常位于 rPr 的首个子元素位置，与 Word 生成的
    XML 保持风格一致）。

    Args:
        rPr_element: w:rPr 元素（run 属性或样式定义中的 rPr）。
        font_name: 中文字体名（``"宋体"`` / ``"微软雅黑"`` 等）。
    """
    rFonts = rPr_element.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr_element.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# =============================================================================
# 11. 首行缩进（字符模式）
# =============================================================================

def set_firstLineChars(paragraph, chars: int = 200):
    """设置段落首行缩进 N 字符（字符模式）。

    Word 优先读取 ``w:firstLineChars`` 属性；长度模式兜底值（Cm/Pt）应由
    调用方（样式构建器 render/styles.py）通过 python-docx 原生 API 同时设置
    ——遵循"各层只做自己该做的事"原则（tech#1 双写策略）。

    字符单位：100 = 1 字符。默认 200 = 2 字符，0 表示无缩进。

    XML 示例::

        <w:ind w:firstLineChars="200"/>

    Args:
        paragraph: 目标段落对象。
        chars: 缩进值（百倍字符数）。默认值取
            config.BODY_FIRST_LINE_CHARS（200 = 2 字符）。
    """
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))


# =============================================================================
# 12. 封面分隔线（M4：机构名上方 1pt 黑色横线）
# =============================================================================

# 封面分隔线居中缩进量（twips）。
# 版心宽 = 21.0 - 3.17 - 2.54 = 15.29 cm（04-interface-spec.md §2.6）
# 单侧缩进 = (15.29 - COVER_SEPARATOR 线宽 5.0) / 2 ≈ 5.145 cm ≈ 2916 twips
_SEPARATOR_INDENT_TWIPS = '2916'


def make_cover_separator_line(paragraph):
    """生成封面分隔线——机构名上方 1pt 黑色横线（M4）。

    通过段落 ``w:pBdr`` 的 bottom 线实现。分隔线规格取自
    config.COVER_SEPARATOR：sz=8（1pt）、颜色 #000000、宽度约 5cm 居中。

    居中效果通过同时设置段落左右缩进实现（缩进量 = (版心宽 - 线宽) / 2），
    使底部边框线在视觉上呈现为一条约 5cm 宽的居中横线。

    Args:
        paragraph: 封面分隔线所在段落（通常为封面机构名上方的空段落）。
    """
    set_paragraph_borders(
        paragraph,
        bottom={
            'val': 'single',
            'sz': str(COVER_SEPARATOR.sz),
            'space': '1',
            'color': _strip_color_prefix(COVER_SEPARATOR.color_hex),
        },
    )

    # 左右缩进使底部边框线宽度 ≈ 5cm 居中
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:left'), _SEPARATOR_INDENT_TWIPS)
    ind.set(qn('w:right'), _SEPARATOR_INDENT_TWIPS)
