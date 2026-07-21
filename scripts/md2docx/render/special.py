"""特殊元素渲染模块（C-07e）：定义框、案例框、引用块等特殊元素的渲染。

本期以防御性实现为主——定义框和案例框在样本报告中极罕见，但需提供完整的
OXML 级渲染能力。所有不支持的降级操作必须记录 Issue（零静默 G-07）。
"""
from __future__ import annotations

from docx.oxml.ns import qn

from ..issues import Issue, IssueCollector, Level

# ---------------------------------------------------------------------------
# 定义框（V3.1 §9.1）
# ---------------------------------------------------------------------------


def render_definition_box(doc, token, styles: dict, oxml_helpers) -> None:
    """渲染定义框/知识卡片（V3.1 §9.1）。

    视觉规格：
      - 浅灰底 (#F2F2F2) + 黑色左边框 (3pt)
      - 术语名 10.5pt Bold，定义内容 10.5pt 正常

    通过段落底纹 ``w:shd`` + 段落左边框 ``w:pBdr/left`` 实现。

    Args:
        doc: python-docx Document 对象。
        token: 定义框 IR（含 term 文本和 definition runs）。
        styles: 样式名 → 样式对象映射。
        oxml_helpers: render/oxml_helpers 模块引用。
    """
    term = getattr(token, 'term', '')
    definition_runs = getattr(token, 'runs', [])

    # 术语名段落：10.5pt Bold + 浅灰底 + 左边框
    p_term = doc.add_paragraph()
    p_term.style = styles.get('Body Text', doc.styles['Normal'])
    run_term = p_term.add_run(term)
    run_term.font.bold = True

    _apply_paragraph_shading(p_term, '#F2F2F2', oxml_helpers)
    oxml_helpers.set_paragraph_borders(
        p_term,
        left={'val': 'single', 'sz': '24', 'space': '4', 'color': '000000'},
    )

    # 定义内容段落：10.5pt 正常 + 浅灰底 + 左边框
    p_def = doc.add_paragraph()
    p_def.style = styles.get('Body Text', doc.styles['Normal'])
    for irun in definition_runs:
        run = p_def.add_run(irun.text)
        if irun.bold:
            run.font.bold = True
        if irun.italic:
            run.font.italic = True

    _apply_paragraph_shading(p_def, '#F2F2F2', oxml_helpers)
    oxml_helpers.set_paragraph_borders(
        p_def,
        left={'val': 'single', 'sz': '24', 'space': '4', 'color': '000000'},
    )


# ---------------------------------------------------------------------------
# 案例框（V3.1 §9.2）
# ---------------------------------------------------------------------------


def render_case_box(doc, token, styles: dict, oxml_helpers) -> None:
    """渲染案例框（V3.1 §9.2）。

    视觉规格：
      - 白底 + 黑色边框 (1.5pt) + 浅灰顶边框 (3pt, #D9D9D9)
      - 案例标题 10.5pt Bold → 标题底部边框兼作分隔线 → 案例细节 10.5pt 正常

    通过段落边框 ``w:pBdr`` 实现：
      - 标题段落：上(D9D9D9 3pt) + 左右下(黑 1.5pt)，底部边框兼作"标题-细节"分隔线
      - 细节段落：左右下(黑 1.5pt)，与标题段落的左右框线对齐形成连续框体

    Args:
        doc: python-docx Document 对象。
        token: 案例框 IR（含 title 文本和 body runs）。
        styles: 样式名 → 样式对象映射。
        oxml_helpers: render/oxml_helpers 模块引用。
    """
    title = getattr(token, 'title', '')
    body_runs = getattr(token, 'runs', [])

    black_12 = {'val': 'single', 'sz': '12', 'space': '4', 'color': '000000'}
    gray_top = {'val': 'single', 'sz': '24', 'space': '4', 'color': 'D9D9D9'}

    # 案例标题段落：Bold + 四边框（浅灰顶，其余黑 1.5pt）
    # 底部黑边框兼作"标题-细节"视觉分隔线
    p_title = doc.add_paragraph()
    p_title.style = styles.get('Body Text', doc.styles['Normal'])
    run_title = p_title.add_run(title)
    run_title.font.bold = True
    oxml_helpers.set_paragraph_borders(
        p_title,
        top=gray_top,
        bottom=black_12,
        left=black_12,
        right=black_12,
    )

    # 案例细节段落：正常 + 左/右/下边框（延续标题段落的左右框线，形成连续框体）
    p_body = doc.add_paragraph()
    p_body.style = styles.get('Body Text', doc.styles['Normal'])
    for irun in body_runs:
        run = p_body.add_run(irun.text)
        if irun.bold:
            run.font.bold = True
        if irun.italic:
            run.font.italic = True

    oxml_helpers.set_paragraph_borders(
        p_body,
        bottom=black_12,
        left=black_12,
        right=black_12,
    )


# ---------------------------------------------------------------------------
# 特殊元素分派
# ---------------------------------------------------------------------------


def render_special_element(
    doc, token, styles: dict, oxml_helpers, issues: IssueCollector
) -> None:
    """特殊元素分派：根据 token 子类型渲染对应元素。

    当前仅占位框架，支持 definition_box 和 case_box 两种子类型。
    未知类型降级为普通段落并记录 I-SPL-01 Issue（零静默 G-07）。

    Args:
        doc: python-docx Document 对象。
        token: 特殊元素 IR（使用 getattr 防御性访问 special_type 字段）。
        styles: 样式名 → 样式对象映射。
        oxml_helpers: render/oxml_helpers 模块引用。
        issues: 全流程 Issue 收集器。
    """
    stype = getattr(token, 'special_type', 'unknown')

    if stype == 'definition_box':
        render_definition_box(doc, token, styles, oxml_helpers)
    elif stype == 'case_box':
        render_case_box(doc, token, styles, oxml_helpers)
    else:
        # 降级渲染为普通段落 + INFO Issue（G-07 零静默）
        issues.append(
            Issue(
                level=Level.INFO,
                code='I-SPL-01',
                stage='render',
                message=f'未知特殊元素类型: {stype}，降级为普通段落',
                source_line=getattr(token, 'source_line', None),
                suggestion=(
                    '若该元素类型确需支持，请在 ir.py 中新增对应的 IR 类型'
                    '并在本模块添加渲染分支'
                ),
            )
        )
        p = doc.add_paragraph()
        p.style = styles.get('Body Text', doc.styles['Normal'])
        for irun in getattr(token, 'runs', []):
            run = p.add_run(irun.text)
            if getattr(irun, 'bold', False):
                run.font.bold = True
            if getattr(irun, 'italic', False):
                run.font.italic = True


# ---------------------------------------------------------------------------
# 内部辅助
# ---------------------------------------------------------------------------


def _apply_paragraph_shading(paragraph, fill_hex: str, oxml_helpers) -> None:
    """对段落应用底纹（通过 pPr → w:shd 实现）。

    先移除已有 shd 元素避免重复，再通过 oxml_helpers.make_shd 追加新底纹。

    Args:
        paragraph: python-docx Paragraph 对象。
        fill_hex: 填充色十六进制字符串（可带或不带 # 前缀）。
        oxml_helpers: render/oxml_helpers 模块引用。
    """
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    oxml_helpers.make_shd(pPr, fill_hex)
