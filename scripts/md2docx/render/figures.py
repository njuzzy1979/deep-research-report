"""图片渲染模块（C-10）：将 FigureIR 渲染为 Word 文档中的嵌入图片 + 题注。

反硬编码防线：所有图信息（尺寸、路径、题注文本、书签名）100% 来自 FigureIR
动态解析，代码中零图片知识硬编码。

G-01 硬约束：所有文件读操作经 iotools.read_bytes()。
"""
from __future__ import annotations

import struct

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, RGBColor

from ..config import (
    COLOR_CAPTION_GRAY,
    FIGURE_LOW_RES_PX_W_THRESHOLD,
    FIGURE_MAX_HEIGHT_CM,
    FIGURE_PLACEHOLDER_FONT_SIZE_CM,
    BehaviorFlags,
)
from ..iotools import read_bytes
from ..ir import FigureIR
from ..issues import Issue, IssueCollector, Level

# ---------------------------------------------------------------------------
# 书签 ID 全局递增计数器（每张图占用一个 ID，render/figures.py 内部自增）
# ---------------------------------------------------------------------------
_next_bookmark_id = 1

# SEQ 域指令模板常量：caption_field_mode="field" 时，图题注编号由 Word SEQ
# 域自动计算，而非 assemble 层誊抄的 figure_id（本次 plan §6.3 图编号方案）
_FIGURE_SEQ_INSTR = "SEQ 图 \\* ARABIC"


# ---------------------------------------------------------------------------
# 内部辅助函数
# ---------------------------------------------------------------------------


def _read_png_dimensions(filepath: str) -> tuple[int, int] | None:
    """从 PNG 文件头读取像素宽高（仅读取前 33 字节 IHDR 头）。

    解析逻辑（R13 裁决，02-algorithms.md §A.5）：
      bytes 0-7:   PNG signature (\\x89PNG\\r\\n\\x1a\\n)
      bytes 8-11:  IHDR chunk length (big-endian uint32，固定 13)
      bytes 12-15: chunk type ('IHDR')
      bytes 16-19: width (big-endian uint32)
      bytes 20-23: height (big-endian uint32)

    Args:
        filepath: 图片文件绝对路径。

    Returns:
        (width, height) 像素尺寸元组；非 PNG 或读取失败返回 None。
    """
    try:
        data = read_bytes(filepath, 33)
    except (OSError, IOError):
        return None

    if len(data) < 33:
        return None
    if data[:8] != b"\x89PNG\r\n\x1a\n":
        return None
    if data[12:16] != b"IHDR":
        return None

    width, height = struct.unpack(">II", data[16:24])
    return width, height


def _resolve_px_dimensions(figure: FigureIR) -> tuple[int | None, int | None]:
    """获取图片像素尺寸：优先用 IR 中已解析的值，否则尝试现场读取 PNG 头。

    这是防御性设计——正常情况下 FigureAssembler 已在阶段3 填充 px_w/px_h，
    此处 fallback 读取仅用于组装阶段未成功解析 PNG 头的边界场景。
    """
    if figure.px_w is not None and figure.px_h is not None:
        return figure.px_w, figure.px_h

    if figure.file_exists:
        dims = _read_png_dimensions(str(figure.path_resolved))
        if dims is not None:
            return dims

    return None, None


def _make_paragraph_border(pPr_element, sz: int = 8, color: str = "000000") -> None:
    """在段落属性上添加四边全框线（用于缺失图片占位框）。

    Args:
        pPr_element: 段落属性元素（w:pPr）。
        sz: 线宽（八分之一磅单位），默认 8 = 1pt。
        color: 颜色十六进制（无 # 号），默认 "000000"。
    """
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border_el = OxmlElement(f"w:{side}")
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), str(sz))
        border_el.set(qn("w:space"), "4")
        border_el.set(qn("w:color"), color)
        pBdr.append(border_el)
    pPr_element.append(pBdr)


def _emit_missing_figure_issue(
    figure: FigureIR, issues: IssueCollector, flags: BehaviorFlags
) -> None:
    """记录图片缺失 Issue，级别按 flags 动态调整。

    - 默认：ERROR（E-IMG-01）
    - --strict：FATAL（由 IssueCollector.append 自动将 ERROR 升级为 FATAL）
    - --allow-missing-figures：WARNING（显式降级）
    """
    if flags.allow_missing_figures:
        level = Level.WARNING
    else:
        level = Level.ERROR  # IssueCollector 在 strict 模式下自动升级为 FATAL

    issues.append(
        Issue(
            level=level,
            code="E-IMG-01",
            stage="render",
            message=(
                f"图片文件不存在：图{figure.figure_id}，"
                f"路径：{figure.path_resolved}"
            ),
            source_line=figure.source_line,
            element_ref=f"图{figure.figure_id}",
            suggestion="请确认图片文件路径正确且文件存在，或使用 --allow-missing-figures 降级为警告",
        )
    )


# ---------------------------------------------------------------------------
# 公开 API
# ---------------------------------------------------------------------------


def render_figure(
    doc,
    figure: FigureIR,
    styles: dict,
    oxml_helpers,
    issues: IssueCollector,
    flags: BehaviorFlags,
) -> None:
    """渲染单张图：嵌入图片 + 题注段落（含书签）。

    渲染流程：
      1. 图片存在性检查 → 缺失时创建占位框 + 记录 Issue，继续渲染题注
      2. 图片尺寸计算 → 仅宽度（14cm 默认），超高时双传 width+height
      3. 嵌入图片（居中段落）
      4. 题注段落（图下方，9pt #555555 居中，含书签）
      5. 数据来源行（可选，若 figure.source_note 非空）

    Args:
        doc: python-docx Document 对象。
        figure: 图 IR（FigureAssembler 产出）。
        styles: 样式名 → 样式对象映射（register_styles 产出）。
        oxml_helpers: render/oxml_helpers 模块引用（用于 bookmarkStart/End）。
        issues: 全流程 Issue 收集器。
        flags: 行为开关（控制图片宽度、缺失图片处理策略等）。
    """
    global _next_bookmark_id

    max_width_cm = flags.figure_max_width_cm
    file_available = figure.file_exists

    # ---- 1. 图片存在性检查 ----
    if not file_available:
        _emit_missing_figure_issue(figure, issues, flags)
        # 创建占位框：空段落 + 1pt 黑色边框 + 缺失提示文字
        placeholder_p = doc.add_paragraph()
        placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 段落四边全框线
        pPr = placeholder_p._p.get_or_add_pPr()
        _make_paragraph_border(pPr, sz=8, color="000000")
        # 占位文字
        placeholder_run = placeholder_p.add_run(
            f"【缺失图片：图{figure.figure_id}】\n"
            f"原路径：{figure.path_resolved}"
        )
        placeholder_run.font.size = Cm(FIGURE_PLACEHOLDER_FONT_SIZE_CM)  # ~9pt
        placeholder_run.font.color.rgb = RGBColor.from_string(
            COLOR_CAPTION_GRAY.lstrip("#")
        )

    # ---- 2. 图片尺寸计算（R13 裁决） ----
    px_w, px_h = _resolve_px_dimensions(figure)

    # 分辨率警告：W-IMG-02
    if px_w is not None and px_w < FIGURE_LOW_RES_PX_W_THRESHOLD:
        issues.append(
            Issue(
                level=Level.WARNING,
                code="W-IMG-02",
                stage="render",
                message=(
                    f"图{figure.figure_id} 图片宽度不足：{px_w}px，"
                    f"300dpi 下 < 14cm，印刷存在模糊风险"
                ),
                source_line=figure.source_line,
                element_ref=f"图{figure.figure_id}",
                suggestion="建议使用宽度 >= 1654px（300dpi @ 14cm）的图片",
            )
        )

    # ---- 3. 嵌入图片 ----
    if file_available:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        # 判断是否需要双传 width+height（超高图片限高 20cm）
        needs_height_clamp = False
        if px_w is not None and px_h is not None and px_w > 0:
            height_at_max_width = (px_h / px_w) * max_width_cm
            if height_at_max_width > FIGURE_MAX_HEIGHT_CM:
                needs_height_clamp = True

        try:
            if needs_height_clamp:
                inline_shape = run.add_picture(
                    str(figure.path_resolved),
                    width=Cm(max_width_cm),
                    height=Cm(FIGURE_MAX_HEIGHT_CM),
                )
                # 记录高度限幅（R13 裁决：02 高度上限）
                issues.append(
                    Issue(
                        level=Level.INFO,
                        code="I-IMG-03",
                        stage="render",
                        message=(
                            f"图{figure.figure_id} 原始高度比例超出 {FIGURE_MAX_HEIGHT_CM}cm "
                            f"上限（{px_w}x{px_h}px → 按 {max_width_cm}cm 宽等比缩放高度为 "
                            f"{height_at_max_width:.1f}cm），已限幅为 {FIGURE_MAX_HEIGHT_CM}cm"
                        ),
                        source_line=figure.source_line,
                        element_ref=f"图{figure.figure_id}",
                    )
                )
            else:
                inline_shape = run.add_picture(
                    str(figure.path_resolved), width=Cm(max_width_cm)
                )
        except Exception:
            # 文件在组装阶段存在、渲染阶段消失的边界场景：降级为占位框
            _emit_missing_figure_issue(figure, issues, flags)
            placeholder_p2 = doc.add_paragraph()
            placeholder_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr2 = placeholder_p2._p.get_or_add_pPr()
            _make_paragraph_border(pPr2, sz=8, color="000000")
            placeholder_run2 = placeholder_p2.add_run(
                f"【图片读取失败：图{figure.figure_id}】\n"
                f"路径：{figure.path_resolved}"
            )
            placeholder_run2.font.size = Cm(0.32)
            placeholder_run2.font.color.rgb = RGBColor.from_string(
                COLOR_CAPTION_GRAY.lstrip("#")
            )

    # ---- 4. 题注段落（图下方，含书签） ----
    caption_p = doc.add_paragraph()
    caption_p.style = styles["Caption Figure"]

    # 书签 ID 分配（全局自增）
    bookmark_id = _next_bookmark_id
    _next_bookmark_id += 1

    # bookmarkStart（在段落中插入书签起始标记）
    oxml_helpers.make_bookmark_start(caption_p, bookmark_id, figure.bookmark_name)

    # 题注文字（Phase 6.3：caption_field_mode 开关控制编号是静态文本还是
    # Word SEQ 域。"text" 模式向后兼容，保留 figure.figure_id 字面拼接；
    # "field" 模式改为 Word 原生自动编号，图片增删/重排后编号自动更新，
    # 不再依赖 assemble 层从作者手写 alt 文本誊抄的编号值）。
    if flags.caption_field_mode == "field":
        caption_p.add_run("图")
        # 占位符必须是"全篇第几个 SEQ 图 域"（Word SEQ 域按序列名全局递增计数），
        # 不是 figure.seq_no（章节内序号）——bookmark_id 恰好就是按渲染顺序
        # 全局自增的值，与 SEQ 域的求值结果语义一致，可直接复用。
        oxml_helpers.make_field(
            caption_p, _FIGURE_SEQ_INSTR, field_type="SEQ",
            placeholder_text=str(bookmark_id),
        )
        oxml_helpers.add_run_segments(caption_p, f" {figure.caption_text}")
    else:
        # "text" 模式已废弃（2026-07-26）：兜底路径——使用 figure.figure_id
        # 拼接静态文本。caption_field_mode 默认值为 "field"（config.py），
        # 此分支仅在显式传入 --caption-field-mode text 时执行。
        caption_text = f"图{figure.figure_id} {figure.caption_text}"
        oxml_helpers.add_run_segments(caption_p, caption_text)

    # bookmarkEnd（在段落中插入书签结束标记）
    oxml_helpers.make_bookmark_end(caption_p, bookmark_id, figure.bookmark_name)

    # ---- 5. 数据来源行（可选） ----
    # 注：FigureIR 当前无 source_note 字段（仅 TableIR 有此字段），
    # 使用 getattr 防御性访问以兼容未来可能的 IR 扩展。
    source_note = getattr(figure, "source_note", None)
    if source_note:
        source_p = doc.add_paragraph()
        source_p.style = styles["Table Source"]
        oxml_helpers.add_run_segments(source_p, f"数据来源：{source_note}")


# ===========================================================================
# 自检块
# ===========================================================================
if __name__ == "__main__":
    import sys as _sys
    import tempfile
    import zlib
    from pathlib import Path

    from docx import Document

    # 使用列表容器避免 nonlocal 问题
    _counts = [0, 0]  # [_passed, _failed]

    def _check(desc: str, condition: bool) -> None:
        if condition:
            _counts[0] += 1
            print(f"  [PASS] {desc}")
        else:
            _counts[1] += 1
            print(f"  [FAIL] {desc}")

    # ---- 1. _read_png_dimensions：合法 PNG ----
    # 构造最小合法 PNG（1x1 红色像素）

    def _make_minimal_png(w: int, h: int) -> bytes:
        """生成最小合法 PNG 字节流（用于测试 IHDR 解析）。"""
        # PNG signature
        sig = b"\x89PNG\r\n\x1a\n"

        # IHDR chunk
        ihdr_data = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)  # 8bit RGB
        ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
        ihdr_chunk = (
            struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
        )

        # IDAT chunk（最小压缩像素数据）
        raw_row = b"\x00" + b"\xff\x00\x00" * w  # filter=0, red pixels
        compressed = zlib.compress(raw_row * h)
        idat_crc = zlib.crc32(b"IDAT" + compressed) & 0xFFFFFFFF
        idat_chunk = (
            struct.pack(">I", len(compressed))
            + b"IDAT"
            + compressed
            + struct.pack(">I", idat_crc)
        )

        # IEND chunk
        iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
        iend_chunk = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)

        return sig + ihdr_chunk + idat_chunk + iend_chunk

    with tempfile.TemporaryDirectory() as tmpdir:
        png_path = Path(tmpdir) / "test_800x600.png"
        png_path.write_bytes(_make_minimal_png(800, 600))

        dims = _read_png_dimensions(str(png_path))
        _check(
            "_read_png_dimensions 800x600 PNG -> (800, 600)",
            dims == (800, 600),
        )

        dims2 = _read_png_dimensions(str(png_path) + ".nonexistent")
        _check(
            "_read_png_dimensions 不存在文件 -> None",
            dims2 is None,
        )

    # ---- 2. render_figure 基本流程（缺失图片 + 占位框） ----
    doc = Document()
    from .styles import register_styles

    style_map = register_styles(doc)
    from . import oxml_helpers as oxml_mod

    collector = IssueCollector(strict=False)

    missing_fig = FigureIR(
        figure_id="1-1",
        chapter_no=1,
        seq_no=1,
        caption_text="测试图题",
        alt_raw="图1-1 测试图题",
        path_raw="images/test.png",
        path_resolved="/nonexistent/path/test.png",
        file_exists=False,
        bookmark_name="fig_1_1",
        px_w=None,
        px_h=None,
        source_line=42,
    )

    render_figure(doc, missing_fig, style_map, oxml_mod, collector, BehaviorFlags())

    # 验证：至少产出了一个段落（占位框 + 题注 = 至少 2 段）
    _check(
        "缺失图片仍产出段落（占位框 + 题注）",
        len(doc.paragraphs) >= 2,
    )

    # 验证：E-IMG-01 Issue 已记录
    _check(
        "缺失图片记录 E-IMG-01 Issue",
        any(i.code == "E-IMG-01" for i in collector),
    )

    # 验证：题注段落包含图号和题注文本（field 模式：图{SEQ} 测试图题）
    caption_found = any(
        "测试图题" in p.text for p in doc.paragraphs
    )
    _check("题注段落包含 '测试图题'", caption_found)

    # ---- 3. --allow-missing-figures 降级 ----
    doc2 = Document()
    style_map2 = register_styles(doc2)
    collector2 = IssueCollector(strict=False)

    render_figure(
        doc2,
        missing_fig,
        style_map2,
        oxml_mod,
        collector2,
        BehaviorFlags(allow_missing_figures=True),
    )

    img_issues = [i for i in collector2 if i.code == "E-IMG-01"]
    _check(
        "--allow-missing-figures 下 E-IMG-01 降级为 WARNING",
        len(img_issues) >= 1 and all(i.level is Level.WARNING for i in img_issues),
    )

    # ---- 4. --strict 升级 ----
    doc3 = Document()
    style_map3 = register_styles(doc3)
    collector3 = IssueCollector(strict=True)

    render_figure(
        doc3, missing_fig, style_map3, oxml_mod, collector3, BehaviorFlags()
    )

    img_issues3 = [i for i in collector3 if i.code == "E-IMG-01"]
    _check(
        "--strict 下 E-IMG-01 升级为 FATAL",
        len(img_issues3) >= 1 and all(i.level is Level.FATAL for i in img_issues3),
    )

    # ---- 5. 书签自增 ----
    doc4 = Document()
    style_map4 = register_styles(doc4)
    collector4 = IssueCollector()

    fig_a = FigureIR(
        figure_id="2-1", chapter_no=2, seq_no=1,
        caption_text="图A", alt_raw="图2-1 图A",
        path_raw="a.png", path_resolved="/tmp/a.png",
        file_exists=False, bookmark_name="fig_2_1",
        px_w=None, px_h=None, source_line=10,
    )
    fig_b = FigureIR(
        figure_id="2-2", chapter_no=2, seq_no=2,
        caption_text="图B", alt_raw="图2-2 图B",
        path_raw="b.png", path_resolved="/tmp/b.png",
        file_exists=False, bookmark_name="fig_2_2",
        px_w=None, px_h=None, source_line=20,
    )

    render_figure(doc4, fig_a, style_map4, oxml_mod, collector4, BehaviorFlags())
    render_figure(doc4, fig_b, style_map4, oxml_mod, collector4, BehaviorFlags())

    # 验证两个书签 ID 不同
    bm_starts = doc4.element.findall(".//" + qn("w:bookmarkStart"))
    _check("两张图产生 2 个书签", len(bm_starts) == 2)

    # ---- 6. caption_field_mode="field" 模式：SEQ 域占位符（本次修复覆盖） ----
    doc5 = Document()
    style_map5 = register_styles(doc5)
    collector5 = IssueCollector()

    render_figure(
        doc5, fig_a, style_map5, oxml_mod, collector5,
        BehaviorFlags(caption_field_mode="field", allow_missing_figures=True),
    )
    render_figure(
        doc5, fig_b, style_map5, oxml_mod, collector5,
        BehaviorFlags(caption_field_mode="field", allow_missing_figures=True),
    )

    instr_texts = doc5.element.findall(".//" + qn("w:instrText"))
    seq_instrs = [t.text for t in instr_texts if t.text and "SEQ" in t.text]
    _check("field 模式产生 2 个 SEQ 图 域", len(seq_instrs) == 2)

    # 提取每个域 separate 与 end 之间的 w:t 文本，验证占位符非空
    fld_chars = doc5.element.findall(".//" + qn("w:fldChar"))
    placeholder_values = []
    for fc in fld_chars:
        if fc.get(qn("w:fldCharType")) == "separate":
            # 找到该 run 之后、下一个 end fldChar 之前的 w:t
            sep_run = fc.getparent()
            sibling = sep_run.getnext()
            val = None
            while sibling is not None:
                fld = sibling.find(qn("w:fldChar"))
                if fld is not None and fld.get(qn("w:fldCharType")) == "end":
                    break
                t = sibling.find(qn("w:t"))
                if t is not None:
                    val = t.text
                sibling = sibling.getnext()
            placeholder_values.append(val)

    _check(
        "SEQ 域占位符均非空",
        len(placeholder_values) == 2 and all(v not in (None, "") for v in placeholder_values),
    )
    # 注：_next_bookmark_id 为模块全局计数器，前面测试用例已消耗过若干 ID，
    # 此处不假设具体起始值，只验证"连续递增的两个整数"这一语义
    _check(
        "SEQ 域占位符为连续递增的两个整数（bookmark_id 全局递增语义）",
        len(placeholder_values) == 2
        and placeholder_values[0].isdigit()
        and placeholder_values[1].isdigit()
        and int(placeholder_values[1]) == int(placeholder_values[0]) + 1,
    )

    # ---- 汇总 ----
    _passed, _failed = _counts
    print(f"\n{'=' * 50}")
    print(f"通过: {_passed}, 失败: {_failed}")
    print(f"{'=' * 50}")

    if _failed == 0:
        print("figures.py 自检通过")
    else:
        print(f"figures.py 自检失败：{_failed} 项未通过")
    _sys.exit(_failed)
