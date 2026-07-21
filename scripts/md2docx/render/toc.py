"""目录与图表目录渲染模块（C-08 render/toc.py）。

将 DocumentIR 中的标题信息通过 Word TOC 域代码 + PAGEREF 域
渲染为目录页与图表目录页（R9 裁决混合方案：文本条目 + PAGEREF 页码域）。

架构约束（00-master-design.md）：
- G-03：所有 Word 域必须通过 oxml_helpers.make_field() 构建，禁止直接拼接域 XML
- G-04：Pt()/Cm() 的 import 仅允许存在于 config.py、render/styles.py、oxml_helpers.py
  三个文件——本文件不 import Pt/Cm，字体/尺寸等 OXML 操作直接经 lxml 元素完成
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ..config import MARGIN_LEFT_CM, MARGIN_RIGHT_CM, PAGE_WIDTH_CM
from ..ir import DocumentIR, FigureIR, TableIR, TableKind

# ---------------------------------------------------------------------------
# 图/表条目排序键（数值排序，避免 "1-10" < "1-2" 的字典序错误）
# ---------------------------------------------------------------------------


def _fig_sort_key(fig: FigureIR) -> tuple[int, int]:
    """按 (章号, 序号) 数值排序 FigureIR。"""
    return (fig.chapter_no, fig.seq_no)


def _tbl_sort_key(tbl: TableIR) -> tuple[int, int]:
    """按 (章号, 序号) 数值排序正文 TableIR。"""
    if tbl.table_id:
        parts = tbl.table_id.split("-")
        return (int(parts[0]), int(parts[1]))
    return (999, 0)  # 无编号的表排末尾（防御：BODY 表总有编号）


# ---------------------------------------------------------------------------
# OXML 辅助（本模块内部，不 import Pt/Cm 以遵守 G-04）
# ---------------------------------------------------------------------------

# 文本区可用宽度：21.0 - 3.17 - 2.54 = 15.29 cm
_TEXT_BODY_WIDTH_CM = PAGE_WIDTH_CM - MARGIN_LEFT_CM - MARGIN_RIGHT_CM
# 制表位位置（twips）：右对齐到文本区右边界
# 1 cm = 567 twips（1 inch = 1440 twips, 1 inch = 2.54 cm）
_RIGHT_TAB_TWIPS = int(_TEXT_BODY_WIDTH_CM * 567)


def _set_run_font(run, size_half_pt: str, color_hex: str, cjk_font: str = "宋体") -> None:
    """通过 OXML 设置单个 run 的字号、颜色与中文字体（遵守 G-04 不 import Pt）。

    Args:
        run: python-docx Run 对象
        size_half_pt: 半磅值字符串（9pt = "18"）
        color_hex: 颜色十六进制（无 # 号），如 "999999"
        cjk_font: 中文字体名，默认 "宋体"
    """
    rPr = run._r.get_or_add_rPr()

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size_half_pt)
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cjk_font)


def _ensure_right_tab_stop(paragraph) -> None:
    """确保段落存在一个右对齐制表位（位于文本区右边界）。

    若段落已有 w:tabs 且在右边界附近已有 right 制表位则不重复添加；
    否则新建 w:tabs 并追加一个 right tab（含点前导符）。
    """
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    else:
        for existing in tabs.findall(qn("w:tab")):
            if existing.get(qn("w:val")) == "right":
                pos_str = existing.get(qn("w:pos"))
                if pos_str and abs(int(pos_str) - _RIGHT_TAB_TWIPS) <= 50:
                    return  # 已存在，不重复添加
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(_RIGHT_TAB_TWIPS))
    tab.set(qn("w:leader"), "dot")
    tabs.append(tab)


def _add_field_with_placeholder(paragraph, instr_text: str, field_type: str,
                                 placeholder_text: str, oxml_helpers) -> None:
    """向段落追加一个 Word 域并插入占位文本（位于 separate 与 end 之间）。

    oxml_helpers.make_field() 创建四态结构（begin/instrText/separate/end），
    但不在 separate 与 end 之间放占位文本。本函数在 separate 后、end 前
    插入一个 <w:r><w:t>placeholder</w:t></w:r>。

    Args:
        paragraph: python-docx Paragraph 对象
        instr_text: 域指令（如 'PAGEREF fig_1_1 \\h'）
        field_type: 域类型标识（'PAGEREF' / 'TOC'）
        placeholder_text: 占位文本（PAGEREF 用 "1"，TOC 可用 ""）
        oxml_helpers: oxml_helpers 模块
    """
    oxml_helpers.make_field(paragraph, instr_text, field_type=field_type)

    if not placeholder_text:
        return

    # make_field 追加了 4 个 run：begin, instrText, separate, end
    # 需要在 separate（倒数第2个）和 end（倒数第1个）之间插入占位 run
    runs = paragraph._p.findall(qn("w:r"))
    if len(runs) < 4:
        return
    end_run = runs[-1]

    placeholder_run = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = placeholder_text
    placeholder_run.append(t_elem)

    paragraph._p.remove(end_run)
    paragraph._p.append(placeholder_run)
    paragraph._p.append(end_run)


# ---------------------------------------------------------------------------
# 公开 API
# ---------------------------------------------------------------------------


def render_toc(doc, ir: DocumentIR, styles: dict, oxml_helpers) -> None:
    """在 Document 当前节中渲染目录页。

    渲染流程：
        1. "目录"标题（TOC Heading 样式，非 Heading 系——不入 TOC、不占编号）
        2. Word TOC 域代码（\\o "1-3" \\h \\z \\u）
        3. 更新提示（9pt 灰色文字）
        4. 图表目录（仅在图表总数 >= 10 时生成，含分页 + PAGEREF 混合方案）

    Args:
        doc: python-docx Document 对象
        ir: DocumentIR 实例（从中提取 FigureIR / TableIR 列表）
        styles: name->style 查找字典（render/styles.py 的 register_styles() 产出）
        oxml_helpers: render/oxml_helpers 模块
    """
    toc_heading_style = styles.get("TOC Heading")

    # ---- 从 elements 中提取图与正文表 ----
    figures: list[FigureIR] = [
        e for e in ir.elements if isinstance(e, FigureIR)
    ]
    body_tables: list[TableIR] = [
        e for e in ir.elements
        if isinstance(e, TableIR) and e.kind == TableKind.BODY and e.table_id is not None
    ]

    # ---- 1. 目录标题 ----
    doc.add_paragraph("目录", style=toc_heading_style)

    # ---- 2. TOC 域 ----
    p_toc = doc.add_paragraph()
    oxml_helpers.make_field(
        p_toc, r'TOC \o "1-3" \h \z \u', field_type="TOC"
    )

    # ---- 3. 更新提示 ----
    p_hint = doc.add_paragraph()
    hint_text = (
        "（打开文档后，右键此处选择"
        "“更新域”或按 F9 键更新目录；"
        "WPS 用户：右键→更新域）"
    )
    hint_run = p_hint.add_run(hint_text)
    _set_run_font(hint_run, size_half_pt="18", color_hex="999999", cjk_font="宋体")

    # ---- 4. 图表目录（仅在图表 >= 10 时生成，R9 auto 模式） ----
    total_charts = len(figures) + len(body_tables)
    if total_charts >= 10:
        render_chart_directory(doc, figures, body_tables, toc_heading_style, oxml_helpers)


def should_render_chart_directory(ir: DocumentIR) -> bool:
    """检查是否需要渲染图表目录（图表总数 >= 10，R9 auto 模式）。"""
    figures = [e for e in ir.elements if isinstance(e, FigureIR)]
    body_tables = [
        e for e in ir.elements
        if isinstance(e, TableIR) and e.kind == TableKind.BODY and e.table_id is not None
    ]
    return (len(figures) + len(body_tables)) >= 10


def render_chart_directory(doc, figures: list, body_tables: list,
                           toc_heading_style, oxml_helpers) -> None:
    """渲染图表目录（图目录 + 表目录，含 PAGEREF 域）。

    调用方负责在调用本函数前插入分页符（G-02：add_page_break 仅 document.py）。
    """
    # 图表目录标题
    doc.add_paragraph("图表目录", style=toc_heading_style)

    # ---- 4a. 图目录条目 ----
    for fig in sorted(figures, key=_fig_sort_key):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

        # 图号文本 + 制表符
        label = f"图{fig.figure_id}\t"
        p.add_run(label)

        # PAGEREF 域（占位 "1"，Word 更新域后回填真实页码）
        instr = f"PAGEREF {fig.bookmark_name} \\h"
        _add_field_with_placeholder(p, instr, "PAGEREF", "1", oxml_helpers)

        # 图题注文本
        if fig.caption_text:
            p.add_run(f"  {fig.caption_text}")

        _ensure_right_tab_stop(p)

    # ---- 4b. 表目录条目 ----
    for tbl in sorted(body_tables, key=_tbl_sort_key):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

        # 表号文本 + 制表符
        label = f"表{tbl.table_id}\t"
        p.add_run(label)

        # PAGEREF 域
        instr = f"PAGEREF {tbl.bookmark_name} \\h"
        _add_field_with_placeholder(p, instr, "PAGEREF", "1", oxml_helpers)

        # 表题注文本
        if tbl.caption_text:
            p.add_run(f"  {tbl.caption_text}")

        _ensure_right_tab_stop(p)


# ===========================================================================
# 自检（验收标准）
# ===========================================================================
if __name__ == "__main__":
    import sys

    from docx import Document

    from ..ir import (
        DocumentIR,
        FigureIR,
        HeadingIR,
        HeadingKind,
        MetadataIR,
        SectionKind,
        SectionPlan,
        SectionSpec,
        TableIR,
        TableKind,
        PageNumFormat,
        HeaderMode,
    )
    from . import oxml_helpers as _oxml

    _failed = 0
    _passed = 0

    def _check(desc: str, condition: bool, detail: str = "") -> None:
        global _passed, _failed
        if condition:
            _passed += 1
            print(f"  [PASS] {desc}")
        else:
            _failed += 1
            print(f"  [FAIL] {desc}  -- {detail}")
            if detail:
                print(f"         详情: {detail}")

    # ---- 构造最小 IR（图表数不足 10，不触发图表目录） ----
    print("\n=== 测试1：图表 < 10 -> 仅目录，无图表目录 ===")
    meta = MetadataIR(
        title="测试报告", subtitle=None, report_type=None,
        organization=None, version_raw="V1.0", version="V1.0",
        date="2026年7月", title_short="测试报告",
    )
    h1 = HeadingIR(
        kind=HeadingKind.CHAPTER, raw_text="第一章 概述", text="概述",
        number="一", display_number="第一章", source_line=10,
    )
    fig1 = FigureIR(
        figure_id="1-1", chapter_no=1, seq_no=1,
        caption_text="测试图", alt_raw="图1-1 测试图",
        path_raw="test.png", path_resolved="/tmp/test.png",
        file_exists=False, bookmark_name="fig_1_1",
        px_w=None, px_h=None, source_line=20,
    )
    section_plan = SectionPlan(sections=[
        SectionSpec(kind=SectionKind.COVER, page_num_fmt=PageNumFormat.NONE,
                     page_num_restart=False, header_mode=HeaderMode.NONE,
                     start_element_index=0),
        SectionSpec(kind=SectionKind.TOC, page_num_fmt=PageNumFormat.LOWER_ROMAN,
                     page_num_restart=True, header_mode=HeaderMode.NONE,
                     start_element_index=1),
        SectionSpec(kind=SectionKind.BODY, page_num_fmt=PageNumFormat.DECIMAL,
                     page_num_restart=True, header_mode=HeaderMode.TITLE_SHORT,
                     start_element_index=2),
    ])
    ir_small = DocumentIR(
        metadata=meta,
        elements=[h1, fig1],
        section_plan=section_plan,
        figure_registry={"1-1": fig1},
        table_registry={},
        xref_registry=[],
    )

    doc_small = Document()
    from .styles import register_styles
    style_map_small = register_styles(doc_small)

    render_toc(doc_small, ir_small, style_map_small, _oxml)

    para_count = len(doc_small.paragraphs)
    _check("图表<10时应至少有3个段落", para_count >= 3, f"实际 {para_count}")
    chart_toc_found = any(
        "图表目录" in p.text for p in doc_small.paragraphs
    )
    _check("图表<10时不应生成图表目录", not chart_toc_found,
           "不应出现'图表目录'")

    # ---- 构造 IR（图表数 >= 10，触发图表目录） ----
    print("\n=== 测试2：图表 >= 10 -> 生成图表目录 ===")
    many_figures: list = []
    fig_registry: dict = {}
    for i in range(1, 9):  # 8 张图
        fid = f"1-{i}"
        f = FigureIR(
            figure_id=fid, chapter_no=1, seq_no=i,
            caption_text=f"测试图{i}",
            alt_raw=f"图{fid} 测试图{i}",
            path_raw=f"test{i}.png", path_resolved=f"/tmp/test{i}.png",
            file_exists=False, bookmark_name=f"fig_1_{i}",
            px_w=None, px_h=None, source_line=20 + i,
        )
        many_figures.append(f)
        fig_registry[fid] = f

    many_tables: list = []
    tbl_registry: dict = {}
    for i in range(1, 5):  # 4 张正文表
        tid = f"2-{i}"
        t = TableIR(
            kind=TableKind.BODY, table_id=tid,
            caption_text=f"测试表{i}", source_note=None,
            header_cells=[], body_rows=[],
            n_cols=2, bookmark_name=f"tbl_2_{i}",
            source_line=100 + i,
        )
        many_tables.append(t)
        tbl_registry[tid] = t

    total = len(many_figures) + len(many_tables)
    _check(f"图表总数={total}（需>=10才生成图表目录）", total >= 10, f"实际 {total}")

    ir_large = DocumentIR(
        metadata=meta,
        elements=many_figures + many_tables,
        section_plan=section_plan,
        figure_registry=fig_registry,
        table_registry=tbl_registry,
        xref_registry=[],
    )

    doc_large = Document()
    style_map_large = register_styles(doc_large)
    render_toc(doc_large, ir_large, style_map_large, _oxml)

    # 验证：应有"图表目录"标题
    chart_toc_paras = [
        p for p in doc_large.paragraphs if "图表目录" in p.text
    ]
    _check("图表>=10时应生成'图表目录'标题", len(chart_toc_paras) >= 1,
           f"找到 {len(chart_toc_paras)} 个")

    # 验证：应包含 PAGEREF 域（instrText 嵌套在 w:r 内，需递归搜索）
    pageref_count = 0
    for p in doc_large.paragraphs:
        for instr in p._p.iter(qn("w:instrText")):
            text = instr.text or ""
            if "PAGEREF" in text:
                pageref_count += 1
    _check(f"PAGEREF 域数量 = 图表数 {total}", pageref_count == total,
           f"实际 {pageref_count}")

    # 验证：TOC 域指令匹配（instrText 嵌套在 w:r 内，需递归搜索）
    toc_instr_found = False
    for p in doc_large.paragraphs:
        for instr in p._p.iter(qn("w:instrText")):
            text = instr.text or ""
            if "TOC" in text:
                # make_field 会包裹空格：' TOC \o "1-3" \h \z \u '
                expected = r' TOC \o "1-3" \h \z \u '
                _check(
                    "TOC 域指令包含 \\o \"1-3\" \\h \\z \\u",
                    r'\o "1-3" \h \z \u' in text,
                    f"实际: {text!r}",
                )
                toc_instr_found = True
    _check("至少存在一个 TOC 域", toc_instr_found)

    # 验证：更新提示可见
    hint_found = any("F9" in p.text for p in doc_large.paragraphs)
    _check("更新提示可见（含'F9'）", hint_found)

    # ---- 测试3：排序键函数 ----
    print("\n=== 测试3：排序键 ===")
    fig_a = FigureIR(
        figure_id="1-10", chapter_no=1, seq_no=10,
        caption_text="", alt_raw="", path_raw="", path_resolved="",
        file_exists=False, bookmark_name="fig_1_10",
        px_w=None, px_h=None, source_line=1,
    )
    fig_b = FigureIR(
        figure_id="1-2", chapter_no=1, seq_no=2,
        caption_text="", alt_raw="", path_raw="", path_resolved="",
        file_exists=False, bookmark_name="fig_1_2",
        px_w=None, px_h=None, source_line=2,
    )
    _check("排序：1-2 < 1-10（数值排序）",
           _fig_sort_key(fig_b) < _fig_sort_key(fig_a))

    # ---- 汇总 ----
    print(f"\n{'=' * 50}")
    print(f"通过: {_passed}, 失败: {_failed}")
    print(f"{'=' * 50}")

    if _failed == 0:
        print("render/toc.py 自检通过！")
    else:
        print(f"render/toc.py 自检失败：{_failed} 项未通过")
        sys.exit(1)
