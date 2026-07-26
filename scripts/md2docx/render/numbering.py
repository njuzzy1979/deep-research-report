"""章节/附录多级列表编号模块（Phase 6.2）。

负责在 ``doc.part.numbering_part`` 上定义 Word 原生"多级列表"编号
（``w:abstractNum`` + ``w:num``），并提供段落级/样式级 ``w:numPr`` 绑定辅助函数，
使章节标题编号由 Word 在打开/打印/F9 更新域时自动计算与渲染，而非
``assemble/headings.py`` 算好字符串后由 render 层拼接成静态文本。

背景与设计取舍（务必先读，避免重复踩坑）：
    空白 ``docx.Document()`` 默认已自带 9 个 ``abstractNum``（id 0~8）与 9 个
    ``num``（id 1~9，供 List Bullet / List Number 等内置样式使用）。本模块新增
    的 ID **必须**从 ``abstractNumId=9``、``numId=10`` 起，避免与默认定义冲突
    （已用真实 python-docx 1.2.0 环境核实）。

    章节编号是否应使用中文数字（"第一章"）？—— 已核实这是一条**真实的 OOXML
    渲染约束，不是我们代码的 bug**：``lvlText`` 模板里的 ``%N`` 占位符，渲染时
    永远按"第 N 级自身的 ``w:numFmt``"取值，而不是按引用它的那一级的
    ``w:numFmt``。也就是说，如果级别 0（章）用
    ``w:numFmt val="chineseCountingThousand"``，那么级别 1（节）的
    ``lvlText="%1.%2"`` 中的 ``%1`` 也必然渲染成中文数字"一"，产出"一.1"而非
    期望的"1.1"——两级无法共用同一个计数器又要求不同数字体系。若拆成两个独立
    ``abstractNum``，又会失去"新增/删除一章时，节编号的章前缀自动跟着变"这一
    动态更新能力（这正是用户提出本需求的核心诉求）。
    因此本模块的章节体系统一使用 ``numFmt="decimal"``，章标题以
    ``lvlText="第%1章"`` 渲染为"第1章"（阿拉伯数字），节/小节渲染为"1.1"
    "1.1.1"——三级共享同一个 ``abstractNum`` 的同一个计数器，保证联动更新，
    这是"自动更新能力"这一验收标准的硬约束下唯一可行的方案。

    附录字母编号（"附录A"）不能与章节共用同一个列表——章节计数器与附录字母
    计数器是两套独立序列，且附录标题也复用 Heading 1 样式（见
    ``render/headings.py::_KIND_TO_LEVEL``），若不显式覆盖会被 Heading 1 样式
    绑定的章节列表"污染"。故附录单独定义一个 1 级 ``abstractNum``
    （``numFmt="upperLetter"``，``lvlText="附录%1"``），并在
    ``render/headings.py`` 中对 APPENDIX 标题显式覆盖为该 numId。

    ABSTRACT / FRONT_MATTER / MAIN_TITLE / PLAIN 四类标题复用 Heading 1/2
    样式但不应显示编号——OOXML 官方约定 ``w:numId w:val="0"`` 表示"关闭编号"，
    对这四类标题在段落级显式设置 numId=0 即可覆盖样式级的默认绑定。
"""
from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# ID 分配（避开 python-docx 空白 Document 默认占用的 abstractNumId 0-8 / numId 1-9）
# ---------------------------------------------------------------------------

CHAPTER_ABSTRACT_NUM_ID = 9   # 章/节/小节 3 级列表
APPENDIX_ABSTRACT_NUM_ID = 10  # 附录字母列表（单级）

CHAPTER_NUM_ID = 10   # 绑定 Heading 1(章)/Heading 2(节)/Heading 3(小节)
APPENDIX_NUM_ID = 11  # 绑定附录标题（覆盖 Heading 1 默认的 CHAPTER_NUM_ID）

NO_NUMBERING_ID = 0  # OOXML 官方"关闭编号"值

# ---------------------------------------------------------------------------
# lvlText 模板常量（本次 plan §6.2 章节/附录编号方案，见上方模块 docstring
# "章节编号是否应使用中文数字"一节的完整推导）
# ---------------------------------------------------------------------------
_CHAPTER_LVL0_TEXT = "第%1章"  # 章标题模板：decimal 数字体系渲染为"第1章"
_APPENDIX_LVL0_TEXT = "附录%1"  # 附录标题模板：upperLetter 数字体系渲染为"附录A"


def _make_lvl(ilvl: int, num_fmt: str, lvl_text: str, start: int = 1) -> OxmlElement:
    """构建单个 ``<w:lvl>`` 元素（多级列表的一个层级定义）。"""
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))

    start_el = OxmlElement("w:start")
    start_el.set(qn("w:val"), str(start))
    lvl.append(start_el)

    fmt_el = OxmlElement("w:numFmt")
    fmt_el.set(qn("w:val"), num_fmt)
    lvl.append(fmt_el)

    text_el = OxmlElement("w:lvlText")
    text_el.set(qn("w:val"), lvl_text)
    lvl.append(text_el)

    jc_el = OxmlElement("w:lvlJc")
    jc_el.set(qn("w:val"), "left")
    lvl.append(jc_el)

    return lvl


def _make_abstract_num(abstract_num_id: int, levels: list[OxmlElement]) -> OxmlElement:
    """构建 ``<w:abstractNum>`` 元素，包裹给定的若干 ``<w:lvl>`` 子元素。"""
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(
        qn("w:val"), "multilevel" if len(levels) > 1 else "singleLevel"
    )
    abstract_num.append(multi_level)

    for lvl in levels:
        abstract_num.append(lvl)

    return abstract_num


def _make_num(num_id: int, abstract_num_id: int) -> OxmlElement:
    """构建 ``<w:num>`` 元素，将具体 numId 关联到某个 abstractNumId。"""
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))

    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)

    return num


def ensure_numbering_defs(doc) -> None:
    """在 ``doc.part.numbering_part`` 上幂等地写入章节/附录两套编号定义。

    幂等性：若 ``CHAPTER_ABSTRACT_NUM_ID`` 对应的 ``abstractNum`` 已存在
    （例如同一 doc 对象被重复调用），直接跳过，不重复插入。

    插入顺序遵循 OOXML ``CT_Numbering`` schema 的元素序列约束——全部
    ``w:abstractNum`` 必须出现在全部 ``w:num`` 之前：新增的 abstractNum
    插在"最后一个已存在 abstractNum"之后、"第一个已存在 num"之前；
    新增的 num 追加在最末尾。

    Args:
        doc: python-docx Document 对象。
    """
    numbering_el = doc.part.numbering_part.element

    # 幂等性检查
    existing_ids = {
        el.get(qn("w:abstractNumId"))
        for el in numbering_el.findall(qn("w:abstractNum"))
    }
    if str(CHAPTER_ABSTRACT_NUM_ID) in existing_ids:
        return

    # ---- 章/节/小节 3 级列表（abstractNumId=9）----
    # 三级统一 numFmt="decimal"：lvlText 的 %N 占位符渲染时固定按第 N 级自身
    # 的 numFmt 取值（OOXML 规则，见模块 docstring），故不能让级别 0 使用
    # chineseCountingThousand 而级别 1+ 想要 arabic——三级必须同一数字体系
    # 才能共享同一个计数器实现"新增/删除章节后编号自动联动"。
    chapter_levels = [
        _make_lvl(0, "decimal", _CHAPTER_LVL0_TEXT),
        _make_lvl(1, "decimal", "%1.%2"),
        _make_lvl(2, "decimal", "%1.%2.%3"),
    ]
    chapter_abstract = _make_abstract_num(CHAPTER_ABSTRACT_NUM_ID, chapter_levels)

    # ---- 附录字母列表（abstractNumId=10，单级）----
    appendix_levels = [
        _make_lvl(0, "upperLetter", _APPENDIX_LVL0_TEXT),
    ]
    appendix_abstract = _make_abstract_num(APPENDIX_ABSTRACT_NUM_ID, appendix_levels)

    # ---- 插入 abstractNum：位于最后一个已存在 abstractNum 之后 ----
    existing_abstract_nums = numbering_el.findall(qn("w:abstractNum"))
    if existing_abstract_nums:
        last_abstract = existing_abstract_nums[-1]
        last_abstract.addnext(appendix_abstract)
        last_abstract.addnext(chapter_abstract)
    else:
        numbering_el.insert(0, appendix_abstract)
        numbering_el.insert(0, chapter_abstract)

    # ---- num：追加在最末尾 ----
    chapter_num = _make_num(CHAPTER_NUM_ID, CHAPTER_ABSTRACT_NUM_ID)
    appendix_num = _make_num(APPENDIX_NUM_ID, APPENDIX_ABSTRACT_NUM_ID)
    numbering_el.append(chapter_num)
    numbering_el.append(appendix_num)


def _get_or_add_numPr(pPr) -> OxmlElement:
    """从给定 ``<w:pPr>`` 中获取或创建 ``<w:numPr>`` 子元素。"""
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    return numPr


def set_heading_numPr(paragraph, ilvl: int, num_id: int) -> None:
    """在段落上设置 ``w:numPr``（段落级直接格式，覆盖样式级绑定）。

    ``num_id=0``（:data:`NO_NUMBERING_ID`）是 OOXML 官方"关闭编号"值，
    用于让不需要编号的标题类型（ABSTRACT/FRONT_MATTER/MAIN_TITLE/PLAIN）
    显式覆盖 Heading 1/2 样式上绑定的默认章节列表。

    Args:
        paragraph: python-docx Paragraph 对象。
        ilvl: 列表层级（0 起始）。
        num_id: 目标 numId（对应 :func:`ensure_numbering_defs` 定义的列表，
            或 0 表示关闭编号）。
    """
    pPr = paragraph._p.get_or_add_pPr()
    numPr = _get_or_add_numPr(pPr)

    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        numId_el = OxmlElement("w:numId")
        numPr.append(numId_el)
    numId_el.set(qn("w:val"), str(num_id))


def bind_style_numPr(style, ilvl: int, num_id: int) -> None:
    """在 Word 命名样式（如 Heading 1/2/3）上绑定默认 ``w:numPr``。

    绑定到样式级而非逐段落设置的好处：任何应用该样式的段落，只要没有
    自己的段落级 ``w:numPr`` 覆盖，就会自动继承样式定义的编号列表与层级——
    这正是 Word 内置"将多级列表关联到标题样式"功能的 OOXML 落地方式，
    CHAPTER/SECTION/SUBSECTION 三类标题因此不需要在 render/headings.py
    中逐段落设置 numPr。

    Args:
        style: python-docx Style 对象（如 ``doc.styles["Heading 1"]``）。
        ilvl: 列表层级（0 起始）。
        num_id: 目标 numId。
    """
    pPr = style.element.get_or_add_pPr()
    numPr = _get_or_add_numPr(pPr)

    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    numId_el = numPr.find(qn("w:numId"))
    if numId_el is None:
        numId_el = OxmlElement("w:numId")
        numPr.append(numId_el)
    numId_el.set(qn("w:val"), str(num_id))


# ===========================================================================
# 自检块
# ===========================================================================
if __name__ == "__main__":
    from docx import Document

    doc = Document()
    ensure_numbering_defs(doc)

    numbering_el = doc.part.numbering_part.element

    # 1. 验证新增的两个 abstractNum 存在且 id 正确（9、10）
    abstract_ids = {
        el.get(qn("w:abstractNumId"))
        for el in numbering_el.findall(qn("w:abstractNum"))
    }
    assert str(CHAPTER_ABSTRACT_NUM_ID) in abstract_ids
    assert str(APPENDIX_ABSTRACT_NUM_ID) in abstract_ids
    # 默认 0-8 仍应保留（未被覆盖/删除）
    for i in range(9):
        assert str(i) in abstract_ids, f"默认 abstractNumId={i} 被意外破坏"

    # 2. 验证新增的两个 num 存在且指向正确的 abstractNumId（10、11）
    num_map = {
        el.get(qn("w:numId")): el.find(qn("w:abstractNumId")).get(qn("w:val"))
        for el in numbering_el.findall(qn("w:num"))
    }
    assert num_map[str(CHAPTER_NUM_ID)] == str(CHAPTER_ABSTRACT_NUM_ID)
    assert num_map[str(APPENDIX_NUM_ID)] == str(APPENDIX_ABSTRACT_NUM_ID)
    for i in range(1, 10):
        assert str(i) in num_map, f"默认 numId={i} 被意外破坏"

    # 3. 幂等性：重复调用不应产生第二份 abstractNumId=9
    ensure_numbering_defs(doc)
    abstract_ids_2 = [
        el.get(qn("w:abstractNumId"))
        for el in numbering_el.findall(qn("w:abstractNum"))
    ]
    assert abstract_ids_2.count(str(CHAPTER_ABSTRACT_NUM_ID)) == 1, "幂等性校验失败：重复插入"

    # 4. 验证 set_heading_numPr 正确写入段落 numPr
    p = doc.add_paragraph("测试段落")
    set_heading_numPr(p, ilvl=1, num_id=CHAPTER_NUM_ID)
    pPr = p._p.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr"))
    assert numPr is not None
    assert numPr.find(qn("w:ilvl")).get(qn("w:val")) == "1"
    assert numPr.find(qn("w:numId")).get(qn("w:val")) == str(CHAPTER_NUM_ID)

    # 5. 验证 bind_style_numPr 正确写入样式 numPr
    h1 = doc.styles["Heading 1"]
    bind_style_numPr(h1, ilvl=0, num_id=CHAPTER_NUM_ID)
    style_pPr = h1.element.find(qn("w:pPr"))
    style_numPr = style_pPr.find(qn("w:numPr"))
    assert style_numPr is not None
    assert style_numPr.find(qn("w:ilvl")).get(qn("w:val")) == "0"
    assert style_numPr.find(qn("w:numId")).get(qn("w:val")) == str(CHAPTER_NUM_ID)

    # 6. 验证 NO_NUMBERING_ID=0 可正确写入（关闭编号场景）
    p2 = doc.add_paragraph("摘要标题测试")
    set_heading_numPr(p2, ilvl=0, num_id=NO_NUMBERING_ID)
    numPr2 = p2._p.find(qn("w:pPr")).find(qn("w:numPr"))
    assert numPr2.find(qn("w:numId")).get(qn("w:val")) == "0"

    print("numbering.py 自检通过：abstractNum/num 定义、段落级/样式级 numPr 绑定均正常")
