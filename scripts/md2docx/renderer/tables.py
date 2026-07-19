"""表格渲染 — 全框线含竖线 + 交替行灰底 + 跨页表头重复。"""
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from .helpers import (set_run_font, set_para_spacing, remove_mark,
                      set_cell_border, set_cell_shading)


def render_table(doc, elem, caption, config):
    fonts = config.fonts
    sizes = config.sizes
    colors = config.colors
    tbl_cfg = config.table

    # 表题注（V3.0 §4.2: 表上方居中，9pt #555555）
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, ls=1.0, before=Pt(6), after=Pt(3))
        cap_text = caption if caption.startswith('表') else f'表 {caption}'
        r = p.add_run(cap_text)
        set_run_font(r, fonts.body.cjk, fonts.body.latin,
                     sizes.caption, color=colors.as_tuple('secondary'))

    rc = len(elem.rows) + 1
    cc = len(elem.header)
    tbl = doc.add_table(rows=rc, cols=cc)
    tbl.autofit = True
    tbl.alignment = 1

    outer_sz = str(int(tbl_cfg.border_outer * 8))
    sep_sz = str(int(tbl_cfg.border_header_sep * 8))
    inner_sz = str(int(tbl_cfg.border_inner * 8))

    thick_outer = {'sz': outer_sz, 'val': 'single', 'color': '000000'}
    medium_sep = {'sz': sep_sz, 'val': 'single', 'color': '000000'}
    thin_inner = {'sz': inner_sz, 'val': 'single', 'color': '000000'}

    # 表头行
    for j, ct in enumerate(elem.header):
        c = tbl.rows[0].cells[j]
        c.text = ''
        r = c.paragraphs[0].add_run(remove_mark(ct))
        set_run_font(r, fonts.table_header.cjk, fonts.table_header.latin,
                     sizes.table_header, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(c, top=thick_outer, bottom=medium_sep,
                        left=thin_inner, right=thin_inner,
                        insideH=thin_inner, insideV=thin_inner)

    # 表体行
    for i, row in enumerate(elem.rows):
        is_even = (i % 2 == 1)
        for j, ct in enumerate(row):
            if j >= cc:
                continue
            c = tbl.rows[i + 1].cells[j]
            c.text = ''
            cell_text = remove_mark(ct)
            is_num = re.match(r'^[\d,.\-%$¥€£]+$', cell_text.strip())
            c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.RIGHT if is_num
                                         else WD_ALIGN_PARAGRAPH.LEFT)
            r = c.paragraphs[0].add_run(cell_text)
            set_run_font(r, fonts.table_body.cjk, fonts.table_body.latin,
                         sizes.table_body)
            if is_even:
                set_cell_shading(c, tbl_cfg.alt_row_color)
            is_last = (i == len(elem.rows) - 1)
            cell_bottom = thick_outer if is_last else thin_inner
            set_cell_border(c, top=thin_inner, bottom=cell_bottom,
                            left=thin_inner, right=thin_inner,
                            insideH=thin_inner, insideV=thin_inner)

    # 跨页表头重复
    if tbl_cfg.repeat_header:
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        tblHeaderEl = OxmlElement('w:tblHeader')
        trPr.append(tblHeaderEl)

    p = doc.add_paragraph()
    set_para_spacing(p, ls=1.0)
