"""
通用渲染辅助函数。

从 scripts/markdown_to_docx.py 提取并重构：所有函数通过 config 参数获取字体/字号/颜色，
不再使用模块级硬编码常量。

原始来源: markdown_to_docx.py lines 325-387 (set_run_font, set_para_spacing, add_inline,
              remove_mark, _make_xml_para, _get_text, _has_image, _get_table_header)
"""

import re
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_run_font(run, cjk, latin, size, bold=False, italic=False, color=None):
    """设置 run 的中英文字体、字号、粗细、斜体、颜色。

    原始来源: markdown_to_docx.py L325-334（已验证稳定，直接搬移）
    """
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)


def set_para_spacing(p, ls=1.5, before=Pt(0), after=Pt(0), first=None, left=None):
    """设置段落间距、行距、首行缩进、左缩进。

    原始来源: markdown_to_docx.py L336-340（已验证稳定，直接搬移）
    """
    pf = p.paragraph_format
    pf.line_spacing = ls
    pf.space_before = before
    pf.space_after = after
    if first:
        pf.first_line_indent = first
    if left:
        pf.left_indent = left


def add_inline(para, text, cjk, latin, size, mono_cjk='宋体', mono_latin='Consolas', mono_size=10.5):
    """解析内联 Markdown 标记（粗体/斜体/代码/链接）并渲染到段落。

    原始来源: markdown_to_docx.py L342-354（重构为接受 mono 参数）
    """
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)|'    # 粗体
        r'(\*(.+?)\*)|'        # 斜体
        r'(`(.+?)`)|'          # 行内代码
        r'(\[(.+?)\]\((.+?)\))' # 链接
    )
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            r = para.add_run(text[last:m.start()])
            set_run_font(r, cjk, latin, size)
        if m.group(2):
            r = para.add_run(m.group(2))
            set_run_font(r, cjk, latin, size, bold=True)
        elif m.group(4):
            r = para.add_run(m.group(4))
            set_run_font(r, cjk, latin, size, italic=True)
        elif m.group(6):
            r = para.add_run(m.group(6))
            set_run_font(r, mono_cjk, mono_latin, mono_size)
        elif m.group(8):
            r = para.add_run(m.group(8))
            set_run_font(r, cjk, latin, size)
            r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            r.underline = True
        last = m.end()
    if last < len(text):
        r = para.add_run(text[last:])
        set_run_font(r, cjk, latin, size)


def remove_mark(s):
    """移除内联 Markdown 标记，返回纯文本。

    原始来源: markdown_to_docx.py L356-357（直接搬移）
    """
    return re.sub(
        r'\*\*(.+?)\*\*', r'\1',
        re.sub(r'\*(.+?)\*', r'\1',
               re.sub(r'`(.+?)`', r'\1', s))
    )


def make_xml_para(text, cjk='宋体', latin='Times New Roman', size_pt=9,
                  bold=False, align='center', after=300, color_hex=None):
    """生成纯 XML 段落（用于后处理阶段的题注插入）。

    原始来源: markdown_to_docx.py L359-376（直接搬移）
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), str(after))
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), cjk)
    rf.set(qn('w:ascii'), latin)
    rf.set(qn('w:hAnsi'), latin)
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement('w:b'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    if color_hex:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), color_hex)
        rPr.append(color_el)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p


def get_para_text(elem):
    """获取 OOXML 段落元素中的纯文本内容。

    原始来源: markdown_to_docx.py L378-379（直接搬移）
    """
    return ''.join(t.text or '' for t in elem.iter(qn('w:t'))).strip()


def has_image(para):
    """判断段落是否包含图片。

    原始来源: markdown_to_docx.py L381-382（直接搬移）
    """
    return para.find(
        './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
    ) is not None


def get_table_header_text(tbl_elem):
    """获取 OOXML 表格元素的表头文本（用于后处理表题注匹配）。

    原始来源: markdown_to_docx.py L384-387（直接搬移）
    """
    rows = list(tbl_elem.iter(qn('w:tr')))
    if not rows:
        return ''
    return ' '.join(
        ''.join(t.text or '' for t in tc.iter(qn('w:t')))
        for tc in rows[0].iter(qn('w:tc'))
    )


def cm_to_twips(cm_val):
    """将厘米转换为 twips (1/20 pt)"""
    return int(cm_val * 567)


def pt_to_half_pt(pt_val):
    """将 pt 转换为 half-points（OOXML w:sz 单位）"""
    return int(pt_val * 2)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None,
                    insideH=None, insideV=None):
    """设置单元格边框。

    原始来源: markdown_to_docx.py L653-672（直接搬移）
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge, spec in [('top', top), ('bottom', bottom), ('left', left),
                        ('right', right), ('insideH', insideH), ('insideV', insideV)]:
        if spec is None:
            continue
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), spec.get('val', 'single'))
        el.set(qn('w:sz'), spec.get('sz', '4'))
        el.set(qn('w:color'), spec.get('color', '000000'))
        el.set(qn('w:space'), '0')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color_hex):
    """设置单元格底色。

    原始来源: markdown_to_docx.py L674-682（直接搬移）
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)
