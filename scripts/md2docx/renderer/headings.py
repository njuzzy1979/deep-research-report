"""
标题编号与渲染。

原始来源: markdown_to_docx.py L394-404 (ChapterNum) + L597-646
"""

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .helpers import set_run_font, set_para_spacing, remove_mark


class HeadingNumbering:
    """标题编号器。支持 1-99 章中文数字编号。"""

    CN_MAP = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
        16: '十六', 17: '十七', 18: '十八', 19: '十九', 20: '二十',
    }

    def __init__(self):
        self.ch = 0
        self.sec = {}
        self.sub = {}

    @classmethod
    def _num_to_cn(cls, n):
        if n in cls.CN_MAP:
            return cls.CN_MAP[n]
        if n <= 0:
            return str(n)
        if n < 100:
            tens = n // 10
            ones = n % 10
            if ones == 0:
                return f'{cls.CN_MAP[tens]}十' if tens > 1 else '十'
            if tens == 0:
                return cls.CN_MAP.get(ones, str(ones))
            if tens == 1:
                return f'十{cls.CN_MAP.get(ones, ones)}'
            return f'{cls.CN_MAP[tens]}十{cls.CN_MAP.get(ones, ones)}'
        return str(n)

    def h1(self):
        self.ch += 1
        self.sec[self.ch] = 0
        return f'第{self._num_to_cn(self.ch)}章'

    def h2(self):
        self.sec[self.ch] = self.sec.get(self.ch, 0) + 1
        sec = self.sec[self.ch]
        self.sub[(self.ch, sec)] = 0
        return f'{self.ch}.{sec}'

    def h3(self):
        sec = self.sec.get(self.ch, 0)
        self.sub[(self.ch, sec)] = self.sub.get((self.ch, sec), 0) + 1
        return f'{self.ch}.{sec}.{self.sub[(self.ch, sec)]}'


def render_h1(doc, text, numbering, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h1_before),
                     after=Pt(config.spacing.h1_after))
    r = p.add_run(f'{numbering.h1()}  {remove_mark(text)}')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h1, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '1')
    p._element.get_or_add_pPr().append(o)


def render_h2(doc, text, numbering, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h2_before),
                     after=Pt(config.spacing.h2_after))
    r = p.add_run(f'{numbering.h2()}  {remove_mark(text)}')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h2, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '2')
    p._element.get_or_add_pPr().append(o)


def render_h3(doc, text, numbering, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h3_before),
                     after=Pt(config.spacing.h3_after))
    r = p.add_run(f'{numbering.h3()}  {remove_mark(text)}')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h3, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '3')
    p._element.get_or_add_pPr().append(o)


def render_h4(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.body,
                     before=Pt(config.spacing.h4_before),
                     after=Pt(config.spacing.h4_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h4, bold=True)


def render_h5(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.body,
                     before=Pt(config.spacing.h5_before),
                     after=Pt(config.spacing.h5_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h5, bold=True, italic=True)


def render_front_h2(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h2_before),
                     after=Pt(config.spacing.h2_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h2, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '2')
    p._element.get_or_add_pPr().append(o)


def render_front_h3(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h3_before),
                     after=Pt(config.spacing.h3_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h3, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '3')
    p._element.get_or_add_pPr().append(o)


def render_back_h1(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h1_before),
                     after=Pt(config.spacing.h1_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h1, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '1')
    p._element.get_or_add_pPr().append(o)


def render_back_h2(doc, text, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h2_before),
                     after=Pt(config.spacing.h2_after))
    r = p.add_run(remove_mark(text))
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin,
                 config.sizes.h2, bold=True)
    o = OxmlElement('w:outlineLvl')
    o.set(qn('w:val'), '2')
    p._element.get_or_add_pPr().append(o)
