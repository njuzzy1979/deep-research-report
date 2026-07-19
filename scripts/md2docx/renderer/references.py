"""参考文献格式化 — GB/T 7714 + 信源分级。"""
import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import set_run_font, set_para_spacing


SOURCE_LEVELS = {
    'L1': '权威来源', 'L2': '可靠来源', 'L3': '可参考来源',
    'L4': '交叉验证来源', 'L5': '待确认来源',
}

L_COLORS = {'L1': '000000', 'L2': '000000', 'L3': '555555',
            'L4': '888888', 'L5': 'AAAAAA'}


def _parse_hex(h):
    h = h.lstrip('#')
    return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def parse_reference_line(line):
    m = re.match(r'^\[(\d+)\]\s*(?:\[(L[1-5])\]\s*)?(.+)', line.strip())
    if not m:
        return None
    text = m.group(3).strip()
    ref_type = 'other'
    if '[J]' in text: ref_type = 'journal'
    elif '[M]' in text: ref_type = 'monograph'
    elif '[R]' in text: ref_type = 'report'
    elif '[EB/OL]' in text: ref_type = 'online'
    return {'num': m.group(1), 'level': m.group(2) or 'L3',
            'text': text, 'type': ref_type}


def render_references(doc, ref_lines, config):
    fonts = config.fonts
    sizes = config.sizes

    for line in ref_lines:
        parsed = parse_reference_line(line)
        if not parsed:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_para_spacing(p, ls=1.5)
            r = p.add_run(line)
            set_run_font(r, fonts.body.cjk, fonts.body.latin, sizes.body)
            continue

        lvl_color = _parse_hex(L_COLORS.get(parsed['level'], '000000'))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p, ls=1.5)

        r1 = p.add_run(f"[{parsed['num']}] ")
        set_run_font(r1, fonts.body.cjk, fonts.body.latin, sizes.body, bold=True)
        r2 = p.add_run(f"[{parsed['level']}] ")
        set_run_font(r2, fonts.body.cjk, fonts.body.latin, sizes.footnote, color=lvl_color)
        r3 = p.add_run(parsed['text'])
        set_run_font(r3, fonts.body.cjk, fonts.body.latin, sizes.body)

    # 信源分级图例
    p_leg = doc.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p_leg, ls=1.0, before=Pt(12), after=Pt(6))
    r = p_leg.add_run('信源分级说明: ')
    set_run_font(r, fonts.body.cjk, fonts.body.latin, sizes.footnote, bold=True)
    for level, desc in SOURCE_LEVELS.items():
        rl = p_leg.add_run(f' [{level}]={desc}')
        set_run_font(rl, fonts.body.cjk, fonts.body.latin, sizes.footnote,
                     color=_parse_hex(L_COLORS.get(level, '000000')))
