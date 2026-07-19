"""
封面渲染 — 参数化版本。

改进：机构名/标题/日期/版本从 config.metadata 读取，不再硬编码。
      顶部留白用 w:spacing w:before 精确控制。

原始来源: markdown_to_docx.py _render_cover (L541-580)
"""

from datetime import date
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .helpers import set_run_font, set_para_spacing


def render_cover(doc, config):
    """封面（V3.0 §7）。

    Args:
        doc: python-docx Document 实例
        config: Config 实例
    """
    meta = config.metadata
    fonts = config.fonts
    sizes = config.sizes
    colors = config.colors
    cover_cfg = config.cover

    # 顶部留白精确 6cm（V3.0 §7.1）
    p = doc.add_paragraph()
    p.alignment = 1  # CENTER
    p.paragraph_format.space_before = Cm(cover_cfg.top_spacing_cm)
    set_para_spacing(p, ls=1.0, before=Pt(0), after=Pt(0))

    # 报告标题：28pt Bold 黑色居中（V3.0: 一号）
    title = meta.title
    if title:
        p1 = doc.add_paragraph()
        p1.alignment = 1
        set_para_spacing(p1, ls=1.2, before=Pt(12), after=Pt(6))
        r1 = p1.add_run(title)
        set_run_font(r1, fonts.heading.cjk, fonts.heading.latin,
                     sizes.cover_title, bold=True)

    # 英文副标题：14pt 黑色居中（V3.0: 四号）
    subtitle = meta.subtitle
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = 1
        set_para_spacing(p2, ls=1.2, before=Pt(8), after=Pt(12))
        r2 = p2.add_run(subtitle)
        set_run_font(r2, fonts.heading.cjk, fonts.heading.latin,
                     sizes.cover_subtitle)

    # 报告类型标签：16pt 黑色居中（V3.0: 三号）
    subtype = meta.subtype or '深度研究报告'
    p3 = doc.add_paragraph()
    p3.alignment = 1
    set_para_spacing(p3, ls=1.2, before=Pt(12), after=Pt(24))
    r3 = p3.add_run(subtype)
    set_run_font(r3, fonts.heading.cjk, fonts.heading.latin,
                 sizes.cover_type)

    # 底部分隔线
    if cover_cfg.show_separator:
        sep = doc.add_paragraph()
        sep.alignment = 1
        set_para_spacing(sep, ls=1.0, before=Pt(96), after=Pt(12))
        sep_char = cover_cfg.separator_char
        sep_repeat = cover_cfg.separator_repeat
        rs = sep.add_run(sep_char * sep_repeat)
        set_run_font(rs, fonts.body.cjk, fonts.body.latin,
                     10, color=colors.as_tuple('separator'))

    # 机构名：14pt Bold 黑色
    org = meta.org
    if org:
        p4 = doc.add_paragraph()
        p4.alignment = 1
        set_para_spacing(p4, ls=1.5)
        r4 = p4.add_run(org)
        set_run_font(r4, fonts.heading.cjk, fonts.heading.latin,
                     14, bold=True)

    # 版本号 + 日期：11pt 黑色
    version = meta.version or 'V1.0'
    date_str = meta.date or date.today().strftime('%Y.%m.%d')
    p5 = doc.add_paragraph()
    p5.alignment = 1
    set_para_spacing(p5, ls=1.5)
    r5 = p5.add_run(f'{version} | {date_str}')
    set_run_font(r5, fonts.body.cjk, fonts.body.latin,
                 sizes.cover_info)

    doc.add_page_break()
