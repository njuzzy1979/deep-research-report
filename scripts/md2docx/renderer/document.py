"""
文档结构管理器 — 管理 python-docx Document 的 section 创建/切换、
页面设置、页码格式、页眉页脚。

对应设计文档 §4.5.1 和 §5.4。
实现约束 DR-06: Section 隔离——封面/目录的页眉页脚不影响正文。

用法:
    from .document import DocumentBuilder

    builder = DocumentBuilder(config)
    sections = builder.build_sections(ir, has_summary=True)
    # sections['cover'] / sections['summary'] / sections['body'] / sections['back']
    doc = builder.doc
"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..config import Config
from .helpers import set_run_font


class DocumentBuilder:
    """文档结构管理器。

    管理 python-docx Document 的 section 生命周期：
    - 创建 Document 实例并初始化默认样式
    - section 的创建/页面设置
    - 页码格式切换（罗马数字 / 阿拉伯数字）
    - 页眉页脚设置（含 Section 隔离）
    """

    def __init__(self, config: Config):
        """创建 Document 实例并设置 A4 页面与 Normal 样式。

        Args:
            config: Config 只读实例，提供所有字体/字号/页面参数。
        """
        self.config = config
        self.doc = Document()

        # ── 页面设置（V3.0 §2）：对默认 section 立即应用 ──
        self.setup_page(self.doc.sections[0])

        # ── Normal 样式（V3.0 §3.3 正文：宋体+TNR，body 字号，body 行距） ──
        self._setup_normal_style()

    # ═══════════════════════════════════════════════════════════
    # Section 生命周期
    # ═══════════════════════════════════════════════════════════

    def add_section_break(self, type_: str = "next_page"):
        """插入分节符并返回新 section。

        使用 python-docx 的 add_section()，默认创建"下一页"类型分节符。
        新 section 自动应用 A4 页面参数。

        Args:
            type_: 分节符类型，目前仅支持 python-docx 默认的 NEW_PAGE。
                   保留此参数以便未来扩展（如 continuous）。

        Returns:
            新创建的 Section 对象。
        """
        new_section = self.doc.add_section()
        self.setup_page(new_section)
        return new_section

    def setup_page(self, section):
        """为 section 设置 A4 页面参数与边距。

        所有尺寸从 config.page 读取，不硬编码。

        Args:
            section: python-docx Section 对象。
        """
        cp = self.config.page
        section.page_width = Cm(cp.width)
        section.page_height = Cm(cp.height)
        section.top_margin = Cm(cp.margin_top)
        section.bottom_margin = Cm(cp.margin_bottom)
        section.left_margin = Cm(cp.margin_left)
        section.right_margin = Cm(cp.margin_right)

    # ═══════════════════════════════════════════════════════════
    # 页码格式
    # ═══════════════════════════════════════════════════════════

    def set_page_number_format(self, section, fmt: str = "lowerRoman", start: int = 1):
        """设置 section 的页码格式与起始值。

        通过 OOXML w:pgNumType 元素直接操作 sectPr，实现罗马数字 /
        阿拉伯数字切换。

        Args:
            section: python-docx Section 对象。
            fmt: 页码格式标识，直接对应 OOXML w:fmt 属性值。
                 'lowerRoman' — 小写罗马数字 (i, ii, iii...)
                 'decimal'    — 阿拉伯数字 (1, 2, 3...)
            start: 起始页码（1-based），映射到 w:start 属性。
        """
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:fmt'), fmt)
        pgNumType.set(qn('w:start'), str(start))

        # 移除已有的 pgNumType（如果存在），避免重复元素
        for old in sectPr.findall(qn('w:pgNumType')):
            sectPr.remove(old)

        sectPr.append(pgNumType)

    # ═══════════════════════════════════════════════════════════
    # 页眉
    # ═══════════════════════════════════════════════════════════

    def setup_header(self, section, text: str, show_line: bool = True):
        """为 section 设置页眉：报告简称右对齐 + 1pt 黑色底线。

        ⚠ 关键约束（DR-06）：
        调用方必须在调用此方法**之前**将 section.header.is_linked_to_previous
        设为 False，否则修改会破坏前一 section 的页眉。

        字体从 config.fonts.heading.cjk 和 config.sizes.header_footer 读取。

        Args:
            section: python-docx Section 对象。
            text: 页眉文本（报告简称）。
            show_line: 是否添加 1pt 黑色段落下边框。
        """
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 清空可能残留的默认空 run
        hp.clear()

        r = hp.add_run(text)
        set_run_font(
            r,
            self.config.fonts.heading.cjk,
            self.config.fonts.heading.latin,
            self.config.sizes.header_footer,
        )

        # 1pt 黑色段落下边框（V3.0 §6.1）
        if show_line:
            self._add_header_bottom_border(hp)

    def _add_header_bottom_border(self, paragraph):
        """为页眉段落添加 1pt 黑色段落下边框（w:pBdr / w:bottom）。

        sz="8" = 8/8 pt = 1pt。

        Args:
            paragraph: 页眉段落对象。
        """
        pPr = paragraph._element.get_or_add_pPr()

        # 移除已有的段落下边框（避免重复）
        for existing in pPr.findall(qn('w:pBdr')):
            pPr.remove(existing)

        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')           # 1pt = 8 eighths of a point
        bottom.set(qn('w:color'), '000000')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ═══════════════════════════════════════════════════════════
    # 页脚
    # ═══════════════════════════════════════════════════════════

    def setup_footer_page_number(self, section):
        """为 section 设置页脚：页码居中（PAGE 域）。

        ⚠ 关键约束（DR-06）：
        调用方必须在调用此方法**之前**将 section.footer.is_linked_to_previous
        设为 False。

        字体从 config.sizes.header_footer 读取。

        Args:
            section: python-docx Section 对象。
        """
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 清空默认残留
        fp.clear()

        # ── 插入 PAGE 域：fldChar(begin) → instrText "PAGE" → fldChar(end) ──
        for tag, txt in [('begin', ''), ('instr', ' PAGE '), ('end', '')]:
            run = fp.add_run()
            if tag == 'instr':
                el = OxmlElement('w:instrText')
                el.set(qn('xml:space'), 'preserve')
                el.text = txt
            else:
                el = OxmlElement('w:fldChar')
                el.set(qn('w:fldCharType'), tag)
            run._element.append(el)

        # 设置页脚文字字体（V3.0 §6.2: 页脚 9pt）
        for run in fp.runs:
            set_run_font(
                run,
                self.config.fonts.body.cjk,
                self.config.fonts.body.latin,
                self.config.sizes.header_footer,
            )

    # ═══════════════════════════════════════════════════════════
    # 分节结构构建（设计文档 §5.4）
    # ═══════════════════════════════════════════════════════════

    def build_sections(self, ir, has_summary: bool = False):
        """构建完整的文档分节结构，返回各 section 引用。

        分节布局（设计文档 §3.4 / §5.4）：

            Section 1: 封面 + 目录 — 无页眉，无页码
            Section 2: 执行摘要   — 罗马数字页码 (i, ii, iii...)（若无摘要则跳过）
            Section 3: 正文       — 阿拉伯数字页码从 1 开始 + 页眉
            Section 4: 参考文献+附录 — 页码续正文，继承正文页眉

        ⚠ 关键约束 DR-06：每个 section 在修改内容之前必须先设置
        is_linked_to_previous = False。

        Args:
            ir: DocumentIR 实例，需提供 ir.metadata（dict，含 header_short / title）。
            has_summary: 是否有执行摘要内容（决定是否创建 Section 2）。

        Returns:
            dict: {'cover': sec1, 'summary': sec2_or_None, 'body': sec3, 'back': sec4}
        """
        # ── Section 1: 封面 + TOC（Document() 自动创建）──
        sec1 = self.doc.sections[0]
        self.setup_page(sec1)
        # 封面 section：清空页眉页脚（无页码）
        sec1.header.is_linked_to_previous = False
        sec1.footer.is_linked_to_previous = False
        for p in sec1.header.paragraphs:
            p.clear()
        for p in sec1.footer.paragraphs:
            p.clear()

        # ── Section 2: 执行摘要（罗马数字页码）──
        if has_summary:
            sec2 = self.doc.add_section()
            self.setup_page(sec2)
            sec2.header.is_linked_to_previous = False
            sec2.footer.is_linked_to_previous = False
            for p in sec2.header.paragraphs:
                p.clear()
            self.set_page_number_format(sec2, 'lowerRoman', 1)
            self.setup_footer_page_number(sec2)
        else:
            sec2 = None

        # ── 正文前插入分节符 ──
        self.doc.add_section()

        # ── Section 3: 正文（阿拉伯数字从 1 开始）──
        sec3 = self.doc.sections[-1]
        self.setup_page(sec3)
        sec3.header.is_linked_to_previous = False
        sec3.footer.is_linked_to_previous = False

        # 页眉文本优先级: header_short > title > 空字符串
        header_text = (
            ir.metadata.get('header_short')
            or ir.metadata.get('title', '')
        )
        if header_text:
            self.setup_header(sec3, header_text)

        self.set_page_number_format(sec3, 'decimal', 1)
        self.setup_footer_page_number(sec3)

        # ── Section 4: 参考文献 + 附录（续正文页码）──
        self.doc.add_section()
        sec4 = self.doc.sections[-1]
        self.setup_page(sec4)
        # 继承正文页眉页脚（linked_to_previous = True）
        sec4.header.is_linked_to_previous = True
        sec4.footer.is_linked_to_previous = True

        return {
            'cover': sec1,
            'summary': sec2,
            'body': sec3,
            'back': sec4,
        }

    # ═══════════════════════════════════════════════════════════
    # 内部辅助
    # ═══════════════════════════════════════════════════════════

    def _setup_normal_style(self):
        """设置 doc.styles['Normal'] 的字体、字号、行距。

        中文字体从 config.fonts.body.cjk 读取，
        西文字体从 config.fonts.body.latin 读取，
        字号从 config.sizes.body 读取，
        行距从 config.line_spacing.body 读取。
        """
        sty = self.doc.styles['Normal']
        fb = self.config.fonts.body
        sty.font.name = fb.latin
        sty.font.size = Pt(self.config.sizes.body)
        sty.paragraph_format.line_spacing = self.config.line_spacing.body

        # 通过 rPr/rFonts 设置 CJK 字体（python-docx 高层 API 无法直接设 CJK）
        se = sty.element
        rPr = se.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), fb.cjk)
        rFonts.set(qn('w:ascii'), fb.latin)
        rFonts.set(qn('w:hAnsi'), fb.latin)
