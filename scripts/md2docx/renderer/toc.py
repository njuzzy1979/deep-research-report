"""TOC（目录）和图表目录渲染。"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .helpers import set_run_font, set_para_spacing


def add_toc(doc, config, levels=None):
    fonts = config.fonts
    if levels is None:
        levels = config.toc.levels

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h1_before),
                     after=Pt(config.spacing.h1_after))
    r = p.add_run('目  录')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin, config.sizes.h1, bold=True)

    pt = doc.add_paragraph()
    for action in ['begin', 'instr', 'separate', 'result', 'end']:
        run = pt.add_run()
        if action == 'instr':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = f'TOC \\o "1-{levels}" \\h \\z'
            run._element.append(el)
        elif action == 'result':
            run.text = ''
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), action)
            run._element.append(el)


def add_figure_toc(doc, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h2_before),
                     after=Pt(config.spacing.h2_after))
    r = p.add_run('图目录')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin, config.sizes.h2, bold=True)
    pt = doc.add_paragraph()
    for action in ['begin', 'instr', 'separate', 'result', 'end']:
        run = pt.add_run()
        if action == 'instr':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = 'TOC \\c "Figure"'
            run._element.append(el)
        elif action == 'result':
            run.text = ''
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), action)
            run._element.append(el)


def add_table_toc(doc, config):
    fonts = config.fonts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=config.line_spacing.heading,
                     before=Pt(config.spacing.h2_before),
                     after=Pt(config.spacing.h2_after))
    r = p.add_run('表目录')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin, config.sizes.h2, bold=True)
    pt = doc.add_paragraph()
    for action in ['begin', 'instr', 'separate', 'result', 'end']:
        run = pt.add_run()
        if action == 'instr':
            el = OxmlElement('w:instrText')
            el.set(qn('xml:space'), 'preserve')
            el.text = 'TOC \\c "Table"'
            run._element.append(el)
        elif action == 'result':
            run.text = ''
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), action)
            run._element.append(el)
