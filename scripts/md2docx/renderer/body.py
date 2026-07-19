"""正文/引用/列表渲染。"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import set_run_font, set_para_spacing, add_inline, remove_mark


def render_body(doc, text, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.body,
                     before=Pt(config.spacing.body_before),
                     after=Pt(config.spacing.body_after),
                     first=Cm(config.indent.body_first_line))
    add_inline(p, text, config.fonts.body.cjk, config.fonts.body.latin, config.sizes.body)


def render_blockquote(doc, lines, config):
    fonts = config.fonts
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_para_spacing(p, ls=config.line_spacing.quote,
                         first=Cm(config.indent.body_first_line),
                         left=Cm(config.indent.quote_left))
        p.paragraph_format.right_indent = Cm(config.indent.quote_left)
        r = p.add_run(remove_mark(line))
        set_run_font(r, fonts.special.cjk, fonts.special.latin,
                     config.sizes.quote, italic=True)


def render_code_block(doc, lines, config):
    fonts = config.fonts
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_spacing(p, ls=1.0, left=Cm(1.0))
        r = p.add_run(line or ' ')
        set_run_font(r, fonts.mono.cjk, fonts.mono.latin, 9)


def render_list_item(doc, elem, config):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, ls=config.line_spacing.body, left=Cm(config.indent.list_left))
    p.paragraph_format.first_line_indent = -Cm(config.indent.list_hang)
    prefix = f'{elem.order_num}. ' if elem.ordered else '● '
    add_inline(p, prefix + elem.text, config.fonts.body.cjk,
               config.fonts.body.latin, config.sizes.body)
