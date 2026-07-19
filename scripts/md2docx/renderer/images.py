"""图片嵌入与图注渲染 — 数据驱动。"""
import re
import os
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .helpers import (set_run_font, set_para_spacing, make_xml_para,
                      get_para_text, has_image)


def embed_image(doc, fig_meta, config, images_base=None):
    fonts = config.fonts
    sizes = config.sizes
    colors = config.colors

    file_path = fig_meta.file_path
    if images_base and not os.path.isabs(file_path):
        file_path = os.path.join(images_base, file_path)
    if not os.path.exists(file_path):
        return False

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(img_p, ls=1.0, before=Pt(2), after=Pt(2))
    run = img_p.add_run()
    try:
        run.add_picture(file_path, width=Inches(5.2))
    except Exception:
        return False

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap_p, ls=config.line_spacing.caption,
                     before=Pt(1), after=Pt(10))
    r = cap_p.add_run(f'图{fig_meta.id}  {fig_meta.caption}')
    set_run_font(r, fonts.body.cjk, fonts.body.latin,
                 sizes.caption, color=colors.as_tuple('secondary'))
    return True


def embed_all_images(doc, figure_registry, config, images_base=None):
    fonts = config.fonts
    sizes = config.sizes
    colors = config.colors

    available = set()
    if images_base and os.path.isdir(images_base):
        available = set(os.listdir(images_base))

    refs = []
    for i, para in enumerate(doc.paragraphs):
        m = re.search(r'图\s*(\d+)[-−]\s*(\d+)', para.text)
        if m:
            key = m.group(1) + '-' + m.group(2)
            if key in figure_registry:
                fig = figure_registry[key]
                fname = os.path.basename(fig.file_path)
                if images_base and fname in available:
                    refs.append((i, key, fig))

    seen_idx = set()
    uniq = []
    for idx, key, fig in refs:
        if idx not in seen_idx:
            seen_idx.add(idx)
            uniq.append((idx, key, fig))

    count = 0
    for para_idx, fig_key, fig_meta in sorted(uniq, key=lambda x: x[0], reverse=True):
        para = doc.paragraphs[para_idx]
        parent = para._element.getparent()
        pos = list(parent).index(para._element)

        file_path = fig_meta.file_path
        if images_base and not os.path.isabs(file_path):
            file_path = os.path.join(images_base, file_path)
        if not os.path.exists(file_path):
            continue

        img_p = OxmlElement('w:p')
        img_pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        img_pPr.append(jc)
        s = OxmlElement('w:spacing')
        s.set(qn('w:before'), '200')
        s.set(qn('w:after'), '100')
        img_pPr.append(s)
        img_p.append(img_pPr)

        tmp = doc.add_paragraph()
        tmp_run = tmp.add_run()
        try:
            tmp_run.add_picture(file_path, width=Inches(5.2))
        except Exception:
            doc.element.body.remove(tmp._element)
            continue
        drawing = tmp._element.find(qn('w:r')).find(qn('w:drawing'))
        img_r = OxmlElement('w:r')
        img_r.append(drawing)
        img_p.append(img_r)
        doc.element.body.remove(tmp._element)

        cap_p = make_xml_para(f'图{fig_meta.id}  {fig_meta.caption}',
                              cjk=fonts.body.cjk, latin=fonts.body.latin,
                              size_pt=sizes.caption, color_hex=colors.secondary)

        parent.insert(pos, cap_p)
        parent.insert(pos, img_p)
        parent.remove(para._element)
        count += 1

    _cleanup_stray_fig_refs(doc)
    return count


def _cleanup_stray_fig_refs(doc):
    body_c = list(doc.element.body)
    to_remove = []
    for i, c in enumerate(body_c):
        if c.tag != qn('w:p'):
            continue
        t = get_para_text(c)
        if not re.match(r'(●\s*)?图\s*\d+[-−]\s*\d+[：:]', t):
            continue
        prev_img = i > 0 and body_c[i - 1].tag == qn('w:p') and has_image(body_c[i - 1])
        next_img = i + 1 < len(body_c) and body_c[i + 1].tag == qn('w:p') and has_image(body_c[i + 1])
        if prev_img or next_img:
            continue
        if 'KaiTi' in c.xml:
            continue
        to_remove.append(i)
    for i in reversed(to_remove):
        body_c[i].getparent().remove(body_c[i])
