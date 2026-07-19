"""特殊元素渲染：定义框、案例框（V3.0 §9）。"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import set_run_font, set_para_spacing, set_cell_border, set_cell_shading


def add_definition_box(doc, box_id, title, content_lines, config):
    fonts = config.fonts
    sizes = config.sizes
    bg = config.colors.bg_light

    p = doc.add_paragraph()
    set_para_spacing(p, ls=1.0, before=Pt(6), after=Pt(3))

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, bg)

    left_bdr = {'sz': '24', 'val': 'single', 'color': '000000'}
    thin = {'sz': '4', 'val': 'single', 'color': '000000'}
    set_cell_border(cell, top=thin, bottom=thin, left=left_bdr, right=thin)

    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = tp.add_run(f'{box_id}  {title}')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin, sizes.h4, bold=True)

    for line in content_lines:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r2 = p2.add_run(line)
        set_run_font(r2, fonts.body.cjk, fonts.body.latin, sizes.body)

    p3 = doc.add_paragraph()
    set_para_spacing(p3, ls=1.0, before=Pt(3), after=Pt(6))


def add_case_box(doc, title, content_lines, config):
    fonts = config.fonts
    sizes = config.sizes

    p = doc.add_paragraph()
    set_para_spacing(p, ls=1.0, before=Pt(6), after=Pt(3))

    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]

    outer = {'sz': '12', 'val': 'single', 'color': '000000'}
    top_bdr = {'sz': '24', 'val': 'single', 'color': 'BBBBBB'}
    set_cell_border(cell, top=top_bdr, bottom=outer, left=outer, right=outer)

    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = tp.add_run(f'案例: {title}')
    set_run_font(r, fonts.heading.cjk, fonts.heading.latin, sizes.h4, bold=True)

    for line in content_lines:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r2 = p2.add_run(line)
        set_run_font(r2, fonts.body.cjk, fonts.body.latin, sizes.body)

    p3 = doc.add_paragraph()
    set_para_spacing(p3, ls=1.0, before=Pt(3), after=Pt(6))
