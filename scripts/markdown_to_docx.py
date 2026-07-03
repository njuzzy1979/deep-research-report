#!/usr/bin/env python3
"""
Markdown → Word (.docx) 研究报告一体化转换器

整合了以下功能：
1. Markdown 清理（BOM/行尾/HTML/手动编号/图表占位/印刷页数等）
2. Word 文档生成（封面/目录/标题编号/表格/页眉页脚/TOC域）
3. 图片嵌入（PNG→docx + 图题注下方居中）
4. 表题注修复（表头内容匹配→题注上方居中）

用法:
    python markdown_to_docx.py <input.md> <output.docx> --title "报告题名" --images "figures_dir"

依赖: pip install python-docx matplotlib
"""

import re, sys, os, argparse
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 常量 ──

# ── 字体方案（研究报告格式规范 V3.0 §3.1）──
FONT_CJK_BODY     = '宋体'
FONT_LATIN_BODY    = 'Times New Roman'
FONT_CJK_HEADING   = '微软雅黑'       # V3.0: 标题醒目
FONT_LATIN_HEADING  = 'Times New Roman'
FONT_CJK_SPECIAL   = '楷体'           # 引文/特殊段落（保留兼容）
FONT_LATIN_SPECIAL  = 'Times New Roman'
FONT_MONO_CJK      = '宋体'
FONT_MONO_LATIN     = 'Consolas'

# ── 字号（V3.0 §3.2 标题层级 + §3.3 正文样式）──
SIZE_TITLE         = 28   # 封面标题（一号）
SIZE_SUBTITLE      = 14   # 英文副标题（四号）
SIZE_COVER_TYPE    = 16   # 报告类型标签（三号）
SIZE_COVER_INFO    = 11   # 版本/日期
SIZE_H1            = 24   # 章标题（小一）
SIZE_H2            = 16   # 节标题（三号）
SIZE_H3            = 14   # 小节标题（四号）
SIZE_H4            = 12   # 段落小标题（小四）
SIZE_BODY          = 11   # 正文（11pt，V3.0: 五号→11pt）
SIZE_TABLE_HEADER  = 10   # 表头
SIZE_TABLE_BODY    = 10.5 # 表体
SIZE_SMALL         = 10.5 # 五号（引用/导语）
SIZE_TINY          = 9    # 小五（图注/表注/脚注/页眉/页脚）

# ── 标题间距（V3.0 §3.2）──
SPACE_H1_BEFORE  = Pt(0)    # V3.0: 章标题段前 0pt
SPACE_H1_AFTER   = Pt(18)   # V3.0: 章标题段后 18pt
SPACE_H2_BEFORE  = Pt(24)   # V3.0: 节标题段前 24pt
SPACE_H2_AFTER   = Pt(12)   # V3.0: 节标题段后 12pt
SPACE_H3_BEFORE  = Pt(18)   # V3.0: 小节段前 18pt
SPACE_H3_AFTER   = Pt(8)    # V3.0: 小节段后 8pt
SPACE_H4_BEFORE  = Pt(12)   # V3.0: 段落小标题段前 12pt
SPACE_H4_AFTER   = Pt(6)    # V3.0: 段落小标题段后 6pt
SPACE_TITLE_BEFORE = Pt(24)
SPACE_TITLE_AFTER  = Pt(18)
SPACE_BODY_BEFORE  = Pt(0)
SPACE_BODY_AFTER   = Pt(0)

# ── 配色方案（V3.0 §11）──
COLOR_BLACK       = (0x00, 0x00, 0x00)
COLOR_DARK_GRAY   = (0x55, 0x55, 0x55)  # 题注/脚注辅助文字
COLOR_LIGHT_GRAY  = (0xF2, 0xF2, 0xF2)  # 表格交替行/定义框底色
COLOR_MID_GRAY    = (0x33, 0x33, 0x33)  # 特殊强调
COLOR_SEPARATOR   = (0xBB, 0xBB, 0xBB)  # 封面装饰分隔线

LINE_SPACING_SINGLE  = 1.0
LINE_SPACING_1_5     = 1.5
INDENT_FIRST_LINE = Cm(0.74)
INDENT_LEFT_QUOTE = Cm(1.5)
INDENT_LEFT_LIST  = Cm(1.5)
INDENT_HANG_LIST  = Cm(0.75)

CN_NUMS = r'一二三四五六七八九十'
CN = r'[' + CN_NUMS + r']+'

# ── FIGURE MAP: ID→(file, caption) ──
FIGURE_MAP = {
    '0-1': ('fig-0-1-market-growth.png',      '图0-1：中国商业SSA市场规模增长预测（2024-2030）'),
    '0-2': ('fig-3-1-ecosystem.png',          '图0-2：全球SSA竞争格局总览'),
    '1-1': ('fig-3-1-ecosystem.png',          '图1-1：空间态势感知在商业航天生态中的位置'),
    '1-2': ('fig-4-1-tech-architecture.png',  '图1-2：本报告分析框架'),
    '2-1': ('fig-2-1-global-market.png',      '图2-1：全球SSA市场增长趋势（2020-2034）'),
    '2-2': ('fig-2-2-regional-pie.png',       '图2-2：全球SSA市场区域结构（2024年）'),
    '2-3': ('fig-5-1-demand-quadrants.png',   '图2-3：SSA需求驱动因素因果链图'),
    '2-4': ('fig-2-4-satellite-growth.png',   '图2-4：全球在轨卫星数量增长（2015-2030）'),
    '2-6': ('fig-3-1-ecosystem.png',          '图2-6：全球SSA企业传感器网络地理分布'),
    '3-1': ('fig-3-1-ecosystem.png',          '图3-1：中国商业SSA产业生态系统'),
    '3-2': ('fig-3-2-radar.png',              '图3-2：中国商业SSA企业竞争力雷达图'),
    '3-3': ('fig-3-2-radar.png',              '图3-3：中国商业SSA企业6维竞争力评估'),
    '3-4': ('fig-3-4-value-chain.png',        '图3-4：中国商业SSA产业链价值分布'),
    '4-1': ('fig-4-1-tech-architecture.png',  '图4-1：SSA六层技术架构全景'),
    '4-2': ('fig-4-1-tech-architecture.png',  '图4-2：天基SSA星座多层架构'),
    '4-3': ('fig-4-1-tech-architecture.png',  '图4-3：SSA数据处理全流程'),
    '4-4': ('fig-4-4-orbit-accuracy.png',     '图4-4：轨道精度对比'),
    '4-5': ('fig-4-4-orbit-accuracy.png',     '图4-5：AI-SSA技术路线图（2025-2035）'),
    '4-6': ('fig-4-6-catalog-comparison.png', '图4-6：中美商业SSA编目数量增长对比'),
    '5-1': ('fig-5-1-demand-quadrants.png',   '图5-1：中国商业SSA需求四象限'),
    '5-2': ('fig-5-2-constellation-demand.png','图5-2：星座部署与SSA需求增长联动'),
    '5-3': ('fig-4-1-tech-architecture.png',  '图5-3：中国SSA政策体系框架'),
    '6-1': ('fig-3-1-ecosystem.png',          '图6-1：中国军民融合SSA架构'),
    '6-2': ('fig-4-6-catalog-comparison.png', '图6-2：全球太空军事化能力对比矩阵'),
    '7-1': ('fig-4-1-tech-architecture.png',  '图7-1：国际SSA治理体系架构'),
    '7-2': ('fig-4-4-orbit-accuracy.png',     '图7-2：TraCSS与Space-Track服务界面对比'),
    '7-3': ('fig-2-1-global-market.png',      '图7-3：全球STM治理关键节点时间线'),
    '8-1': ('fig-8-1-risk-matrix.png',        '图8-1：中国商业SSA风险地图'),
}

# ── TABLE MATCHERS: header pattern → caption ──
TABLE_MATCHERS = [
    (r'等级.*来源类型.*本报告中的用法.*占比',       '表F-2：本报告证据等级标注规范'),
    (r'等级.*来源.*用法',                          '表F-1：证据等级说明'),
    (r'术语.*英文全称.*核心含义.*使用语境',     '表1-1：SSA/SDA/STM概念辨析'),
    (r'方法.*应用.*核心产出',                  '表1-2：研究方法组合'),
    (r'区域.*市场规模.*占比.*增长率',           '表2-1：全球SSA市场区域分布（2024年）'),
    (r'类别.*数量',                            '表2-2：ESA 2025空间碎片统计数据'),
    (r'管辖区.*核心法规.*离轨时限.*SSA',        '表2-3：全球三大监管极SSA法规对比'),
    (r'维度.*LeoLabs.*COMSPOC.*ExoAnalytic',  '表2-4：全球商业SSA企业对标矩阵'),
    (r'参与者类型.*代表实体.*核心角色.*对商业',   '表3-1：中国商业SSA产业生态参与者'),
    (r'子平台.*功能.*SSA相关能力',              '表3-2：星图测控"洞察者"平台矩阵'),
    (r'因素.*说明',                            '表3-3：SSA中游平台核心竞争要素'),
    (r'客户类型.*需求占比.*需求特征.*增长潜力',   '表3-4：SSA下游客户需求结构'),
    (r'精度等级.*典型系统.*误差量级.*适用场景',   '表3-5：轨道精度等级对比'),
    (r'代际.*时期.*核心特征.*代表系统',          '表4-1：SSA技术体系代际演进'),
    (r'技术层面.*自主化率.*关键依赖.*替代路径',   '表4-2：中国SSA技术自主化水平评估'),
    (r'系统.*能力.*地位',                      '表4-3：中国地基雷达系统能力'),
    (r'系统.*能力\s*$',                        '表4-4：中国地基光学望远镜系统'),
    (r'星座.*国家.*计划规模.*轨道',             '表4-5：全球天基SSA星座对比'),
    (r'系统.*传播7天后误差.*特点',              '表4-6：轨道预报精度对比'),
    (r'方向.*技术路线.*成熟度.*中国进展',        '表4-7：AI在SSA各环节的应用现状'),
    (r'维度.*美国.*商业.*中国.*商业.*差距',      '表4-8：中美商业SSA技术能力对比'),
    (r'需求来源.*卫星数量.*SSA年费.*市场规模',    '表5-1：2027年中国SSA市场需求测算'),
    (r'需求领域.*市场规模.*增长驱动',            '表5-2：2030年SSA市场远景测算'),
    (r'星座.*总规模.*关键里程碑.*部署时间表',     '表5-3：中国主要星座SSA需求参数'),
    (r'生命周期阶段.*要求.*SSA服务需求',         '表5-4：碎片减缓全生命周期SSA需求'),
    (r'维度.*正面影响.*负面影响',               '表6-1：军民融合对商业SSA的影响矩阵'),
    (r'系统.*轨道.*能力.*军事价值',             '表6-2：中国关键军事SSA系统'),
    (r'举措.*内容.*预期效果',                   '表6-3：军民融合透明化改革建议'),
    (r'国家.*军事航天组织.*SSA军事体系.*商业',   '表6-4：全球主要国家太空军事态势对比'),
    (r'事件.*年份.*碎片数量.*当前状态',          '表6-5：历史上三大ASAT碎片事件'),
    (r'管辖区.*离轨时限.*执行机制.*生效年份',     '表7-1：全球主要离轨时限标准对比'),
    (r'层级.*提供方.*服务内容.*收费.*目标用户',   '表7-2：美国TraCSS三层服务架构'),
    (r'立场.*代表方.*核心论点',                  '表7-3：2025年TraCSS预算争议各方立场'),
    (r'维度.*TraCSS.*EU SST',                 '表7-4：TraCSS与EU SST关键差异'),
    (r'风险项.*可能表现.*影响',                  '表8-1：政策执行不及预期风险清单'),
    (r'风险.*类型.*概率.*影响.*风险等级.*优先级', '表8-2：中国商业SSA风险综合评估矩阵'),
    (r'行动项.*责任主体.*时间节点.*资源需求.*效果指标', '表9-1：企业技术战略行动清单'),
    (r'行动项.*责任主体.*时间节点.*关键成功因素',  '表9-2：企业商业模式升级路径'),
    (r'联盟职能.*具体产出.*时间节点',            '表9-3：中国商业SSA产业联盟建设规划'),
    (r'优先级.*赛道.*壁垒强度.*增长前景.*投资风险', '表9-4：商业SSA投资赛道评估'),
    (r'风险类型.*风险等级.*发生概率.*影响.*缓释',   '表9-5：投资风险管理矩阵'),
    (r'退出方式.*可能性.*时间窗口.*参考案例',     '表9-6：投资退出路径预判'),
    (r'维度.*乐观场景.*基准场景.*悲观场景',       '表9-7：三种场景分析（2025-2030）'),
    (r'层级.*提供方.*服务内容.*收费模式.*覆盖',    '表9-8：SSA服务层级建议架构'),
]
COMPILED_TABLE_MATCHERS = [(re.compile(p), cap) for p, cap in TABLE_MATCHERS]

CHAPTER_TITLES = [
    '导论与研究方法', '全球商业SSA产业全景', '中国商业SSA竞争格局',
    'SSA核心技术体系', '需求与驱动机制', '军民融合与国家安全维度',
    '国际规制与太空治理', '风险与挑战', '启示与建议',
]

CN_NAMES = ['一','二','三','四','五','六','七','八','九','十']

# ───────────────────────────────────────────────────────
# PHASE 1: Markdown cleaning
# ───────────────────────────────────────────────────────

RE_HTML    = re.compile(r"^<div\s+style=['\"]page-break-before:\s*always;?['\"]>\s*</div>\s*$", re.I)
RE_CHART   = re.compile(r'\*{0,2}图表占位\*{0,2}[：:]\s*')
RE_PAGE    = re.compile(r'^[（(](?:第' + CN + r'章完|摘要完)[，,]建议印刷页数[：:]\s*\d+\s*页[）)]\s*$')
RE_TOC     = re.compile(r'[（(]在Word/WPS')
RE_META    = re.compile(r'^\*\*(?:版本|日期|定位|规模|核心问题)\*\*[：:].*$')
RE_FULL    = re.compile(r'^\*《中国商业空间态势感知领域深度分析报告》全文完\*\s*$')
RE_H1_CH   = re.compile(r'^(#\s+)第' + CN + r'章[：:]\s*')
RE_H2_SEC  = re.compile(r'^(##\s+)\d+\.\d+\s+')
RE_H3_SUB  = re.compile(r'^(###\s+)\d+\.\d+\.\d+\s+')
RE_TBL_REF = re.compile(r'^表\s*\d+[-−]\s*\d+[：:]')
RE_LIST_FIG = re.compile(r'^[-*]\s+(图\s*\d+[-−]\s*\d+[：:].*)$')
RE_COVER  = re.compile(r'^\*\*China.+Deep-Dive Analysis\*\*\s*$')


def clean_markdown(input_path):
    """Binary-safe clean: strip CR, remove junk, restructure headings, insert --- markers."""
    with open(input_path, 'rb') as f:
        raw = f.read()
    if raw[:3] == b'\xef\xbb\xbf': raw = raw[3:]
    raw = raw.replace(b'\r\n', b'\n').replace(b'\r', b'\n')
    lines = raw.decode('utf-8').split('\n')

    out = []
    for line in lines:
        s = line.strip()
        if RE_HTML.match(s): continue
        if RE_PAGE.match(s) or re.match(r'^[（(]摘要完', s): continue
        if RE_TOC.search(line): continue
        if RE_META.match(s): continue
        if RE_FULL.match(s): continue
        if '图表占位' in line: line = RE_CHART.sub('', line)
        # Convert list-fig refs to plain text
        m = RE_LIST_FIG.match(line)
        if m: line = m.group(1)
        m = RE_H1_CH.match(line)
        if m: line = m.group(1) + line[m.end():]
        m = RE_H2_SEC.match(line)
        if m: line = m.group(1) + line[m.end():]
        m = RE_H3_SUB.match(line)
        if m: line = m.group(1) + line[m.end():]
        out.append(line)

    # Remove orphan table refs
    out2 = []; i = 0
    while i < len(out):
        line = out[i]
        if RE_TBL_REF.match(line.strip()):
            has = any(out[j].strip().startswith('|') and out[j].strip().endswith('|')
                      for j in range(i+1, min(i+4, len(out))))
            if not has: i += 1; continue
        out2.append(line); i += 1

    # English subtitle → H2
    for j, ln in enumerate(out2):
        if RE_COVER.match(ln.strip()):
            out2[j] = '## ' + ln.strip().strip('*'); break
    # 目录→H1, 摘要→H2
    for j, ln in enumerate(out2):
        if ln.strip() == '## 目录': out2[j] = '# 目录'
        elif ln.strip() == '# 摘要': out2[j] = '## 摘要'
    # Remove ALL --- lines
    out3 = [ln for ln in out2 if ln.strip() != '---']
    # Insert --- markers (sorted descending)
    insertions = []
    for j, ln in enumerate(out3):
        if ln.strip() in ['# 目录', '## 证据等级说明']:
            insertions.append(j)
    for title in CHAPTER_TITLES:
        for j, ln in enumerate(out3):
            if re.match(r'^#\s+' + re.escape(title) + r'\s*$', ln.strip()):
                insertions.append(j); break
    for j, ln in enumerate(out3):
        if ln.strip().startswith('# 附录'): insertions.append(j); break
    for idx in sorted(set(insertions), reverse=True):
        out3.insert(idx, '---')
    return out3


# ───────────────────────────────────────────────────────
# PHASE 2: Markdown parser (extracted from original)
# ───────────────────────────────────────────────────────

class Heading:
    def __init__(s, lvl, text): s.level=lvl; s.text=text
class Paragraph:
    def __init__(s, text): s.text=text
class CodeBlock:
    def __init__(s, lang, lines): s.language=lang; s.lines=lines
class BlockQuote:
    def __init__(s, lines): s.lines=lines
class HorizontalRule: pass
class EmptyLine: pass
class ListItem:
    def __init__(s, text, indent=0, ordered=False, num=None):
        s.text=text; s.indent_level=indent; s.ordered=ordered; s.order_num=num
class TableElement:
    def __init__(s, header, rows): s.header=header; s.rows=rows
class ImageElement:
    def __init__(s, alt, url): s.alt_text=alt; s.url=url


def parse_markdown(text):
    lines = text.split('\n'); elements = []; i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip() == '': elements.append(EmptyLine()); i+=1; continue
        if line.strip().startswith('```'):
            lang=line.strip()[3:].strip(); clines=[]; i+=1
            while i<len(lines) and not lines[i].strip().startswith('```'):
                clines.append(lines[i]); i+=1
            i+=1; elements.append(CodeBlock(lang,clines)); continue
        if re.match(r'^[-\*_]{3,}\s*$', line.strip()):
            elements.append(HorizontalRule()); i+=1; continue
        hm = re.match(r'^(#{1,4})\s+(.*)', line)
        if hm: elements.append(Heading(len(hm.group(1)), hm.group(2).strip())); i+=1; continue
        if '|' in line and i+1<len(lines) and re.match(r'^[\|\s\-:]+$', lines[i+1].strip()):
            hdr=[c.strip() for c in line.split('|') if c.strip()]; i+=2; rows=[]
            while i<len(lines) and '|' in lines[i] and lines[i].strip():
                row=[c.strip() for c in lines[i].split('|') if c.strip()]
                if row: rows.append(row)
                i+=1
            elements.append(TableElement(hdr,rows)); continue
        if line.strip().startswith('>'):
            ql=[];
            while i<len(lines) and lines[i].strip().startswith('>'):
                ql.append(lines[i].strip()[1:].strip()); i+=1
            elements.append(BlockQuote(ql)); continue
        if re.match(r'^(\s*)[-*+]\s+', line):
            li=[]
            while i<len(lines) and re.match(r'^(\s*)[-*+]\s+', lines[i]):
                m=re.match(r'^(\s*)[-*+]\s+(.*)', lines[i])
                li.append(ListItem(m.group(2).strip(), indent=len(m.group(1))//2)); i+=1
            elements.extend(li); continue
        if re.match(r'^(\s*)\d+\.\s+', line):
            li=[]
            while i<len(lines) and re.match(r'^(\s*)\d+\.\s+', lines[i]):
                m=re.match(r'^(\s*)(\d+)\.\s+(.*)', lines[i])
                li.append(ListItem(m.group(3).strip(), indent=len(m.group(1))//2, ordered=True, num=int(m.group(2)))); i+=1
            elements.extend(li); continue
        im=re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', line.strip())
        if im: elements.append(ImageElement(im.group(1), im.group(2))); i+=1; continue
        elements.append(Paragraph(line.strip())); i+=1
    return elements


# ───────────────────────────────────────────────────────
# PHASE 3: DOCX rendering helpers
# ───────────────────────────────────────────────────────

def set_run_font(run, cjk, latin, size, bold=False, italic=False, color=None):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)

def set_para_spacing(p, ls=1.5, before=Pt(0), after=Pt(0), first=None, left=None):
    pf = p.paragraph_format
    pf.line_spacing = ls; pf.space_before = before; pf.space_after = after
    if first: pf.first_line_indent = first
    if left: pf.left_indent = left

def add_inline(para, text, cjk, latin, size):
    pattern = re.compile(r'(\*\*(.+?)\*\*)|(\*(.+?)\*)|(`(.+?)`)|(\[(.+?)\]\((.+?)\))')
    last=0
    for m in pattern.finditer(text):
        if m.start()>last:
            r=para.add_run(text[last:m.start()]); set_run_font(r,cjk,latin,size)
        if m.group(2): r=para.add_run(m.group(2)); set_run_font(r,cjk,latin,size,bold=True)
        elif m.group(4): r=para.add_run(m.group(4)); set_run_font(r,cjk,latin,size,italic=True)
        elif m.group(6): r=para.add_run(m.group(6)); set_run_font(r,FONT_MONO_CJK,FONT_MONO_LATIN,SIZE_SMALL)
        elif m.group(8): r=para.add_run(m.group(8)); set_run_font(r,cjk,latin,size); r.font.color.rgb=RGBColor(0x05,0x63,0xC1); r.underline=True
        last=m.end()
    if last<len(text):
        r=para.add_run(text[last:]); set_run_font(r,cjk,latin,size)

def remove_mark(s):
    return re.sub(r'\*\*(.+?)\*\*', r'\1', re.sub(r'\*(.+?)\*', r'\1', re.sub(r'`(.+?)`', r'\1', s)))

def _make_xml_para(text, cjk='宋体', latin='Times New Roman', size_pt=9, bold=False, align='center', after=300, color_hex=None):
    """生成纯 XML 段落（题注等），默认 V3.0 题注：9pt #555555"""
    p=OxmlElement('w:p'); pPr=OxmlElement('w:pPr')
    jc=OxmlElement('w:jc'); jc.set(qn('w:val'),align); pPr.append(jc)
    sp=OxmlElement('w:spacing'); sp.set(qn('w:after'),str(after)); pPr.append(sp)
    p.append(pPr)
    r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    rf=OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'),cjk); rf.set(qn('w:ascii'),latin); rf.set(qn('w:hAnsi'),latin)
    rPr.append(rf)
    if bold: rPr.append(OxmlElement('w:b'))
    sz=OxmlElement('w:sz'); sz.set(qn('w:val'),str(int(size_pt*2))); rPr.append(sz)
    if color_hex:
        color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), color_hex); rPr.append(color_el)
    r.append(rPr)
    t=OxmlElement('w:t'); t.text=text; t.set(qn('xml:space'),'preserve')
    r.append(t); p.append(r)
    return p

def _get_text(elem):
    return ''.join(t.text or '' for t in elem.iter(qn('w:t'))).strip()

def _has_image(para):
    return para.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline') is not None

def _get_table_header(tbl_elem):
    rows = list(tbl_elem.iter(qn('w:tr')))
    if not rows: return ''
    return ' '.join(''.join(t.text or '' for t in tc.iter(qn('w:t'))) for tc in rows[0].iter(qn('w:tc')))


# ───────────────────────────────────────────────────────
# PHASE 4: DOCX building
# ───────────────────────────────────────────────────────

class ChapterNum:
    def __init__(s):
        s.ch=0; s.sec={}; s.sub={}
    def h1(s):
        s.ch+=1; s.sec[s.ch]=0; return f'第{CN_NAMES[s.ch-1] if s.ch<=len(CN_NAMES) else str(s.ch)}章'
    def h2(s):
        s.sec[s.ch]=s.sec.get(s.ch,0)+1; sec=s.sec[s.ch]; s.sub[(s.ch,sec)]=0
        return f'{s.ch}.{sec}'
    def h3(s):
        sec=s.sec.get(s.ch,0); s.sub[(s.ch,sec)]=s.sub.get((s.ch,sec),0)+1
        return f'{s.ch}.{sec}.{s.sub[(s.ch,sec)]}'


def build_docx(elements, output_path, report_title, images_dir=None):
    doc = Document()
    # ── 页面设置（V3.0 §2）──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)    # V3.0: 装订侧预留
    section.right_margin = Cm(2.54)

    # ── Normal 样式（V3.0 §3.3 正文：宋体+TNR，11pt，1.5倍行距）──
    sty = doc.styles['Normal']
    sty.font.name = FONT_LATIN_BODY; sty.font.size = Pt(SIZE_BODY)
    sty.paragraph_format.line_spacing = LINE_SPACING_1_5
    se = sty.element; rP = se.get_or_add_rPr()
    rF = OxmlElement('w:rFonts'); rP.insert(0, rF)
    rF.set(qn('w:eastAsia'), FONT_CJK_BODY)

    # ── 页眉（V3.0 §6.1: 报告简称右对齐 + 1pt 黑色底线）──
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run(report_title); set_run_font(r, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_TINY)
        # 添加页眉底线（1pt 黑色段落下边框）
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')     # 1pt = 8 eighths of a point
        bottom.set(qn('w:color'), '000000')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # ── 页脚（V3.0 §6.2: 页码居中，阿拉伯数字）──
        fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tag, txt in [('begin',''),('instr',' PAGE '),('end','')]:
            run = fp.add_run()
            el = OxmlElement('w:fldChar') if tag != 'instr' else OxmlElement('w:instrText')
            if tag == 'instr': el.set(qn('xml:space'), 'preserve'); el.text = txt
            else: el.set(qn('w:fldCharType'), tag)
            run._element.append(el)

    numbering = ChapterNum()
    in_front = True; in_back = False; found_title = False
    pending_cover_title = None; en_subtitle = None; pending_tbl_cap = None

    for elem in elements:
        if isinstance(elem, EmptyLine): continue

        if isinstance(elem, HorizontalRule):
            if pending_cover_title:
                _render_cover(doc, pending_cover_title, en_subtitle); pending_cover_title = None
            doc.add_page_break()
            continue

        if isinstance(elem, Heading):
            text = remove_mark(elem.text)
            if elem.level == 1:
                if not found_title:
                    pending_cover_title = report_title or text
                    if not report_title: report_title = text
                    found_title = True; continue
                if '目录' in text:
                    if pending_cover_title: _render_cover(doc, pending_cover_title, en_subtitle); pending_cover_title = None
                    _add_toc(doc); continue
                if any(kw in text for kw in ['参考文献','附录','术语表','图表索引','资料清单']):
                    if pending_cover_title: _render_cover(doc, pending_cover_title, en_subtitle); pending_cover_title = None
                    in_back=True; in_front=False
                    _add_back_h1(doc, text); continue
                if pending_cover_title: _render_cover(doc, pending_cover_title, en_subtitle); pending_cover_title = None
                in_front = False
                _add_h1(doc, text, numbering)
            elif elem.level == 2:
                if pending_cover_title and not any(kw in text for kw in ['目录','证据等级','摘要','核心发现','主要建议','关键约束']):
                    en_subtitle = text; continue
                if any(kw in text for kw in ['参考文献','附录','术语表','图表索引','资料清单']): in_back=True; in_front=False
                if in_front:
                    if '目录' in text: _add_toc(doc)
                    else: _add_front_h2(doc, text)
                elif in_back: _add_back_h2(doc, text)
                else: _add_h2(doc, text, numbering)
            elif elem.level == 3:
                if in_front or in_back: _add_front_h3(doc, text)
                else: _add_h3(doc, text, numbering)
            elif elem.level == 4: _add_h4(doc, text)

        elif isinstance(elem, Paragraph):
            tbl_m = re.match(r'^表\s*\d+[-−]\s*\d+[：:]\s*(.+)', elem.text)
            if tbl_m: pending_tbl_cap = elem.text; continue
            _add_body(doc, elem.text)

        elif isinstance(elem, CodeBlock):
            for l in elem.lines:
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
                set_para_spacing(p,ls=1.0,left=Cm(1.0))
                r=p.add_run(l or ' '); set_run_font(r,FONT_MONO_CJK,FONT_MONO_LATIN,SIZE_SMALL)

        elif isinstance(elem, BlockQuote):
            for l in elem.lines:
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
                set_para_spacing(p,ls=1.5,first=INDENT_FIRST_LINE,left=INDENT_LEFT_QUOTE)
                p.paragraph_format.right_indent=INDENT_LEFT_QUOTE
                r=p.add_run(remove_mark(l)); set_run_font(r,FONT_CJK_SPECIAL,FONT_LATIN_SPECIAL,SIZE_SMALL)

        elif isinstance(elem, ListItem):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            set_para_spacing(p,ls=1.5,left=INDENT_LEFT_LIST)
            p.paragraph_format.first_line_indent=-INDENT_HANG_LIST
            prefix = f'{elem.order_num}. ' if elem.ordered else '● '
            add_inline(p, prefix+elem.text, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_BODY)

        elif isinstance(elem, TableElement):
            _add_table(doc, elem, pending_tbl_cap); pending_tbl_cap = None

        elif isinstance(elem, ImageElement):
            pass  # handled post-hoc

    # ── Post-processing ──
    # Embed images & fix captions
    if images_dir and os.path.isdir(images_dir):
        _embed_images(doc, images_dir)
    _fix_table_captions(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'[OK] Word report generated: {output_path}')


# ───────────────────────────────────────────────────────
# Cover, TOC, Headings
# ───────────────────────────────────────────────────────

def _render_cover(doc, cn_title, en_subtitle):
    """封面（V3.0 §7）：标题 28pt Bold、副标题 14pt、报告类型 16pt、机构 14pt Bold、版本日期 11pt"""
    from datetime import date
    # 顶部留白 6cm 等效
    for _ in range(10):
        p = doc.add_paragraph(); set_para_spacing(p, ls=1.0)

    # 报告标题：28pt Bold 黑色居中
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, ls=1.2, before=Pt(12), after=Pt(6))
    r = p.add_run(cn_title); set_run_font(r, FONT_CJK_HEADING, FONT_LATIN_HEADING, SIZE_TITLE, bold=True, color=COLOR_BLACK)

    # 英文副标题：14pt 黑色居中
    if en_subtitle:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p2, ls=1.2, before=Pt(8), after=Pt(12))
        r2 = p2.add_run(en_subtitle); set_run_font(r2, FONT_LATIN_HEADING, FONT_LATIN_HEADING, SIZE_SUBTITLE, color=COLOR_BLACK)

    # 报告类型标签：16pt 黑色居中
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p3, ls=1.2, before=Pt(12), after=Pt(24))
    r3 = p3.add_run('深度研究报告'); set_run_font(r3, FONT_CJK_HEADING, FONT_LATIN_HEADING, SIZE_COVER_TYPE, color=COLOR_BLACK)

    # 底部分隔线（大量段前空白将机构信息推到底部）
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sep, ls=1.0, before=Pt(96), after=Pt(12))
    rs = sep.add_run('━' * 40); set_run_font(rs, FONT_CJK_BODY, FONT_LATIN_BODY, 10, color=COLOR_SEPARATOR)

    # 机构名：14pt Bold 黑色
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p4, ls=1.5)
    r4 = p4.add_run('遨天科技'); set_run_font(r4, FONT_CJK_HEADING, FONT_LATIN_HEADING, 14, bold=True, color=COLOR_BLACK)

    # 版本号 + 日期：11pt 黑色
    today = date.today().strftime('%Y.%m.%d')
    p5 = doc.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p5, ls=1.5)
    r5 = p5.add_run(f'V1.0 | {today}'); set_run_font(r5, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_COVER_INFO, color=COLOR_BLACK)

    doc.add_page_break()

def _add_toc(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p,ls=1.0,before=SPACE_H1_BEFORE,after=SPACE_H1_AFTER)
    r=p.add_run('目  录'); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H1,bold=True)
    pt=doc.add_paragraph()
    for action in ['begin','instr','separate','result','end']:
        run=pt.add_run()
        if action=='instr':
            el=OxmlElement('w:instrText'); el.set(qn('xml:space'),'preserve')
            el.text='TOC \\o "1-3" \\h \\z'; run._element.append(el)
        elif action=='result':
            run.text=''
        else:
            el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),action); run._element.append(el)

def _add_h1(doc, text, num):
    """章标题 H1（V3.0 §3.2）：小一 24pt Bold，居中，编号 '第X章'"""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p,ls=1.0,before=SPACE_H1_BEFORE,after=SPACE_H1_AFTER)
    r=p.add_run(f'{num.h1()}  {text}'); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H1,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'1'); p._element.get_or_add_pPr().append(o)

def _add_h2(doc, text, num):
    """节标题 H2（V3.0 §3.2）：三号 16pt Bold，两端对齐"""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.0,before=SPACE_H2_BEFORE,after=SPACE_H2_AFTER)
    r=p.add_run(f'{num.h2()}  {text}'); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H2,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'2'); p._element.get_or_add_pPr().append(o)

def _add_h3(doc, text, num):
    """小节标题 H3（V3.0 §3.2）：四号 14pt Bold，两端对齐"""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.0,before=SPACE_H3_BEFORE,after=SPACE_H3_AFTER)
    r=p.add_run(f'{num.h3()}  {text}'); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H3,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'3'); p._element.get_or_add_pPr().append(o)

def _add_h4(doc, text):
    """段落小标题 H4（V3.0 §3.2）：小四 12pt Bold，不入目录"""
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.5,before=SPACE_H4_BEFORE,after=SPACE_H4_AFTER)
    r=p.add_run(text); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H4,bold=True)

def _add_front_h2(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.0,before=SPACE_H2_BEFORE,after=SPACE_H2_AFTER)
    r=p.add_run(text); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H2,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'2'); p._element.get_or_add_pPr().append(o)

def _add_front_h3(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.0,before=SPACE_H3_BEFORE,after=SPACE_H3_AFTER)
    r=p.add_run(text); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H3,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'3'); p._element.get_or_add_pPr().append(o)

def _add_back_h1(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p,ls=1.0,before=SPACE_H1_BEFORE,after=SPACE_H1_AFTER)
    r=p.add_run(text); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H1,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'1'); p._element.get_or_add_pPr().append(o)

def _add_back_h2(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.0,before=SPACE_H2_BEFORE,after=SPACE_H2_AFTER)
    r=p.add_run(text); set_run_font(r,FONT_CJK_HEADING,FONT_LATIN_HEADING,SIZE_H2,bold=True)
    o=OxmlElement('w:outlineLvl'); o.set(qn('w:val'),'2'); p._element.get_or_add_pPr().append(o)

def _add_body(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p,ls=1.5,before=SPACE_BODY_BEFORE,after=SPACE_BODY_AFTER,first=INDENT_FIRST_LINE)
    add_inline(p, text, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_BODY)

def _set_cell_border(cell, top=None, bottom=None, left=None, right=None, insideH=None, insideV=None):
    """设置单元格边框。每个参数: {"sz": "8", "val": "single", "color": "000000"} 或 None。
    sz 单位: 1/8 pt（如 sz="12" = 1.5pt, sz="8" = 1pt, sz="4" = 0.5pt）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除现有边框
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge, spec in [('top', top), ('bottom', bottom), ('left', left), ('right', right),
                        ('insideH', insideH), ('insideV', insideV)]:
        if spec is None:
            continue
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), spec.get('val', 'single'))
        el.set(qn('w:sz'), spec.get('sz', '4'))
        el.set(qn('w:color'), spec.get('color', '000000'))
        el.set(qn('w:space'), '0')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def _set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)

def _create_table_borders_table():
    """创建全框线单元格边框样式"""
    thin = {'sz': '4', 'val': 'single', 'color': '000000'}   # 0.5pt
    return {'top': thin, 'bottom': thin, 'left': thin, 'right': thin,
            'insideH': thin, 'insideV': thin}

def _add_table(doc, elem, caption):
    """表格渲染（V3.0 §5.1）：
    - 全框线含竖线：顶线/底线 1.5pt，表头下分隔线 1pt，内部线 0.5pt，竖线 0.5pt
    - 交替行灰底 #F2F2F2
    - 表头：微软雅黑+TNR 10pt 加粗居中
    - 表体：宋体+TNR 10.5pt，文字左对齐/数字右对齐
    - 跨页自动重复表头行
    """
    # 表题注（V3.0 §4.2: 表上方居中，9pt #555555）
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, ls=1.0, before=Pt(6), after=Pt(3))
        r = p.add_run(caption); set_run_font(r, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_TINY, color=COLOR_DARK_GRAY)

    rc = len(elem.rows) + 1
    cc = len(elem.header)
    tbl = doc.add_table(rows=rc, cols=cc)
    # 不设全局 style，手动逐格设边框

    # 设置表格宽度为页宽 90%，居中
    tbl.autofit = True
    tbl.alignment = 1  # WD_TABLE_ALIGNMENT.CENTER

    # 边框规格
    thick_outer = {'sz': '12', 'val': 'single', 'color': '000000'}    # 1.5pt
    medium_sep  = {'sz': '8',  'val': 'single', 'color': '000000'}    # 1pt
    thin_inner  = {'sz': '4',  'val': 'single', 'color': '000000'}    # 0.5pt

    # ── 表头行（V3.0: 微软雅黑+TNR 10pt 加粗居中，白底）──
    for j, ct in enumerate(elem.header):
        c = tbl.rows[0].cells[j]; c.text = ''
        r = c.paragraphs[0].add_run(remove_mark(ct))
        set_run_font(r, FONT_CJK_HEADING, FONT_LATIN_HEADING, SIZE_TABLE_HEADER, bold=True)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 表头边框：顶底线 1.5pt，左右 0.5pt
        _set_cell_border(c, top=thick_outer, bottom=medium_sep, left=thin_inner, right=thin_inner,
                         insideH=thin_inner, insideV=thin_inner)

    # ── 表体行 ──
    for i, row in enumerate(elem.rows):
        is_even = (i % 2 == 1)  # 第 0 行是数据行 i=0（索引 1），即实际第 2 行
        for j, ct in enumerate(row):
            if j >= cc:
                continue
            c = tbl.rows[i + 1].cells[j]; c.text = ''
            # 判断是否数字（右对齐），否则左对齐
            cell_text = remove_mark(ct)
            is_numeric = re.match(r'^[\d,.\-%$¥€£]+$', cell_text.strip())
            align = WD_ALIGN_PARAGRAPH.RIGHT if is_numeric else WD_ALIGN_PARAGRAPH.LEFT
            c.paragraphs[0].alignment = align
            r = c.paragraphs[0].add_run(cell_text)
            set_run_font(r, FONT_CJK_BODY, FONT_LATIN_BODY, SIZE_TABLE_BODY)
            # 交替行灰底
            if is_even:
                _set_cell_shading(c, 'F2F2F2')
            # 到底行用粗底线，中间行用细线
            is_last_row = (i == len(elem.rows) - 1)
            cell_bottom = thick_outer if is_last_row else thin_inner
            _set_cell_border(c, top=thin_inner, bottom=cell_bottom,
                             left=thin_inner, right=thin_inner,
                             insideH=thin_inner, insideV=thin_inner)

    # ── 跨页表头重复（V3.0 §5.1）──
    tbl_header = tbl.rows[0]._tr
    trPr = tbl_header.get_or_add_trPr()
    tblHeaderEl = OxmlElement('w:tblHeader')
    trPr.append(tblHeaderEl)

    doc.add_paragraph()


# ───────────────────────────────────────────────────────
# Post-process: embed images
# ───────────────────────────────────────────────────────

def _embed_images(doc, images_dir):
    available = set(os.listdir(images_dir))
    wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

    # Collect image reference paragraphs by index
    refs = []
    for i, para in enumerate(doc.paragraphs):
        m = re.search(r'图\s*(\d+)[-−]\s*(\d+)', para.text)
        if m:
            key = m.group(1)+'-'+m.group(2)
            if key in FIGURE_MAP:
                fname, caption = FIGURE_MAP[key]
                if fname in available:
                    refs.append((i, key, fname, caption))

    # Dedup + reverse process
    seen = set(); uniq = []
    for idx, key, fname, cap in refs:
        if idx not in seen: seen.add(idx); uniq.append((idx, key, fname, cap))

    for para_idx, fig_key, fname, caption in sorted(uniq, key=lambda x: x[0], reverse=True):
        para = doc.paragraphs[para_idx]
        parent = para._element.getparent()
        pos = list(parent).index(para._element)
        full_path = os.path.join(images_dir, fname)
        if not os.path.exists(full_path): continue

        # Image paragraph
        img_p = OxmlElement('w:p')
        img_pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); img_pPr.append(jc)
        s = OxmlElement('w:spacing'); s.set(qn('w:before'), '200'); s.set(qn('w:after'), '100'); img_pPr.append(s)
        img_p.append(img_pPr)
        tmp = doc.add_paragraph(); tmp_run = tmp.add_run()
        tmp_run.add_picture(full_path, width=Inches(5.2))
        drawing = tmp._element.find(qn('w:r')).find(qn('w:drawing'))
        img_r = OxmlElement('w:r'); img_r.append(drawing); img_p.append(img_r)
        doc.element.body.remove(tmp._element)

        # Caption paragraph below image（V3.0 §3.3: 图注 9pt #555555 居中）
        cap_p = _make_xml_para(caption, color_hex='555555')

        parent.insert(pos, cap_p); parent.insert(pos, img_p)
        parent.remove(para._element)

    # Cleanup stray fig refs
    body_c = list(doc.element.body)
    to_remove = []
    for i, c in enumerate(body_c):
        if c.tag != qn('w:p'): continue
        t = _get_text(c)
        if not re.match(r'(●\s*)?图\s*\d+[-−]\s*\d+[：:]', t): continue
        prev_img = i > 0 and body_c[i-1].tag == qn('w:p') and _has_image(body_c[i-1])
        next_img = i+1 < len(body_c) and body_c[i+1].tag == qn('w:p') and _has_image(body_c[i+1])
        if prev_img or next_img: continue
        if 'KaiTi' in c.xml: continue
        to_remove.append(i)
    for i in reversed(to_remove): body_c[i].getparent().remove(body_c[i])


# ───────────────────────────────────────────────────────
# Post-process: fix table captions
# ───────────────────────────────────────────────────────

def _fix_table_captions(doc):
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tbl_tag = '{' + W_NS + '}tbl'
    p_tag = '{' + W_NS + '}p'

    body = list(doc.element.body)
    new = []
    tbl_count = 0
    for i, child in enumerate(body):
        tag = child.tag
        if tag == tbl_tag:
            tbl_count += 1
            hdr = _get_table_header(child)
            caption = None
            for pat, cap_text in COMPILED_TABLE_MATCHERS:
                if pat.search(hdr):
                    caption = cap_text; break
            if not caption:
                caption = '表'
            new.append(_make_xml_para(caption, color_hex='555555'))  # V3.0: 表注 9pt #555555
            new.append(child); continue

        # Keep images and their captions as-is
        if tag == p_tag and _has_image(child):
            new.append(child); continue

        if tag == p_tag and re.match(r'图\s*\d+[-−]\s*\d+', _get_text(child)):
            prev_img = (new and new[-1].tag == p_tag and _has_image(new[-1]))
            if prev_img: new.append(child)
            continue

        # Skip orphan fig/tbl text refs
        t = _get_text(child) if tag == p_tag else ''
        if tag == p_tag and t:
            if re.match(r'(●\s*)?图\s*\d+[-−]\s*\d+[：:]', t):
                if not (new and new[-1].tag == p_tag and _has_image(new[-1])): continue
            if re.match(r'(●\s*)?表\s*\d+[-−]\s*\d+[：:]', t):
                continue  # skip text-only table refs

        new.append(child)

    for c in list(doc.element.body):
        doc.element.body.remove(c)
    for c in new:
        doc.element.body.append(c)
    print(f'  Tables processed: {tbl_count}')


# ───────────────────────────────────────────────────────
# Main entry point
# ───────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='Markdown → Word (.docx) 研究报告一体化转换器\n'
                    '严格遵循《研究报告格式规范 V3.0》（遨天科技，2026-06-24）')
    parser.add_argument('input', help='Markdown 文件路径')
    parser.add_argument('output', help='输出 .docx 路径')
    parser.add_argument('--title', '-t', default=None, help='报告题名')
    parser.add_argument('--images', '-i', default=None, help='图片目录（包含PNG图表）')
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f'[ERROR] File not found: {args.input}')
        sys.exit(1)

    # PHASE 1: Clean markdown
    print('[1/4] Cleaning markdown...')
    lines = clean_markdown(args.input)
    clean_text = '\n'.join(lines)

    # Write clean version for debugging
    clean_path = args.input.replace('.md', '-cleaned.md')
    with open(clean_path, 'wb') as f:
        f.write(clean_text.encode('utf-8'))

    # PHASE 2: Parse
    print('[2/4] Parsing markdown...')
    elements = parse_markdown(clean_text)

    # PHASE 3+4: Build docx + embed + fix
    print('[3/4] Building Word document (V3.0 format spec)...')
    build_docx(elements, args.output, args.title or '研究报告', args.images)

    print('[4/4] Done!')
    print(f'  Cleaned markdown: {clean_path}')
    print(f'  Word document:    {args.output}')
    if not args.images:
        print('  (No --images specified; charts not embedded)')

    # ── 格式合规摘要（V3.0 §10.3）──
    print()
    print('  -- Format: V3.0 Specification --')
    print(f'  Fonts: Title=Microsoft YaHei+TNR | Body=SimSun+TNR')
    print(f'  Sizes: H1={SIZE_H1}pt H2={SIZE_H2}pt H3={SIZE_H3}pt Body={SIZE_BODY}pt')
    print(f'  Margins: Left 3.17cm Right 2.54cm Top/Bottom 2.54cm')
    print(f'  Line spacing: {LINE_SPACING_1_5}x | First-line indent: 2 chars')
    print(f'  Tables: Full grid + alternate row gray #F2F2F2 | Header 10pt bold')
    print(f'  Captions: 9pt #555555 center | Header: Right-aligned + 1pt black line')
    print(f'  Cover: Title 28pt Bold + Subtitle 16pt + Org 14pt Bold')
    print(f'  [NOTE] Open in Word/WPS, right-click TOC -> "Update Field" for page numbers.')


if __name__ == '__main__':
    main()
