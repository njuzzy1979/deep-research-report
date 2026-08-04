# -*- coding: utf-8 -*-
"""tests/test_outline_skeleton.py —— D1-8 骨架 docx 预确认测试。

重点覆盖 §9.2 的"致命前提"：section 数据全空时**必须拒绝产出**，避免制造
"我在阶段 4 已经确认过了"的虚假安全感。
"""
from __future__ import annotations

from pathlib import Path

import pytest

import outline_skeleton as sk

_GOOD = '''---
struct_template: research
title: 测试报告题名
structure:
  frontmatter: []
  bodymatter:
    - chapter_no: 1
      chapter_title: 空间智能演化态势
      sections:
        - section_no: "1.1"
          section_title: 研究背景与意义
        - section_no: "1.2"
          section_title: 研究目标与范围
    - chapter_no: 2
      chapter_title: 理论体系构建
      sections:
        - section_no: "2.1"
          section_title: 六层认知升维模型
        - section_no: "2.2"
          section_title: 七维问题空间
  appendix: []
---
## 第 1 章：空间智能演化态势

## 第 2 章：理论体系构建
'''

# 真实事故形态：声明了章但 sections 全空（实测 16/16）
_EMPTY_SECTIONS = '''---
title: 测试报告题名
structure:
  bodymatter:
    - chapter_no: 1
      chapter_title: 空间智能演化态势
      subsections: []
    - chapter_no: 2
      chapter_title: 理论体系构建
      subsections: []
  appendix: []
---
'''


def _w(tmp_path, body, name="outline.md"):
    p = tmp_path / name
    p.write_text(body, encoding="utf-8")
    return str(p)


# ── §9.2 致命前提：section 全空必须拒绝 ─────────────────────


def test_refuses_when_no_sections_declared(tmp_path):
    """未经 D1-9 逼出 section 数据时，骨架只有章标题、节层空白——
    用户确认这样一份骨架比不做本功能更危险，故默认拒绝产出。"""
    r = sk.generate_skeleton(_w(tmp_path, _EMPTY_SECTIONS))
    assert r["passed"] is False
    assert "0 个节" in r["error"]
    assert "虚假安全感" in r["error"]
    assert not (tmp_path / sk.SKELETON_DOCX_RELNAME).exists()


def test_allow_empty_sections_flag_forces_output(tmp_path):
    """显式放行时才产出，且骨架内对无节的章就地标注告警。"""
    r = sk.generate_skeleton(_w(tmp_path, _EMPTY_SECTIONS), allow_empty_sections=True)
    assert r["passed"] is True
    md = Path(r["skeleton_md"]).read_text(encoding="utf-8")
    assert "未声明任何节" in md
    assert r["warnings"], "应给出无节章的告警"


# ── 骨架 Markdown 合成 ───────────────────────────────────────


def test_builds_markdown_with_correct_levels(tmp_path):
    import yaml
    from md2docx.assemble.outline_reader import extract_yaml_front_matter
    parsed, _ = extract_yaml_front_matter(_GOOD)
    text, stats = sk.build_skeleton_markdown(parsed["structure"], "测试报告题名")

    assert text.startswith("# 测试报告题名")
    # 章 -> ##（docx Heading 1）、节 -> ###（docx Heading 2）
    assert "## 空间智能演化态势" in text
    assert "### 研究背景与意义" in text
    assert stats["chapter_count"] == 2
    assert stats["section_count"] == 4
    assert stats["chapters_without_sections"] == []
    _ = yaml


def test_placeholder_line_present_under_each_section(tmp_path):
    """占位提示行是必需的：既让用户看到"此处待填充"语义，也使节下不是纯空白。"""
    from md2docx.assemble.outline_reader import extract_yaml_front_matter
    parsed, _ = extract_yaml_front_matter(_GOOD)
    text, stats = sk.build_skeleton_markdown(parsed["structure"], "T")
    assert text.count(sk.SECTION_PLACEHOLDER) == stats["section_count"]


def test_no_body_text_beyond_placeholders(tmp_path):
    """骨架**不写任何正文**——只有标题与占位行。

    参考文献节的说明行是一处刻意的例外：它不是 SECTION_PLACEHOLDER 常量
    （因为内容不同——说明该节由阶段9自动生成，而非"待阶段7填充"），但同样
    是引用块占位说明而非正文，故单独放行。
    """
    from md2docx.assemble.outline_reader import extract_yaml_front_matter
    parsed, _ = extract_yaml_front_matter(_GOOD)
    text, _stats = sk.build_skeleton_markdown(parsed["structure"], "T")
    references_note_line = (
        "> （本节由阶段9 convert_references.py/finalize_pipeline.py "
        "自动生成并插入统一参考文献列表，无需在 outline.md 中声明或在阶段7手动撰写）"
    )
    for line in text.split("\n"):
        s = line.strip()
        if s and not s.startswith("#") and s != sk.SECTION_PLACEHOLDER and s != references_note_line:
            pytest.fail(f"骨架含非占位正文: {s!r}")


def test_placeholder_contains_no_secrecy_keyword():
    """gate3._check_secrecy 是剔除型门禁（出现密级词即 FATAL）。

    本 skill 既定立场是产物一律不带密级标注，骨架占位文案须避开这类词。
    """
    for word in sk._FORBIDDEN_IN_PLACEHOLDER:
        assert word not in sk.SECTION_PLACEHOLDER


def test_appendix_rendered_with_letter(tmp_path):
    from md2docx.assemble.outline_reader import extract_yaml_front_matter
    body = _GOOD.replace(
        "  appendix: []",
        '  appendix:\n    - appendix_letter: A\n      appendix_title: 缩略语对照\n',
    )
    parsed, _ = extract_yaml_front_matter(body)
    text, _s = sk.build_skeleton_markdown(parsed["structure"], "T")
    assert "## 附录A：缩略语对照" in text


# ── 产物落位约定（D3 §5.3 三条强制约束）────────────────────


def test_artifact_naming_and_placement(tmp_path):
    r = sk.generate_skeleton(_w(tmp_path, _GOOD))
    assert r["passed"] is True

    # 1. 文件名含 skeleton 与 preview 双重语义标识
    assert "skeleton" in sk.SKELETON_DOCX_RELNAME
    assert "preview" in sk.SKELETON_DOCX_RELNAME

    # 2. 中间 md 为点号前缀隐藏件，且**不得命名为 final-report***
    assert sk.SKELETON_MD_RELNAME.startswith(".")
    assert not sk.SKELETON_MD_RELNAME.startswith("final-report")

    # 3. docx 落 research/（outline 同级），**不进 output/**
    assert Path(r["skeleton_docx"]).parent == tmp_path
    assert "output" not in Path(r["skeleton_docx"]).parts
    # 中间 md 落 drafts/
    assert Path(r["skeleton_md"]).parent == tmp_path / "drafts"


def test_docx_hierarchy_maps_chapters_to_heading1(tmp_path):
    """层级映射实测：## -> Heading 1（章）、### -> Heading 2（节）。

    骨架生成器固定在正文章节之后、附录之前插入"参考文献"占位节。
    2026-08-04：HeadingKind.REFERENCES 已从 _KIND_TO_LEVEL 移除，
    参考文献不再渲染为 Heading 1（改为普通 H2→Heading 2），h1 列表不包含它。
    """
    pytest.importorskip("docx")
    from docx import Document
    r = sk.generate_skeleton(_w(tmp_path, _GOOD))
    assert r["passed"] is True, r.get("error")
    d = Document(r["skeleton_docx"])
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    assert h1 == ["空间智能演化态势", "理论体系构建"]
    assert h2 == ["研究背景与意义", "研究目标与范围", "六层认知升维模型", "七维问题空间"]


def test_skeleton_generation_is_idempotent(tmp_path):
    """同一 outline 重复生成，骨架 md 内容逐字相同（无时间戳/随机数）。"""
    o = _w(tmp_path, _GOOD)
    a = sk.generate_skeleton(o)
    first = Path(a["skeleton_md"]).read_text(encoding="utf-8")
    b = sk.generate_skeleton(o)
    second = Path(b["skeleton_md"]).read_text(encoding="utf-8")
    assert first == second


# ── 用法层面错误 ─────────────────────────────────────────────


def test_missing_outline_returns_error(tmp_path):
    r = sk.generate_skeleton(str(tmp_path / "nope.md"))
    assert r["passed"] is False and "error" in r


def test_missing_structure_node_returns_error(tmp_path):
    r = sk.generate_skeleton(_w(tmp_path, "---\ntitle: 无 structure\n---\n"))
    assert r["passed"] is False and "structure" in r["error"]


# ── D2-9 接口对齐：骨架生成不得自行调用 python-docx ──────────


def test_does_not_import_python_docx_directly():
    """§9.4.5 约束：骨架只合成 Markdown，docx 一律交给 python -m md2docx。

    这是登记在案的**合法调用正样本**（规则一的两条不命中路径之一）；若本脚本
    自行 `from docx import Document`，会与 D2-9 规则一的语义冲突。
    """
    src = Path(sk.__file__).read_text(encoding="utf-8")
    # 允许注释中提及，检查实际 import 语句
    for line in src.split("\n"):
        s = line.strip()
        if s.startswith(("from docx", "import docx")):
            pytest.fail(f"骨架生成器不得直接 import python-docx: {s!r}")
    assert "md2docx" in src
