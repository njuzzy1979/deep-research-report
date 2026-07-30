# -*- coding: utf-8 -*-
"""tests/test_finalize_pipeline.py —— D5 定稿顺序管道测试。

覆盖：
  6 个 failure_step 在各自步骤失败时被正确返回（strip_markers/h1_check/merge/
  convert_refs/contract_check/delivery_checklist），且 steps 字典中该步骤
  status=fail、reason 非空
  成功路径全绿：构造最小但完整的 drafts/outline/source-index 三件套，走完整
  6 步，断言 overall_pass=True、failure_step=None、output_path 指向实际生成
  的文件
  CLI 退出码语义（0/1/2）
  台账隔离（不触碰真实项目台账文件）

注：contract_check 步骤的成功路径刻意构造为"无引用"样本（不含 [SRC-XXX]），
因为已实测确认 contract_check.py 的 C6（`_check_c6_references`）对纯数字
引用 `[N]` 本身判负（`pure_num_hits` 命中即 `pass=False`），与 convert_refs
步骤转换后必然产生纯数字引用这一事实存在真实冲突（详见实现报告"发现的问题"
一节）。为了不让这个既有的 contract_check.py 行为污染"D5 管道本身工作正常"
这一测试目标，成功路径样本不引入任何引用；C6 冲突单独用一个明确标注为
"已知问题重现"的测试记录下来，断言其确实复现（而非静默略过）。
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import pytest

import finalize_pipeline as fp


def _write_outline(tmp_path):
    """frontmatter-only 结构（无 bodymatter 编号章节）。

    刻意不使用 bodymatter 章节：已实测确认 merge_drafts.assemble_merged()
    按规范插入的章容器 `## 第 X 章：<chapter_title>` 本身会命中
    contract_check.py 的 C2（MANUAL_NUMBER_PATTERN 匹配"第\\s*N...章"），
    且 merged=True + stage="stage9" 下 severity=fatal——这是 merge_drafts.py
    与 contract_check.py 两个既有组件之间的真实冲突（此前被 merge_drafts.py
    "阶段E只WARN不阻断"的反模式掩盖，D5 管道使其第一次真正可见，见实现报告
    "发现的问题"）。用 frontmatter-only 结构隔离"管道机制本身是否正常"这一
    测试目标；该 C2 冲突单独由 test_failure_step_contract_check_reproduces_
    chapter_container_c2_conflict 明确记录复现。
    """
    outline = tmp_path / "outline.md"
    outline.write_text(
        "---\n"
        "struct_template: research\n"
        "title: 测试报告\n"
        "structure:\n"
        "  frontmatter:\n"
        "    - chapter_title: 前言/导论\n"
        "      sections:\n"
        "        - section_no: \"\"\n"
        "          section_title: 问题提出与研究背景\n"
        "  bodymatter: []\n"
        "  appendix: []\n"
        "---\n",
        encoding="utf-8",
    )
    return outline


def _write_outline_with_chapter(tmp_path):
    """含 bodymatter 编号章节的结构，专供 C2 冲突复现测试使用。"""
    outline = tmp_path / "outline.md"
    outline.write_text(
        "---\n"
        "struct_template: research\n"
        "title: 测试报告\n"
        "structure:\n"
        "  frontmatter: []\n"
        "  bodymatter:\n"
        "    - chapter_no: 1\n"
        "      chapter_title: 第一章测试\n"
        "      sections:\n"
        "        - section_no: \"1.1\"\n"
        "          section_title: 测试节\n"
        "  appendix: []\n"
        "---\n",
        encoding="utf-8",
    )
    return outline


def _write_source_index(tmp_path):
    csv_path = tmp_path / "source-index.csv"
    csv_path.write_text(
        "source_id,title,author_or_org,publisher,publish_date,source_type,url_or_path\n"
        "SRC-001,测试来源标题,测试作者,测试出版社,2024,book,\n",
        encoding="utf-8",
    )
    return csv_path


def _write_clean_draft(drafts_dir):
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 测试节\n\n正文内容，无任何引用残留，干净通过，字数足够长一些用于测试。\n",
        encoding="utf-8",
    )


# ── 成功路径全绿 ─────────────────────────────────────────────

def test_success_path_all_steps_pass(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)
    output = tmp_path / "final-report.md"

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(output),
    )

    assert result["overall_pass"] is True
    assert result["failure_step"] is None
    assert result["output_path"] is not None
    assert output.exists()
    # verify_docx（第 7 步，D2-7）是可选步：只在传入 docx 路径时执行，
    # 本用例不传，故不参与"每步都 pass"的断言。
    for step in fp.FAILURE_STEPS:
        if step == "verify_docx":
            assert step not in result["steps"]
            continue
        assert result["steps"][step]["status"] == "pass"


# ── 6 个 failure_step 枚举分别验证 ───────────────────────────

def test_failure_step_strip_markers_when_drafts_dir_missing(tmp_path):
    result = fp.run_finalize_pipeline(
        drafts_dir=str(tmp_path / "does-not-exist"),
        outline_path=str(tmp_path / "outline.md"),
        source_index_path=str(tmp_path / "source-index.csv"),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    assert result["failure_step"] == "strip_markers"


def test_failure_step_merge_when_outline_yaml_missing_structure(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    # outline.md 缺少 structure 节点 —— parse_outline_yaml 内部 sys.exit(2)
    outline = tmp_path / "outline.md"
    outline.write_text("---\ntitle: 无 structure 节点\n---\n", encoding="utf-8")
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    assert result["failure_step"] == "merge"
    assert "sys.exit" in result["failure_reason"] or "解析失败" in result["failure_reason"]


def test_failure_step_merge_when_outline_yaml_malformed(tmp_path, monkeypatch):
    """outline.md YAML 语法错误（非结构缺失，而是解析异常）同样应捕获为
    merge 步骤失败，而不是让 SystemExit(2) 击穿整个管道进程。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = tmp_path / "outline.md"
    # 制造非法 YAML：未闭合的方括号
    outline.write_text("---\nstructure: [unclosed\n---\n", encoding="utf-8")
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    assert result["failure_step"] == "merge"


def test_failure_step_convert_refs_when_source_index_missing(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(tmp_path / "does-not-exist-source-index.csv"),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    assert result["failure_step"] == "convert_refs"


def test_failure_step_convert_refs_when_slash_refs_present(tmp_path, monkeypatch):
    """斜杠分隔 SRC 引用（[SRC-001/026]）不支持自动转换，应在 convert_refs
    步骤被检出并阻断（复用 convert_references.find_slash_refs_in_file）。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 测试节\n\n正文引用了斜杠分隔格式 [SRC-001/SRC-002]。\n",
        encoding="utf-8",
    )
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    assert result["failure_step"] == "convert_refs"


def test_failure_step_contract_check_reproduces_pure_num_conflict(tmp_path, monkeypatch):
    """回归测试（原为"已知问题重现"，缺陷修复后转为正向验证）：

    历史问题（G7 端到端验证发现的 P0，与 C2 章容器冲突同构）：drafts 中含
    ``[SRC-001]`` 引用，``convert_refs`` 步骤按 GB/T 7714 顺序编码制正确将其
    转换为纯数字 ``[1]``，但紧接着的 ``contract_check`` 步骤又把"纯数字引用"
    本身判为 C6 违规——**定稿管道否定自己上一步的正确产出**。后果：任何含至少
    一条参考文献的真实报告都无法走完管道，``delivery_checklist``（含降级台账
    确认、红队确认、全文通读确认）永远不可达。

    修复（编排器裁决）：``_check_c6_references()`` 改为**分阶段**判定，与 C7 已有的
    stage7/stage9 对称设计保持一致——
    - stage7 分章草稿：纯数字引用仍是违规（作者不应提前写死编号，会与自动编号冲突）
    - stage9 + merged 合并终稿：纯数字引用是 ``convert_references.py`` 的预期产出，不判负
    ``slash_src`` / ``s_variant`` 两类真正的格式错误在任何阶段都仍然判负。

    本测试验证：含引用的报告能顺利通过 contract_check 的 C6 判定并走完管道。
    对"stage7 仍拦截纯数字引用"的验证见
    test_c6_still_blocks_pure_num_in_stage7_draft。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 测试节\n\n正文引用了数据来源[SRC-001]，内容详实可靠。\n",
        encoding="utf-8",
    )
    outline = _write_outline_with_chapter(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    c6 = result["steps"]["contract_check"]["detail"]["contract"]["C6_reference_format"]
    # 纯数字引用在 stage9 合并终稿中被识别为预期产出，不再判负
    assert c6["pure_num_expected"] is True
    assert c6["pass"] is True, f"C6 仍判负，命中：{c6.get('pure_num_hits')}"
    # convert_refs 确实做了转换（存在纯数字引用），只是不再被判为违规
    assert c6["pure_num_count"] >= 1
    # contract_check 不再是阻断点，管道得以继续
    assert result["failure_step"] != "contract_check"


def test_c6_still_blocks_pure_num_in_stage7_draft():
    """豁免精确性回归：stage9 豁免**不得**削弱 C6 对分章草稿的拦截。

    直接调用 ``check_contract()`` 验证三种情形：
    - stage7 草稿：纯数字引用仍判负（作者提前写死编号）
    - stage9 合并终稿：纯数字引用放行
    - stage9 合并终稿：斜杠 SRC / S 变体等真正的格式错误**仍然**判负
    """
    import contract_check as cc

    draft = "# T\n\n## 第 1 章：绪论\n\n正文引用[1]与[2,3]。\n"
    # stage7：仍判负
    r7 = cc.check_contract(draft, merged=False, expect_figures=None, stage="stage7")
    c6_7 = r7["contract"]["C6_reference_format"]
    assert c6_7["pure_num_expected"] is False
    assert c6_7["pass"] is False, "stage7 草稿的纯数字引用未被拦截——豁免范围过宽"

    # stage9 + merged：放行
    r9 = cc.check_contract(draft, merged=True, expect_figures=None, stage="stage9")
    assert r9["contract"]["C6_reference_format"]["pass"] is True

    # stage9 + merged 但含真正的格式错误：仍判负
    bad = "# T\n\n## 第 1 章：绪论\n\n正文[SRC-001/026]与[S001]。\n"
    r9b = cc.check_contract(bad, merged=True, expect_figures=None, stage="stage9")
    c6_9b = r9b["contract"]["C6_reference_format"]
    assert c6_9b["pass"] is False, "斜杠/S变体格式错误被错误放行"
    assert c6_9b["slash_src_hits"] and c6_9b["s_variant_hits"]


def test_contract_check_reproduces_chapter_container_c2_conflict(tmp_path, monkeypatch):
    """回归测试（原为"已知问题重现"，缺陷修复后转为正向验证）：

    历史问题：``merge_drafts.assemble_merged()`` 按 references/stage-7-writing.md
    规范在阶段9合并时统一插入的标准章容器 ``## 第 {c_no} 章：{c_title}``，其字面
    结构必然命中 ``contract_check.MANUAL_NUMBER_PATTERN``，且在 merged=True +
    stage="stage9" 下 C2 severity 升级为 "fatal"（标注"不可降级放行"）。这导致
    **任何含至少一个编号章节的报告**在 contract_check 步骤必然失败——即检查器
    把自家合并管道的标准输出判为致命错误。

    修复（编排器裁决）：``contract_check.PIPELINE_CHAPTER_CONTAINER_PATTERN`` 在
    merged 模式下豁免管道自动生成的章容器。C2 的立法意图是禁止**作者手写**编号
    前缀（应交由 Word 自动编号域生成），而章容器是管道自身按规范产出的结构性标记，
    下游 md2docx 会将其识别为 CHAPTER 并接管编号。

    本测试验证：含编号章节的报告能顺利通过 contract_check 的 C2 判定。
    对"作者手写编号仍被拦截"的验证见
    test_c2_still_blocks_author_written_manual_numbering。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 测试节\n\n正文内容干净无残留，不含任何引用，字数足够长一些用于测试。\n",
        encoding="utf-8",
    )
    outline = _write_outline_with_chapter(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    c2 = result["steps"]["contract_check"]["detail"]["contract"]["C2_manual_number"]
    # 章容器已被豁免：C2 不再因管道自身产出的 `## 第 1 章：xxx` 而判负
    assert c2["pass"] is True, f"章容器未被豁免，C2 命中：{c2['hits']}"
    assert not any("第 1 章" in h or "第1章" in h for h in c2["hits"])
    # contract_check 步骤不再是阻断点
    assert result["failure_step"] != "contract_check"


def test_c2_still_blocks_author_written_manual_numbering(tmp_path, monkeypatch):
    """豁免精确性回归：章容器豁免**不得**削弱 C2 对作者手写编号的拦截。

    构造一份正文含 `### 1.1 手写编号节` 的草稿，确认在 merged+stage9 下
    C2 仍判负且 severity 为 fatal——证明豁免范围严格限定为"管道章容器"
    这一种确定格式，没有把 C2 整体放宽。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 1.1 手写编号节\n\n正文内容干净无残留，字数足够长一些用于测试。\n",
        encoding="utf-8",
    )
    outline = _write_outline_with_chapter(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    c2 = result["steps"]["contract_check"]["detail"]["contract"]["C2_manual_number"]
    assert c2["pass"] is False, "作者手写编号未被拦截——豁免范围过宽"
    assert c2["severity"] == "fatal"
    assert any("1.1" in h for h in c2["hits"])


def test_delivery_checklist_step_unreachable_via_pipeline_due_to_c2_conflict(tmp_path, monkeypatch):
    """回归测试（原为"已知问题重现"，C2 冲突修复后恢复其原始测试意图）：

    历史问题：因 C2 章容器冲突（见
    test_contract_check_reproduces_chapter_container_c2_conflict），任何含
    bodymatter 的报告都在 contract_check 步骤被提前拦下，导致 delivery_checklist
    这道最后关卡经由完整管道**根本不可达**。

    C2 豁免修复后，本测试恢复原始意图：验证 delivery_checklist 确实是管道
    最后一道独立关卡——构造一个能通过 contract_check、但会被 D6 清单第 06 项
    （写作者自声明剥离）拦下的样本，确认 failure_step 正确指向
    delivery_checklist 而非更早的步骤。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    (drafts_dir / "ch01-1-1-测试节.md").write_text(
        "### 测试节\n\n正文内容干净无残留，字数足够长一些用于测试。\n\n"
        "## 写作者自声明\n\n本文由AI生成，不代表任何官方立场。\n",
        encoding="utf-8",
    )
    outline = _write_outline_with_chapter(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is False
    # delivery_checklist 现在可达，且正确成为阻断点
    assert result["failure_step"] == "delivery_checklist", (
        f"delivery_checklist 未被触达或未正确阻断，实际 failure_step="
        f"{result['failure_step']}，已执行步骤={list(result['steps'].keys())}"
    )
    assert "delivery_checklist" in result["steps"]
    assert "06_writer_selfclaim_stripped" in (
        result["steps"]["delivery_checklist"]["detail"]["failed_items"]
    )


def test_delivery_checklist_function_detects_writer_selfclaim_directly(tmp_path, monkeypatch):
    """绕开完整管道（不经过 contract_check 这道会因 C2 冲突提前阻断的关卡），
    直接调用 delivery_checklist_check.run_delivery_checklist() 验证该函数
    自身对写作者自声明残留的检测逻辑是正确的——证明 D6 聚合脚本本身没有
    缺陷，问题完全出在"管道内它被 contract_check 步骤挡在前面、根本触达
    不到"这一编排层面。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    merged = tmp_path / "final-report.md"
    merged.write_text(
        "### 测试节\n\n正文内容干净无残留，字数足够长一些用于测试。\n\n"
        "## 写作者自声明\n\n本文由AI生成，不代表任何官方立场。\n",
        encoding="utf-8",
    )
    from delivery_checklist_check import run_delivery_checklist

    result = run_delivery_checklist(str(merged))
    assert result["overall_pass"] is False
    assert "06_writer_selfclaim_stripped" in result["failed_items"]


# ── h1_check 步骤：正向验证 H1 被正确替换为 H2 ──────────────────

def test_h1_check_replaces_h1_with_h2(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    draft_file = drafts_dir / "ch01-1-1-测试节.md"
    draft_file.write_text(
        "# 误用的一级标题\n\n### 测试节\n\n正文内容干净无残留，字数足够长一些用于测试。\n",
        encoding="utf-8",
    )
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["steps"]["h1_check"]["status"] == "pass"
    assert str(draft_file) in result["steps"]["h1_check"]["files_with_h1_replaced"]
    # 文件已被写回，H1 应变为 H2
    new_text = draft_file.read_text(encoding="utf-8")
    assert "## 误用的一级标题" in new_text
    assert not new_text.startswith("# 误用的一级标题")


# ── strip_markers 步骤：.bak 备份 + 清洗生效 ─────────────────────

def test_strip_markers_creates_bak_and_cleans(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    draft_file = drafts_dir / "ch01-1-1-测试节.md"
    original_text = (
        "### 测试节\n\n正文内容干净无残留，字数足够长一些用于测试。\n\n"
        "全文约 1200 字\n"
    )
    draft_file.write_text(original_text, encoding="utf-8")
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)

    fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(tmp_path / "final-report.md"),
    )

    bak_file = drafts_dir / "ch01-1-1-测试节.md.bak"
    assert bak_file.exists()
    assert bak_file.read_text(encoding="utf-8") == original_text
    cleaned_text = draft_file.read_text(encoding="utf-8")
    assert "全文约 1200 字" not in cleaned_text


# ── CLI 退出码语义 ───────────────────────────────────────────────

def test_cli_exit_2_when_drafts_dir_missing(tmp_path, monkeypatch):
    monkeypatch.setattr(sys, "argv", [
        "finalize_pipeline.py",
        "--drafts-dir", str(tmp_path / "does-not-exist"),
        "--outline", str(tmp_path / "outline.md"),
        "--source-index", str(tmp_path / "source-index.csv"),
        "--output", str(tmp_path / "final-report.md"),
    ])
    with pytest.raises(SystemExit) as exc_info:
        fp.main()
    assert exc_info.value.code == 2


def test_cli_exit_0_when_success(tmp_path, monkeypatch, capsys):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)
    output = tmp_path / "final-report.md"

    monkeypatch.setattr(sys, "argv", [
        "finalize_pipeline.py",
        "--drafts-dir", str(drafts_dir),
        "--outline", str(outline),
        "--source-index", str(source_index),
        "--output", str(output),
        "--json",
    ])
    with pytest.raises(SystemExit) as exc_info:
        fp.main()
    assert exc_info.value.code == 0
    out = capsys.readouterr().out
    data = json.loads(out)
    assert data["overall_pass"] is True


def test_cli_exit_1_when_step_fails(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)

    monkeypatch.setattr(sys, "argv", [
        "finalize_pipeline.py",
        "--drafts-dir", str(drafts_dir),
        "--outline", str(outline),
        "--source-index", str(tmp_path / "does-not-exist-source-index.csv"),
        "--output", str(tmp_path / "final-report.md"),
    ])
    with pytest.raises(SystemExit) as exc_info:
        fp.main()
    assert exc_info.value.code == 1


# ── D2-8：失败时不留半成品 ───────────────────────────────────


def test_d2_8_failure_leaves_no_official_output_but_keeps_partial(tmp_path, monkeypatch):
    """核心不变量：正式产物名的存在本身即等价于 overall_pass=True。

    事故第 2 步留下 388 字符的空 final-report.md，直接诱发第 3 步"我来手动
    修一下"。此用例断言失败时正式产物名**不存在**、而 .partial **保留**供诊断。
    """
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    output = tmp_path / "final-report.md"

    # source-index 缺失 → convert_refs 步失败（此时 merge 已写出 .partial）
    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(tmp_path / "no-such-index.csv"),
        output_path=str(output),
    )

    assert result["overall_pass"] is False
    assert not output.exists(), "失败时不得留下正式产物名的半成品"
    partial = tmp_path / "final-report.md.partial"
    assert partial.exists(), ".partial 失败时须保留供诊断"
    assert result["partial_path"] is not None


def test_d2_8_success_promotes_partial_and_removes_it(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    source_index = _write_source_index(tmp_path)
    output = tmp_path / "final-report.md"

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(source_index),
        output_path=str(output),
    )

    assert result["overall_pass"] is True
    assert output.exists()
    assert not (tmp_path / "final-report.md.partial").exists(), "转正后 .partial 应已消失"


def test_d2_8_run_id_is_deterministic(tmp_path):
    """run_id 必须由内容派生、不含随机数/时间戳（否则打破 md2docx G-11 幂等）。"""
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    a = fp._derive_run_id(str(outline), str(drafts_dir))
    b = fp._derive_run_id(str(outline), str(drafts_dir))
    assert a == b
    assert len(a) == 12 and all(c in "0123456789abcdef" for c in a)


def test_d2_8_stale_previous_output_is_renamed_on_failure(tmp_path, monkeypatch):
    """失败且上次成功产物仍在时，主动改名为 .stale-<run_id>，而非仅告警。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    outline = _write_outline(tmp_path)
    output = tmp_path / "final-report.md"
    output.write_text("上一次成功的正式产物", encoding="utf-8")

    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(outline),
        source_index_path=str(tmp_path / "no-such-index.csv"),
        output_path=str(output),
    )

    assert result["overall_pass"] is False
    assert not output.exists(), "旧产物须被改名，避免被当作本次结果"
    assert result["staled_previous_output"] is not None
    assert Path(result["staled_previous_output"]).exists()


# ── D2-7：docx 回读校验 ──────────────────────────────────────

pytest.importorskip("docx", reason="python-docx 未安装时跳过 D2-7 用例")


def _mk_docx(path, paras):
    from docx import Document
    d = Document()
    for text, style in paras:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    d.save(str(path))
    return str(path)


def test_d2_7_catches_accident_form_h1_followed_by_h1(tmp_path):
    """事故形态：章标题下 0 字符，正文全部过继给章内第一个小节标题。"""
    p = _mk_docx(tmp_path / "a.docx", [
        ("第 1 章：导论", "Heading 1"),
        ("本章结论", "Heading 1"),
        ("这里有几千字正文。", None),
    ])
    r = fp.verify_docx_structure(p, 1)
    assert r["pass"] is False
    assert "第 1 章：导论" in r["empty_headings"]


def test_d2_7_catches_skeleton_only_form(tmp_path):
    """D1 §9.4.4 发现的原设计漏检：只有 H1/H2、完全无正文的骨架 docx。

    原实现 `elif prev is not None: buf.append(p.text)` 会把 Heading 2 的标题
    文本当正文收集，实测得 pass=True。加 not startswith("Heading") 后修复。
    """
    p = _mk_docx(tmp_path / "b.docx", [
        ("第 1 章", "Heading 1"),
        ("1.1 某节", "Heading 2"),
        ("第 2 章", "Heading 1"),
        ("2.1 某节", "Heading 2"),
    ])
    r = fp.verify_docx_structure(p, 2)
    assert r["pass"] is False, "只有骨架无正文的 docx 必须被判失败"
    assert len(r["empty_headings"]) == 2


def test_d2_7_catches_duplicate_chapter_headings(tmp_path):
    p = _mk_docx(tmp_path / "c.docx", [
        ("本章结论", "Heading 1"), ("正文若干。", None),
        ("本章结论", "Heading 1"), ("正文若干。", None),
    ])
    r = fp.verify_docx_structure(p, 2)
    assert r["pass"] is False
    assert r["duplicate_headings"] == ["本章结论"]


def test_d2_7_passes_healthy_docx(tmp_path):
    p = _mk_docx(tmp_path / "d.docx", [
        ("第 1 章：甲", "Heading 1"), ("1.1 节", "Heading 2"), ("本节正文内容。", None),
        ("第 2 章：乙", "Heading 1"), ("2.1 节", "Heading 2"), ("本节正文内容。", None),
    ])
    r = fp.verify_docx_structure(p, 2)
    assert r["pass"] is True
    assert r["h1_count"] == 2


def test_d2_7_catches_chapter_count_mismatch(tmp_path):
    """A5：Heading 1 数量必须 == outline 声明章数（防伪章）。"""
    p = _mk_docx(tmp_path / "e.docx", [
        ("第 1 章：甲", "Heading 1"), ("正文。", None),
    ])
    r = fp.verify_docx_structure(p, 5)
    assert r["pass"] is False
    assert r["h1_count"] == 1 and r["expected"] == 5


def test_d2_7_step_is_skipped_when_no_docx_path_given(tmp_path, monkeypatch):
    """第 7 步为可选步：未传 verify_docx_path 时不执行、不影响 overall_pass。"""
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(_write_outline(tmp_path)),
        source_index_path=str(_write_source_index(tmp_path)),
        output_path=str(tmp_path / "final-report.md"),
    )
    assert result["overall_pass"] is True
    assert "verify_docx" not in result["steps"]


# ── D3-2：provenance sidecar + emit_delivery 命名下沉 ────────


def test_d3_2_provenance_written_on_success(tmp_path, monkeypatch):
    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "empty-log.jsonl"))
    drafts_dir = tmp_path / "drafts"
    drafts_dir.mkdir()
    _write_clean_draft(drafts_dir)
    output = tmp_path / "final-report.md"
    result = fp.run_finalize_pipeline(
        drafts_dir=str(drafts_dir),
        outline_path=str(_write_outline(tmp_path)),
        source_index_path=str(_write_source_index(tmp_path)),
        output_path=str(output),
    )
    assert result["overall_pass"] is True
    sidecar = tmp_path / ".provenance.jsonl"
    assert sidecar.exists(), "成功交付后须写 provenance sidecar"
    rec = json.loads(sidecar.read_text(encoding="utf-8").splitlines()[0])
    assert rec["produced_by"] == "finalize_pipeline.py"
    assert rec["run_id"] == result["run_id"]


def test_d3_2_provenance_is_append_only(tmp_path):
    """sidecar 是 append-only，多次运行累积而非覆盖（同 .degradation-log.jsonl）。"""
    fp.append_provenance(tmp_path, "abc123", [str(tmp_path / "a.docx")], str(tmp_path / "m.md"))
    fp.append_provenance(tmp_path, "def456", [str(tmp_path / "b.docx")], str(tmp_path / "m.md"))
    lines = (tmp_path / ".provenance.jsonl").read_text(encoding="utf-8").strip().splitlines()
    assert len(lines) == 2


def test_d3_2_provenance_rejects_non_hex_run_id(tmp_path):
    """抗伪造：run_id 须符合既有 nonce 原语格式 [0-9a-f]{6,16}。"""
    assert fp.append_provenance(tmp_path, "NOT-A-NONCE!", [], "m.md") is None
    assert not (tmp_path / ".provenance.jsonl").exists()


def test_d3_2_sanitize_filename_stem_truncates_by_gbk_bytes(tmp_path):
    """截断阈值按 GBK 字节而非字符数——中文在 GBK 下占 2 字节。"""
    stem = fp.sanitize_filename_stem("中" * 200)
    assert len(stem.encode("gbk", errors="replace")) <= 120
    # Windows 非法字符被替换
    assert fp.sanitize_filename_stem('a<b>c:d"e/f\\g|h?i*j') == "a_b_c_d_e_f_g_h_i_j"


def test_d3_2_emit_delivery_keeps_docx_and_report_stem_identical(tmp_path):
    """真风险不是中文字符，而是转换报告内硬编码引用 docx 路径——两者 stem 须严格一致。"""
    src_docx = tmp_path / "final-report.docx"
    src_docx.write_bytes(b"PK\x03\x04fake")
    src_report = tmp_path / "final-report.conversion-report.md"
    src_report.write_text(f"转换产物: {src_docx.resolve()}\n文件名: final-report.docx", encoding="utf-8")
    ddir = tmp_path / "output"

    r = fp.emit_delivery(str(src_docx), str(ddir), "空间态势认知智能框架研究", "1.0",
                         conversion_report_src=str(src_report))
    names = sorted(Path(p).name for p in r["delivery_paths"])
    assert names == [
        "空间态势认知智能框架研究_v1.0.conversion-report.md",
        "空间态势认知智能框架研究_v1.0.docx",
    ]
    # 报告内的路径引用已回写为重命名后的文件
    txt = (ddir / "空间态势认知智能框架研究_v1.0.conversion-report.md").read_text(encoding="utf-8")
    assert "空间态势认知智能框架研究_v1.0.docx" in txt
    assert "final-report.docx" not in txt


# ── D3-4：归档机制（首版只报告不移动）────────────────────────


def test_d3_4_never_judges_by_filename_pattern(tmp_path):
    """核心陷阱：按文件名判定会**颠倒**——违规的 SCIF_V1.0.docx 看起来更像正式
    交付物，合规的 final-report.docx 反而不匹配 *_v<版本>*。"""
    ddir = tmp_path / "output"
    ddir.mkdir()
    violating = ddir / "SCIF_V1.0.docx"      # 违规手写产物，事故物证
    violating.write_bytes(b"fake")
    compliant = ddir / "final-report.docx"    # 真正的合规产物
    compliant.write_bytes(b"fake")

    # 只有 compliant 有 provenance 记录
    fp.append_provenance(tmp_path, "abc123def", [str(compliant)], str(tmp_path / "m.md"))

    r = fp.archive_stale_outputs(str(ddir), str(tmp_path), [], archive_stale=False)
    assert str(compliant.resolve()) in r["known_artifacts"]
    assert str(violating.resolve()) in r["unknown_files"], (
        "违规产物必须落在 unknown_files（交人判断），不得因文件名像正式产物而被归档"
    )
    assert r["archived"] == [], "首版只报告不移动"
    assert violating.exists() and compliant.exists(), "不得移动或删除任何文件"


def test_d3_4_archive_stale_flag_still_does_not_move(tmp_path):
    ddir = tmp_path / "output"
    ddir.mkdir()
    f = ddir / "some.docx"
    f.write_bytes(b"fake")
    r = fp.archive_stale_outputs(str(ddir), str(tmp_path), [], archive_stale=True)
    assert r["archived"] == []
    assert f.exists()


def test_d3_4_current_delivery_paths_count_as_known(tmp_path):
    ddir = tmp_path / "output"
    ddir.mkdir()
    f = ddir / "cur.docx"
    f.write_bytes(b"fake")
    r = fp.archive_stale_outputs(str(ddir), str(tmp_path), [str(f)], archive_stale=False)
    assert str(f.resolve()) in r["known_artifacts"]
    assert r["unknown_files"] == []
