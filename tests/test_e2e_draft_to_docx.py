# -*- coding: utf-8 -*-
"""tests/test_e2e_draft_to_docx.py —— 分章草稿 → 合并 → docx 端到端测试（D1-7）。

**存在意义**：故障链上每个组件单独看都符合自己的契约，**从未有任何测试跑过
完整链路**。现有 fixture ``scripts/md2docx/tests/test_fixtures/multi-chapter.md``
是「H2 章标题 → 正文 → H3」结构，**完全不含"两个相邻 H2"这一致命组合**——
这是缺陷从未被测出的直接原因。

本文件的 fixture（``tests/fixtures/e2e-merge/``）**刻意逐字复刻
``references/writer-template.md`` 的骨架**：每份分章草稿首个 H2 逐字为
``## 本章结论``、不写章容器 H2（由合并器生成），防止为让测试变绿而削弱真实性。

Owner 裁决：本文件由 **D1 为唯一 owner**（D3 曾计划建同一路径文件，已裁定
D3 只提供 fixture 不建测试文件，避免同一文件两套用例互相覆盖）。

断言清单（对应 D1 §七 A1-A7 与 D3 §7.4 I1-I7）：
  I1/A4  每个 Heading 1 到下一个 Heading 1 之间必须有非空正文  ← 直接对应用户投诉
  I2/A5  Heading 1 数量 == outline 声明章数（防伪章）
  I3/A6  不得出现文本重复的 Heading 1（防 13 个"本章结论"）
  I4/A2  各 ch{XX}-*.md 内容在合并产物中出现次数**恰为 1**（防重复拼接）
  I6     lookup size > 0 且 CHAPTER 数 == 声明章数（防白名单空转）
  I7/A7  归一化对已合规 outline 为恒等变换（幂等性）
  A3     含 ``## 本章结论`` 的草稿经合并后其 kind == SECTION
"""
from __future__ import annotations

import collections
import re
import subprocess
import sys
from pathlib import Path

import pytest

import merge_drafts as md
from md2docx.assemble.outline_reader import (
    _build_structure_lookup,
    extract_yaml_front_matter,
    normalize_outline_structure,
)
from md2docx.ir import HeadingKind

FIXTURE_DIR = Path(__file__).parent / "fixtures" / "e2e-merge"
OUTLINE = FIXTURE_DIR / "outline.md"
DRAFTS = FIXTURE_DIR / "drafts"
SOURCE_INDEX = FIXTURE_DIR / "source-index.csv"

EXPECTED_CHAPTERS = 2


@pytest.fixture()
def merged_text():
    structure = md.parse_outline_yaml(str(OUTLINE))
    return md.assemble_merged(structure, str(DRAFTS))


# ── fixture 真实性自检（防止测试被"改简单"而失去意义）──────────


def test_fixture_faithfully_reproduces_writer_template_skeleton():
    """fixture 必须真实还原 merge_drafts 的**输入**形态。

    这条用例存在的目的是：若未来有人为了让测试变绿而修改 fixture（比如把
    ``## 本章结论`` 改成 ``### 本章结论``、或给草稿补上章容器 H2），本用例会
    立即失败——真实性不可被静默削弱。
    """
    drafts = sorted(DRAFTS.glob("ch*.md"))
    assert len(drafts) == EXPECTED_CHAPTERS, "fixture 应含多章草稿"
    for f in drafts:
        lines = [ln for ln in f.read_text(encoding="utf-8").split("\n") if ln.strip()]
        h2s = [ln for ln in lines if re.match(r"^##\s+\S", ln)]
        assert h2s, f"{f.name} 应含 H2"
        assert h2s[0].strip() == "## 本章结论", (
            f"{f.name} 的首个 H2 必须逐字为「## 本章结论」（writer-template R1 红线）"
        )
        # 草稿中**不得**自带章容器 H2
        for ln in h2s:
            assert not re.match(r"^##\s+第\s*\d+\s*章", ln), (
                f"{f.name} 不得自带章容器 H2（由合并器生成）"
            )


def test_fixture_contains_the_fatal_adjacent_h2_precondition():
    """现有 md2docx fixture 不含"两个相邻 H2"组合，故缺陷从未被测出。

    本用例断言：合并**前**，草稿首个 H2 与合并器将插入的章容器 H2 会构成
    相邻 H2——即本 fixture 确实覆盖了致命前置条件。
    """
    first = sorted(DRAFTS.glob("ch*.md"))[0]
    body = first.read_text(encoding="utf-8").lstrip()
    assert body.startswith("## 本章结论"), "草稿应以 H2 开头，与章容器 H2 构成相邻组合"


# ── I1/A4：每个 Heading 1 下必须有非空正文（用户投诉本体）─────


def test_i1_no_empty_chapter_after_merge(merged_text):
    """治愈用户投诉的核心断言：章容器下不得紧跟另一个同级标题而无正文。"""
    lines = [ln for ln in merged_text.split("\n") if ln.strip()]
    for a, b in zip(lines, lines[1:]):
        if re.match(r"^##\s+\S", a) and re.match(r"^##\s+\S", b):
            pytest.fail(f"出现相邻 H2（事故形态）: {a!r} 紧跟 {b!r}")


def test_i1_chapter_container_is_followed_by_content(merged_text):
    """章容器 H2 之后到下一个 H2 之间必须有非标题正文。

    作用域限定为**正文章**（`## 第 N 章：`）：``assemble_merged`` 对 appendix
    只输出标题行、不拼接任何内容（附录无对应草稿文件，这是既有设计），故附录
    标题天然"无正文"，不属于本不变量的约束对象。
    """
    blocks = re.split(r"^##\s+(?!#)", merged_text, flags=re.MULTILINE)[1:]
    chapter_blocks = [b for b in blocks if re.match(r"^第\s*\d+\s*章[：:]", b)]
    assert len(chapter_blocks) == EXPECTED_CHAPTERS
    for blk in chapter_blocks:
        body = [
            ln for ln in blk.split("\n")[1:]
            if ln.strip() and not ln.lstrip().startswith("#")
        ]
        assert body, f"章容器下无正文: {blk.split(chr(10))[0]!r}"


# ── I2/A5 + I3/A6：章数与去重 ────────────────────────────────


def test_i2_chapter_container_count_equals_declared(merged_text):
    containers = re.findall(r"^##\s+第\s*\d+\s*章[：:]", merged_text, re.MULTILINE)
    assert len(containers) == EXPECTED_CHAPTERS


def test_i3_no_duplicate_h2_after_merge(merged_text):
    """D1-5 落地后，13 个并列的"本章结论"H2 应全部下沉为 H3。"""
    h2s = re.findall(r"^##\s+(.+?)\s*$", merged_text, re.MULTILINE)
    dup = [t for t, c in collections.Counter(h2s).items() if c > 1]
    assert dup == [], f"出现重复 H2: {dup}"
    assert "本章结论" not in [t.strip() for t in h2s], "「本章结论」不应再是 H2"


def test_a3_conclusion_heading_demoted_to_h3(merged_text):
    """A3：``## 本章结论`` 经合并后应为 H3（节级），且每章各一个。"""
    h3s = [t.strip() for t in re.findall(r"^###\s+(.+?)\s*$", merged_text, re.MULTILINE)]
    assert h3s.count("本章结论") == EXPECTED_CHAPTERS


def test_original_h3_sections_demoted_to_h4(merged_text):
    """草稿内的 H3 正文节应下沉为 H4（docx Heading 3）。"""
    h4s = [t.strip() for t in re.findall(r"^####\s+(.+?)\s*$", merged_text, re.MULTILINE)]
    assert "技术代际划分依据" in h4s
    assert "六层认知升维模型" in h4s


# ── I4/A2：重复拼接（现有用例无一覆盖）──────────────────────


def test_i4_each_draft_appears_exactly_once(merged_text):
    """捕获 D1 §3.4 的重复拼接 bug：sections 声明多条而草稿回落到章级通配符时，
    同一份 ch{XX}-*.md 会被每个 section 各命中一次。"""
    for f in sorted(DRAFTS.glob("ch*.md")):
        body = f.read_text(encoding="utf-8")
        sig = next(
            ln.strip() for ln in body.split("\n")
            if len(ln.strip()) > 40 and not ln.lstrip().startswith(("#", ">"))
        )
        assert merged_text.count(sig) == 1, (
            f"{f.name} 的内容在合并产物中出现 {merged_text.count(sig)} 次，应恰为 1"
        )


# ── I6：白名单空转 ───────────────────────────────────────────


def test_i6_lookup_is_not_empty_and_chapter_count_matches():
    parsed, _ = extract_yaml_front_matter(OUTLINE.read_text(encoding="utf-8"), str(OUTLINE))
    lookup = _build_structure_lookup(parsed["structure"], str(OUTLINE))
    assert len(lookup) > 0, "lookup 不得为空（白名单空转）"
    chapters = [k for k, (kind, _n) in lookup.items() if kind == HeadingKind.CHAPTER]
    assert len(chapters) == EXPECTED_CHAPTERS
    sections = [k for k, (kind, _n) in lookup.items() if kind == HeadingKind.SECTION]
    assert len(sections) == 4, "fixture 声明 4 个节，应全部入表"


# ── I7/A7：归一化幂等 ───────────────────────────────────────


def test_i7_normalize_is_identity_on_compliant_outline():
    """fixture 已用权威键名 → 归一化应为恒等变换（现成的幂等性回归证据）。"""
    parsed, _ = extract_yaml_front_matter(OUTLINE.read_text(encoding="utf-8"), str(OUTLINE))
    structure = parsed["structure"]
    once = normalize_outline_structure(structure)
    assert once["bodymatter"] == structure["bodymatter"]
    assert normalize_outline_structure(once) == once


# ── 全链路：草稿 → finalize_pipeline → docx 回读 ─────────────


def test_full_chain_merge_then_docx_readback(tmp_path, monkeypatch):
    """端到端：合并 → md2docx → python-docx 回读，断言 A4/A5/A6 全部成立。

    这是本方案中唯一一条真正跨越"合并器 + 转换器 + 交付物"三层的用例。
    """
    pytest.importorskip("docx")
    from docx import Document
    import finalize_pipeline as fp

    monkeypatch.setenv("DRR_DEGRADATION_LOG", str(tmp_path / "log.jsonl"))

    # 复制 fixture 到 tmp（管线会就地清洗/改写草稿，不能污染 fixture）
    work_drafts = tmp_path / "drafts"
    work_drafts.mkdir()
    for f in DRAFTS.glob("ch*.md"):
        (work_drafts / f.name).write_text(f.read_text(encoding="utf-8"), encoding="utf-8")
    outline = tmp_path / "outline.md"
    outline.write_text(OUTLINE.read_text(encoding="utf-8"), encoding="utf-8")
    src_index = tmp_path / "source-index.csv"
    src_index.write_text(SOURCE_INDEX.read_text(encoding="utf-8"), encoding="utf-8")

    merged_md = tmp_path / "final-report.md"
    structure = md.parse_outline_yaml(str(outline))
    merged = md.assemble_merged(structure, str(work_drafts))
    # 保留一个 H1 作为报告主标题（md2docx 要求）
    merged_md.write_text("# 端到端合并测试报告\n\n" + merged, encoding="utf-8")

    docx_out = tmp_path / "e2e.docx"
    proc = subprocess.run(
        [sys.executable, "-m", "md2docx", str(merged_md), str(docx_out)],
        cwd=str(Path(md.__file__).resolve().parent),
        capture_output=True, text=True, encoding="utf-8",
    )
    assert docx_out.exists(), f"md2docx 未产出 docx: {(proc.stderr or proc.stdout)[-600:]}"

    # A4/A5/A6 —— 复用 D2-7 的运行时门禁函数，测试与门禁同口径。
    # 期望章数为「正文章 + 附录」：md2docx 把附录也渲染为 Heading 1。
    result = fp.verify_docx_structure(str(docx_out), EXPECTED_CHAPTERS + 1)

    # 附录在 assemble_merged 中只输出标题、无正文（既有设计，附录无草稿文件），
    # 故从"空标题"断言中排除——本用例要守的是**正文章不得为空**这一投诉本体。
    empty_non_appendix = [
        t for t in result["empty_headings"] if "缩略语对照" not in t and "附录" not in t
    ]
    assert empty_non_appendix == [], (
        f"docx 中存在无正文的正文章标题（用户投诉形态）: {empty_non_appendix}"
    )
    assert result["duplicate_headings"] == [], (
        f"docx 中存在重复的章标题: {result['duplicate_headings']}"
    )

    d = Document(str(docx_out))
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    h2 = [p.text for p in d.paragraphs if p.style.name == "Heading 2"]
    # 正文章各一个 Heading 1（附录另计）
    chapter_h1 = [t for t in h1 if "缩略语对照" not in t]
    assert len(chapter_h1) == EXPECTED_CHAPTERS, f"正文章 Heading 1 应为 {EXPECTED_CHAPTERS} 个，实际 {h1}"
    # 「本章结论」下沉后应出现在 Heading 2 层，且不占用 Heading 1
    assert "本章结论" not in h1, "「本章结论」不得是 Heading 1"
    assert h2.count("本章结论") == EXPECTED_CHAPTERS
